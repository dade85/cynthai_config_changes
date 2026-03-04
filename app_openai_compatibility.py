# -*- coding: utf-8 -*-
"""
Engineering Change Intelligence — Streamlit (ECR/ECO + Impact + Risk + Visual Diff)
=================================================================================

What this app DOES (working end-to-end today):
- Ingest engineering documents (PDF, DOCX, TXT, CSV/JSON, images) into a local library
- Extract engineering entities (parts, revisions, requirements-like statements, dimensions-like tokens)
- Build a configuration-aware engineering graph (parts ↔ assemblies ↔ processes ↔ docs ↔ changes)
- Create ECRs/ECOs and auto-generate:
  - Change request definition (structured)
  - Impact assessment across parts/assemblies/process/cost/schedule/compliance (graph-propagated)
  - Risk & DFMEA/PFMEA-style table (rule-based; OpenAI optional)
  - Release checklist + approval workflow states + audit log
  - Acceptance & validation criteria
  - Traceable evidence links to source documents/chunks
- Visual compare:
  - Image ↔ Image diff
  - PDF page (rendered) ↔ PDF page diff (first page by default)
- Predictive/simulation:
  - Simple “delay/quality risk” predictor trained from imported change-history outcomes (optional)
  - What-if sliders to simulate schedule/cost/quality risk

What this app PROVIDES AS PLUG-IN POINTS (stubs you can extend to “ALL OF THE ABOVE” at enterprise grade):
- Teamcenter/PLM connector interface (import/export hooks)
- CAD/JT/STEP rendering integration points (viewer hooks)
- Enterprise identity/roles + e-signature integration points
- Embeddings / vector search (local TF-IDF now; swap in OpenAI embeddings or your own later)

Security:
- API keys are kept in Streamlit session_state unless you explicitly export env vars.
- For production, deploy in VPC/on-prem and use a secret manager.

Run:
  pip install -r requirements.txt
  streamlit run app.py
"""

from __future__ import annotations

import os
import io
import re
import json
import time
import math
import base64
import gzip
import difflib
import uuid
import textwrap
import hashlib
import sqlite3
import threading
import subprocess
import tempfile
import zipfile
import datetime as dt
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st

try:
    import requests
    REQUESTS_AVAILABLE = True
except Exception:
    requests = None
    REQUESTS_AVAILABLE = False

import streamlit.components.v1 as components

import networkx as nx
import plotly.graph_objects as go
import plotly.express as px

# Optional 3D/CAD/BIM analyzers (handled gracefully)
try:
    import trimesh  # type: ignore
    TRIMESH_AVAILABLE = True
except Exception:
    trimesh = None
    TRIMESH_AVAILABLE = False

try:
    import ifcopenshell  # type: ignore
    IFCOPEN_AVAILABLE = True
except Exception:
    ifcopenshell = None
    IFCOPEN_AVAILABLE = False


# Optional deps (handled gracefully)
try:
    import pdfplumber  # type: ignore
except Exception:
    pdfplumber = None

try:
    import fitz  # PyMuPDF  # type: ignore
except Exception:
    fitz = None

try:
    from PIL import Image, ImageChops, ImageEnhance, ImageDraw, ImageFont  # type: ignore
except Exception:
    Image = None  # type: ignore

try:
    import cv2  # type: ignore
    CV2_AVAILABLE = True
except Exception:
    cv2 = None  # type: ignore
    CV2_AVAILABLE = False

try:
    from skimage.metrics import structural_similarity as skimage_ssim  # type: ignore
    SKIMAGE_AVAILABLE = True
except Exception:
    skimage_ssim = None  # type: ignore
    SKIMAGE_AVAILABLE = False


try:
    import pytesseract  # type: ignore
    PYTESSERACT_AVAILABLE = True
except Exception:
    pytesseract = None  # type: ignore
    PYTESSERACT_AVAILABLE = False


try:
    from docx import Document  # python-docx  # type: ignore
except Exception:
    Document = None  # type: ignore

try:
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
except Exception:
    A4 = None  # type: ignore
    canvas = None  # type: ignore

try:
    from sklearn.feature_extraction.text import TfidfVectorizer  # type: ignore
    from sklearn.linear_model import LogisticRegression  # type: ignore
    from sklearn.model_selection import train_test_split  # type: ignore
    from sklearn.metrics import roc_auc_score  # type: ignore
except Exception:
    TfidfVectorizer = None  # type: ignore
    LogisticRegression = None  # type: ignore
    train_test_split = None  # type: ignore
    roc_auc_score = None  # type: ignore

# OpenAI (optional)
OPENAI_AVAILABLE = False
try:
    from openai import OpenAI  # type: ignore
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False


# ----------------------------
# Configuration
# ----------------------------
APP_NAME = "CynthAI© Engineering Change Copilot"

# Streamlit page config (must be the first Streamlit command)
st.set_page_config(page_title=APP_NAME, layout="wide")

DATA_DIR = os.path.join(os.path.dirname(__file__), "ei_data")
FILES_DIR = os.path.join(DATA_DIR, "files")
DB_PATH = os.path.join(DATA_DIR, "ei.db")

DB_WRITE_LOCK = threading.Lock()

EXPORTS_DIR = os.path.join(DATA_DIR, "exports")
ARCHIVES_DIR = os.path.join(DATA_DIR, "archives")

# Supported CAD uploads (stored as evidence; optional rendering via connectors)
CAD_FILE_EXTS = ["step","stp","iges","igs","jt","stl","obj","dwg","dxf","sldprt","sldasm","ply","glb","gltf"]

DEFAULT_MODEL = "gpt-4o-mini"  # if you have access; otherwise use a model available in your org

# Mistral (optional) — text LLM routing
MISTRAL_CHAT_URL = "https://api.mistral.ai/v1/chat/completions"
DEFAULT_MISTRAL_MODEL = "mistral-small-latest"



# ----------------------------
# Utilities
# ----------------------------
def ensure_dirs() -> None:
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(FILES_DIR, exist_ok=True)
    os.makedirs(EXPORTS_DIR, exist_ok=True)
    os.makedirs(ARCHIVES_DIR, exist_ok=True)


def now_iso() -> str:
    return dt.datetime.now().replace(microsecond=0).isoformat()


def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return h.hexdigest()


def safe_int(x: Any, default: int = 0) -> int:
    try:
        return int(float(x))
    except Exception:
        return default


def safe_float(x: Any, default: float = 0.0) -> float:
    try:
        return float(x)
    except Exception:
        return default


# -----------------------------
# Robust tabular ingestion
# -----------------------------
def read_tabular_file(uploaded_file) -> pd.DataFrame:
    """
    Bullet-proof CSV/XLSX reader:
    - auto-detect delimiter (comma/semicolon/tab)
    - handles UTF-8 and latin-1 fallbacks
    - supports Excel (xlsx)
    """
    if uploaded_file is None:
        return pd.DataFrame()

    name = getattr(uploaded_file, "name", "") or ""
    lower = name.lower()

    # Streamlit UploadedFile behaves like BytesIO; keep pointer safe
    def _reset():
        try:
            uploaded_file.seek(0)
        except Exception:
            pass

    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        _reset()
        try:
            return pd.read_excel(uploaded_file)
        except Exception:
            _reset()
            return pd.read_excel(uploaded_file, engine="openpyxl")

    # CSV / TXT
    _reset()
    raw = uploaded_file.read()
    if isinstance(raw, str):
        raw_bytes = raw.encode("utf-8", errors="ignore")
    else:
        raw_bytes = raw or b""

    # Try decodings
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = raw_bytes.decode(enc)
            break
        except Exception:
            text = None
    if text is None:
        text = raw_bytes.decode("utf-8", errors="ignore")

    # Use pandas with python engine + delimiter sniffing
    from io import StringIO
    buf = StringIO(text)

    try:
        df = pd.read_csv(buf, engine="python")
        if df.shape[1] <= 1:
            raise ValueError("likely wrong delimiter")
        return df
    except Exception:
        buf = StringIO(text)
        return pd.read_csv(buf, sep=None, engine="python", on_bad_lines="skip")


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip().lower().replace(" ", "_").replace("-", "_") for c in df.columns]
    return df

def coerce_numeric_series(s: pd.Series) -> pd.Series:
    """
    Coerce numbers robustly from strings:
    - Handles € and EUR suffix/prefix
    - Handles EU style 1.234,56 and US style 1,234.56
    - Strips spaces and non-numeric symbols (keeps . , - +)
    """
    if s is None:
        return pd.Series(dtype="float64")
    ss = s.astype(str).str.replace("\u00a0", " ", regex=False)
    ss = ss.str.replace("€", "", regex=False).str.replace("EUR", "", regex=False).str.replace("euro", "", regex=False)
    ss = ss.str.replace(r"[^\d,\.\-\+]", "", regex=True)

    eu_like = ss.str.contains(r"\d+[\.\s]\d{3},\d{1,2}$", regex=True)
    ss_eu = ss.copy()
    ss_eu[eu_like] = ss_eu[eu_like].str.replace(".", "", regex=False).str.replace(",", ".", regex=False)

    us_like = ss.str.contains(r"\d+,\d{3}(\.\d{1,2})?$", regex=True)
    ss_us = ss.copy()
    ss_us[us_like] = ss_us[us_like].str.replace(",", "", regex=False)

    out = pd.to_numeric(ss_us, errors="coerce")
    out[eu_like] = pd.to_numeric(ss_eu[eu_like], errors="coerce")
    return out



def map_bom_schema(df: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, str]]:
    """
    Returns (normalized_df, colmap) where colmap maps canonical fields to actual columns in df.
    Supports:
      - edge-list: parent_part_number, child_part_number, quantity, uom, effectivity, notes
      - rich BOM: parent_part_number + part_number + quantity + cost/lead/weight/etc
    """
    df = normalize_columns(df)

    # canonical -> candidate columns in priority order
    candidates = {
        "parent_part_number": ["parent_part_number", "parent_part_no", "parent_pn", "parent", "parent_part"],
        "child_part_number": ["child_part_number", "component_part_number", "part_number", "child_part_no", "child_pn", "child"],
        "part_number": ["part_number", "child_part_number", "component_part_number"],
        "serial_number": ["serial_number", "serial_no", "serial", "sn"],
        "bom_description": ["bom_description", "description", "part_description", "item_description"],
        "bom_level": ["bom_level", "level", "lvl"],
        "quantity": ["quantity", "qty", "component_qty", "bom_quantity"],
        "uom": ["uom", "unit", "unit_of_measure"],
        "unit_cost_eur": ["unit_cost_eur","unit_cost","cost_unit_eur","unit_cost_(euro)","unit_cost_(eur)","unit_cost_euro","unitprice","unit_price"],
        "extended_cost_eur": ["extended_cost_eur","extended_cost_euro","extended_cost_(euro)","extended_cost_(eur)","extended_cost","total_cost_eur","total_cost_euro","ext_cost_eur","ext_cost_euro","ext_cost"],
        "lead_time_days": ["lead_time_days", "lead_time", "lead_days", "lt_days"],
        "eco_id": ["eco_id", "ecr_id", "eco", "change_id"],
        "weight_kg": ["weight_kg", "weight", "mass_kg", "mass"],
        "effectivity_date": ["effectivity_date", "effectivity", "effective_from", "valid_from", "date_effective"],
        "effective_from": ["effective_from", "effectivity_from", "valid_from"],
        "effective_to": ["effective_to", "effectivity_to", "valid_to"],
        "material": ["material", "material_spec"],
        "process": ["process", "process_route", "routing", "work_center", "process_step"],
        "supplier": ["supplier", "supplier_name"],
        "supplier_part_no": ["supplier_part_no", "supplier_part_no_", "supplier_part_number", "supplier_part_no", "supplier_part_no"],
        "revision": ["revision", "parent_revision", "parent_rev", "rev"],
        "lifecycle_state": ["lifecycle_state", "lifecycle_status", "lifecycle", "status"],
        "change_reason": ["change_reason", "reason", "change_reason_code"],
        "compliance_tags": ["compliance_tags", "compliance_reach", "reach", "rohs", "compliance"],
        "notes": ["notes", "note", "remarks", "comment"],
        "line_no": ["line_no", "line", "item_no", "position"],
        "make_buy": ["make_buy", "make_or_buy", "procurement_type"],
        "manufacturer_name": ["manufacturer_name", "mfr_name"],
        "manufacturer_part_no": ["manufacturer_part_no", "mfr_part_no", "manufacturer_part_number"],
        "created_by": ["created_by", "author", "owner"],
        "creation_time": ["creation_time", "created_at", "created_on"],
        "dimensions_mm": ["dimensions_mm", "dims_mm", "dimensions"],
        "tolerance_mm": ["tolerance_mm", "tol_mm", "tolerance"],
        "criticality": ["criticality", "crit", "safety_criticality"],
        "effectivity_code": ["effectivity_code", "eff_code"],
        "compliance_reach": ["compliance_reach", "reach_compliance", "reach"],
        "finish_spec": ["finish_spec", "finish"],
        "process_route": ["process_route", "routing", "route"],
        "work_center": ["work_center", "workcentre"],
    }

    colmap: Dict[str, str] = {}
    for canon, opts in candidates.items():
        for o in opts:
            if o in df.columns:
                colmap[canon] = o
                break

    return df, colmap


# -----------------------------
# 3D model ingestion & analysis
# -----------------------------
THREED_EXTS = {"stl","obj","ply","glb","gltf","stp","step","igs","iges"}

def is_3d_filename(name: str) -> bool:
    if not name:
        return False
    ext = name.lower().split(".")[-1]
    return ext in THREED_EXTS



# -----------------------------
# FreeCADCmd fallback for STEP/IGES metrics
# -----------------------------
FREECAD_METRICS_SCRIPT = r"""
import json, sys

try:
    import FreeCAD
    import Part
except Exception as e:
    print("JSON_METRICS:" + json.dumps({"ok": False, "error": str(e)}))
    sys.exit(0)

fn = sys.argv[1]

try:
    doc = FreeCAD.newDocument("Doc")
except Exception:
    doc = None

shape = None

# Try direct Shape.read
try:
    s = Part.Shape()
    s.read(fn)
    shape = s
except Exception:
    pass

# Fallback insert into doc
if shape is None:
    try:
        if doc is None:
            doc = FreeCAD.newDocument("Doc")
        Part.insert(fn, doc.Name)
        for o in getattr(doc, "Objects", []):
            if hasattr(o, "Shape"):
                shape = o.Shape
                break
    except Exception as e:
        print("JSON_METRICS:" + json.dumps({"ok": False, "error": "import_failed", "detail": str(e)}))
        sys.exit(0)

if shape is None:
    print("JSON_METRICS:" + json.dumps({"ok": False, "error": "no_shape"}))
    sys.exit(0)

try:
    bb = shape.BoundBox
    out = {
        "ok": True,
        "bbox_dx": float(bb.XLength),
        "bbox_dy": float(bb.YLength),
        "bbox_dz": float(bb.ZLength),
        "volume": float(getattr(shape, "Volume", None)) if hasattr(shape, "Volume") else None,
        "surface_area": float(getattr(shape, "Area", None)) if hasattr(shape, "Area") else None,
    }
    try:
        com = shape.CenterOfMass
        out["com_x"] = float(com.x)
        out["com_y"] = float(com.y)
        out["com_z"] = float(com.z)
    except Exception:
        pass
    print("JSON_METRICS:" + json.dumps(out))
except Exception as e:
    print("JSON_METRICS:" + json.dumps({"ok": False, "error": str(e)}))
"""


def _find_freecadcmd_candidates() -> List[str]:
    cands: List[str] = []
    ui_path = str(st.session_state.get("freecad_cmd_path", "") or "").strip()
    if ui_path:
        cands.append(ui_path)
    env_path = os.environ.get("FREECADCMD_PATH", "").strip()
    if env_path:
        cands.append(env_path)

    # Common locations (Windows paths expressed with forward slashes to avoid escaping issues)
    common = [
        "C:/Users/d.adewunmi/AppData/Local/Programs/FreeCAD 1.0/bin/FreeCADCmd.exe",
        "C:/Program Files/FreeCAD 1.0/bin/FreeCADCmd.exe",
        "C:/Program Files/FreeCAD 0.21/bin/FreeCADCmd.exe",
    ]
    cands.extend(common)

    # Also try FreeCADCmd in PATH
    cands.append("FreeCADCmd")
    cands.append("freecadcmd")

    # De-dup
    seen=set()
    out=[]
    for c in cands:
        if c and c not in seen:
            seen.add(c)
            out.append(c)
    return out


def _resolve_freecadcmd() -> Optional[str]:
    for c in _find_freecadcmd_candidates():
        if c in {"FreeCADCmd","freecadcmd"}:
            return c
        try:
            if os.path.exists(c):
                return c
        except Exception:
            continue
    return None


def _freecad_step_metrics(file_bytes: bytes, filename: str) -> Tuple[Dict[str, Any], Optional[str]]:
    """Return STEP/IGES geometry metrics using FreeCADCmd when available.

    Returns (metrics, error). If FreeCADCmd isn't available, returns status with hint.
    """
    metrics: Dict[str, Any] = {"filename": filename}
    freecad_cmd = _resolve_freecadcmd()
    if not freecad_cmd:
        metrics["status"] = "freecad_unavailable"
        metrics["hint"] = "FreeCADCmd not found. Set a valid path in the Geometry/BIM page (FreeCADCmd.exe) or set env FREECADCMD_PATH."
        return metrics, "freecad_unavailable"

    ext = (filename.lower().split(".")[-1] if filename else "").lower()
    if ext not in {"stp","step","igs","iges"}:
        metrics["status"] = "unsupported_for_freecad"
        return metrics, "unsupported"

    try:
        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / ("model." + ext)
            inp.write_bytes(file_bytes)
            script = Path(td) / "metrics_freecad.py"
            script.write_text(FREECAD_METRICS_SCRIPT, encoding="utf-8")

            cmd = [freecad_cmd, str(script), str(inp)]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            out = (proc.stdout or "") + "\n" + (proc.stderr or "")
            # Parse last JSON_METRICS line
            line = None
            for l in out.splitlines()[::-1]:
                if l.startswith("JSON_METRICS:"):
                    line = l
                    break
            if not line:
                metrics["status"] = "freecad_parse_failed"
                metrics["raw"] = out[-2000:]
                return metrics, "parse_failed"

            payload = line.split("JSON_METRICS:",1)[1].strip()
            data = json.loads(payload)
            if not data.get("ok"):
                metrics["status"] = "freecad_error"
                metrics["error"] = data.get("error")
                metrics["detail"] = data.get("detail")
                return metrics, "freecad_error"

            metrics.update({k:v for k,v in data.items() if k != "ok"})
            metrics["status"] = "ok_freecad"
            metrics["backend"] = "FreeCADCmd"
            return metrics, None
    except Exception as e:
        metrics["status"] = "freecad_exception"
        metrics["error"] = str(e)
        return metrics, str(e)

def analyze_3d_model_bytes(file_bytes: bytes, filename: str) -> Tuple[Dict[str, Any], Any]:
    """Analyze a 3D model from bytes.

    Returns (metrics, plotly_fig_or_None).
    - Uses trimesh for mesh-like formats (STL/OBJ/PLY/GLB/GLTF).
    - For STEP/IGES: tries trimesh loaders; if unavailable, returns a clear status.

    NOTE: This is a demo-grade geometry proxy, not a full CAD feature tree.
    """
    metrics: Dict[str, Any] = {"filename": filename}
    ext = (filename.lower().split(".")[-1] if filename else "").lower()

    if (not TRIMESH_AVAILABLE) or go is None:
        # FreeCADCmd fallback for STEP/IGES if available
        if ext in {"stp","step","igs","iges"}:
            m2, err = _freecad_step_metrics(file_bytes if isinstance(file_bytes,(bytes,bytearray)) else bytes(file_bytes), filename)
            if m2.get("status","").startswith("ok"): 
                return m2, None
            # If FreeCAD unavailable, fall through to informative status
            metrics.update({k:v for k,v in m2.items() if k not in {"filename"}})
        metrics["status"] = "trimesh_unavailable"
        metrics["hint"] = "Install 'trimesh' to enable mesh previews/metrics. For STEP/IGES, configure FreeCADCmd path for exact KPIs."
        return metrics, None

    # Attempt load
    mesh = None
    try:
        bio = io.BytesIO(file_bytes if isinstance(file_bytes, (bytes, bytearray)) else bytes(file_bytes))
        # Provide file_type hint for reliability
        file_type = ext if ext else None
        mesh = trimesh.load(bio, file_type=file_type, force='mesh')
    except Exception as e:
        mesh = None
        metrics["status"] = "load_failed"
        metrics["error"] = str(e)

    if mesh is None:
        # For STEP/IGES, attempt FreeCADCmd exact KPIs
        if ext in {"stp","step","igs","iges"}:
            m2, err = _freecad_step_metrics(file_bytes if isinstance(file_bytes,(bytes,bytearray)) else bytes(file_bytes), filename)
            if m2.get("status","").startswith("ok"): 
                return m2, None
            # Provide best available status/hint
            metrics.update({k:v for k,v in m2.items() if k not in {"filename"}})
            if metrics.get("status") in {"freecad_unavailable","freecad_parse_failed","freecad_error","freecad_exception"}:
                return metrics, None
            metrics["status"] = "cad_kernel_required"
            metrics["hint"] = "STEP/IGES needs CAD kernel. Configure FreeCADCmd path for exact KPIs."
            return metrics, None
        return metrics, None

    # Some loaders return a Scene
    try:
        if hasattr(mesh, 'geometry') and isinstance(getattr(mesh, 'geometry', None), dict):
            geoms = list(mesh.geometry.values())
            if geoms:
                mesh = trimesh.util.concatenate([g for g in geoms if g is not None])
    except Exception:
        pass

    # Ensure we have vertices/faces
    try:
        vertices = getattr(mesh, 'vertices', None)
        faces = getattr(mesh, 'faces', None)
        if vertices is None or len(vertices) == 0:
            metrics["status"] = "no_vertices"
            return metrics, None

        # Basic metrics
        bounds = mesh.bounds if hasattr(mesh, 'bounds') else None
        if bounds is not None and len(bounds) == 2:
            extents = (bounds[1] - bounds[0]).tolist()
            metrics["bbox_dx"] = float(extents[0])
            metrics["bbox_dy"] = float(extents[1])
            metrics["bbox_dz"] = float(extents[2])

        metrics["n_vertices"] = int(len(vertices))
        metrics["n_faces"] = int(len(faces)) if faces is not None else 0

        # Surface/volume proxies
        try:
            metrics["surface_area"] = float(mesh.area) if hasattr(mesh, 'area') else None
        except Exception:
            metrics["surface_area"] = None
        try:
            metrics["volume"] = float(mesh.volume) if hasattr(mesh, 'volume') else None
        except Exception:
            metrics["volume"] = None

        metrics["status"] = "ok"

        # Build preview figure (cap size for responsiveness)
        fig = None
        try:
            max_faces = 20000
            if faces is not None and len(faces) and len(faces) <= max_faces:
                v = vertices
                f = faces
                fig = go.Figure(data=[go.Mesh3d(
                    x=v[:,0], y=v[:,1], z=v[:,2],
                    i=f[:,0], j=f[:,1], k=f[:,2],
                    opacity=0.8
                )])
                fig.update_layout(margin=dict(l=0,r=0,t=0,b=0))
            else:
                # Fallback to point cloud
                v = vertices
                step = max(1, int(len(v) / 8000))
                vv = v[::step]
                fig = go.Figure(data=[go.Scatter3d(x=vv[:,0], y=vv[:,1], z=vv[:,2], mode='markers', marker=dict(size=1))])
                fig.update_layout(margin=dict(l=0,r=0,t=0,b=0))
        except Exception as e:
            metrics["preview_error"] = str(e)
            fig = None

        return metrics, fig

    except Exception as e:
        metrics["status"] = "analysis_failed"
        metrics["error"] = str(e)
        return metrics, None


# ----------------------------
# Unified Engineering Data Model (UEDM) + Snapshots + BIM analysis
# ----------------------------

BIM_EXTS = {"ifc"}

def is_bim_filename(name: str) -> bool:
    if not name:
        return False
    ext = name.lower().split(".")[-1]
    return ext in BIM_EXTS


try:
    import ifcopenshell  # type: ignore
except Exception:
    ifcopenshell = None  # type: ignore


def analyze_ifc_model_bytes(file_bytes: bytes, filename: str) -> Tuple[Dict[str, Any], Any]:
    """Return basic IFC BIM metrics. Geometry extraction depends on ifcopenshell build.

    For demo: entity counts, storeys/spaces/products. Always degrades gracefully.
    """
    metrics: Dict[str, Any] = {"filename": filename}
    if ifcopenshell is None:
        metrics["status"] = "ifcopenshell_unavailable"
        return metrics, None

    try:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".ifc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        model = ifcopenshell.open(tmp_path)
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

        def cnt(t: str) -> int:
            try:
                return len(model.by_type(t))
            except Exception:
                return 0

        metrics.update({
            "status": "ok",
            "ifc_products": cnt("IfcProduct"),
            "ifc_storeys": cnt("IfcBuildingStorey"),
            "ifc_spaces": cnt("IfcSpace"),
            "ifc_walls": cnt("IfcWall") + cnt("IfcWallStandardCase"),
            "ifc_slabs": cnt("IfcSlab"),
            "ifc_beams": cnt("IfcBeam"),
            "ifc_columns": cnt("IfcColumn"),
        })
        return metrics, None

    except Exception as e:
        metrics["status"] = "ifc_parse_failed"
        metrics["error"] = str(e)
        return metrics, None


def _ensure_uedm_state() -> None:
    if "uedm" not in st.session_state:
        st.session_state["uedm"] = {
            "items": pd.DataFrame(),
            "relationships": pd.DataFrame(),
            "attributes": pd.DataFrame(),
            "events": pd.DataFrame(),
            "sources": {},
        }


def _uedm_pack(uedm: Dict[str, Any]) -> str:
    payload = {
        "items": uedm.get("items", pd.DataFrame()).to_dict("records"),
        "relationships": uedm.get("relationships", pd.DataFrame()).to_dict("records"),
        "attributes": uedm.get("attributes", pd.DataFrame()).to_dict("records"),
        "events": uedm.get("events", pd.DataFrame()).to_dict("records"),
        "sources": uedm.get("sources", {}),
    }
    raw = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    gz = gzip.compress(raw)
    return base64.b64encode(gz).decode("ascii")


def _uedm_unpack(b64: str) -> Dict[str, Any]:
    raw = gzip.decompress(base64.b64decode(b64.encode("ascii")))
    payload = json.loads(raw.decode("utf-8"))
    return {
        "items": pd.DataFrame(payload.get("items", [])),
        "relationships": pd.DataFrame(payload.get("relationships", [])),
        "attributes": pd.DataFrame(payload.get("attributes", [])),
        "events": pd.DataFrame(payload.get("events", [])),
        "sources": payload.get("sources", {}),
    }


def init_snapshot_tables() -> None:
    con = sqlite3.connect(DB_PATH, check_same_thread=False)
    cur = con.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS uedm_snapshots (
            snapshot_id TEXT PRIMARY KEY,
            name TEXT,
            created_at TEXT,
            payload_b64 TEXT
        )
        """
    )
    con.commit()
    con.close()


def save_uedm_snapshot(name: str, uedm: Dict[str, Any]) -> str:
    init_snapshot_tables()
    snapshot_id = f"snap_{uuid.uuid4().hex[:12]}"
    payload_b64 = _uedm_pack(uedm)

    def _write(cur):
        cur.execute(
            "INSERT OR REPLACE INTO uedm_snapshots(snapshot_id, name, created_at, payload_b64) VALUES (?,?,?,?)",
            (snapshot_id, name, now_iso(), payload_b64),
        )

    _db_write_retry(lambda cur: _write(cur))
    return snapshot_id


def list_uedm_snapshots() -> pd.DataFrame:
    init_snapshot_tables()
    con = sqlite3.connect(DB_PATH, check_same_thread=False)
    df = pd.read_sql_query("SELECT snapshot_id, name, created_at FROM uedm_snapshots ORDER BY created_at DESC", con)
    con.close()
    return df


def load_uedm_snapshot(snapshot_id: str) -> Optional[Dict[str, Any]]:
    init_snapshot_tables()
    con = sqlite3.connect(DB_PATH, check_same_thread=False)
    cur = con.cursor()
    cur.execute("SELECT payload_b64 FROM uedm_snapshots WHERE snapshot_id=?", (snapshot_id,))
    row = cur.fetchone()
    con.close()
    if not row:
        return None
    return _uedm_unpack(row[0])


def delete_uedm_snapshot(snapshot_id: str) -> None:
    init_snapshot_tables()

    def _write(cur):
        cur.execute("DELETE FROM uedm_snapshots WHERE snapshot_id=?", (snapshot_id,))

    _db_write_retry(lambda cur: _write(cur))


def build_uedm_from_current_sources() -> Dict[str, Any]:
    """Build a unified, queryable engineering model from what is currently loaded in the app."""
    _ensure_uedm_state()

    items: List[Dict[str, Any]] = []
    rels: List[Dict[str, Any]] = []
    attrs: List[Dict[str, Any]] = []
    events: List[Dict[str, Any]] = []
    sources: Dict[str, Any] = {}

    # Documents (includes drawings, procedures, CAD/BIM files)
    try:
        docs_df = list_documents()
    except Exception:
        docs_df = pd.DataFrame()

    sources["documents"] = int(len(docs_df))
    for _, r in docs_df.iterrows():
        doc_id = str(r.get("doc_id"))
        filename = str(r.get("filename"))
        ext = str(r.get("file_ext") or "").lower()
        items.append({
            "item_id": f"doc:{doc_id}",
            "item_type": "document",
            "name": filename,
            "source": "library",
            "ref_id": doc_id,
        })
        attrs.append({"item_id": f"doc:{doc_id}", "key": "file_ext", "value": ext, "value_type": "str"})
        attrs.append({"item_id": f"doc:{doc_id}", "key": "uploaded_at", "value": str(r.get("uploaded_at")), "value_type": "str"})

    # BOM dashboard dataframe (preferred for rich schema)
    bom_df = st.session_state.get("bom_dashboard_df", pd.DataFrame())
    if isinstance(bom_df, pd.DataFrame) and len(bom_df) > 0:
        sources["bom_rows"] = int(len(bom_df))
        # ensure canonical names exist
        if "parent_part_number" not in bom_df.columns and "parent_part_id" in bom_df.columns:
            bom_df = bom_df.rename(columns={"parent_part_id": "parent_part_number"})
        if "part_number" not in bom_df.columns and "child_part_id" in bom_df.columns:
            bom_df = bom_df.rename(columns={"child_part_id": "part_number"})

        for _, r in bom_df.iterrows():
            p = str(r.get("parent_part_number", "")).strip()
            c = str(r.get("part_number", "")).strip()
            if p:
                items.append({"item_id": f"part:{p}", "item_type": "part", "name": p, "source": "bom"})
            if c:
                items.append({"item_id": f"part:{c}", "item_type": "part", "name": c, "source": "bom"})
            if p and c:
                rels.append({"src": f"part:{p}", "dst": f"part:{c}", "rel_type": "contains", "quantity": r.get("quantity", r.get("bom_quantity", None)), "uom": r.get("uom", "")})
            # attributes (rich schema)
            for col, val in r.items():
                if col in {"parent_part_number","part_number"}:
                    continue
                if val is None or (isinstance(val, float) and np.isnan(val)):
                    continue
                vtype = "num" if isinstance(val, (int,float,np.integer,np.floating)) else "str"
                tgt = f"part:{c}" if c else (f"part:{p}" if p else None)
                if tgt:
                    attrs.append({"item_id": tgt, "key": str(col), "value": val, "value_type": vtype})

    

    # Extra BOM sources (EBOM/MBOM/etc.) uploaded in UEDM hub
    extra = st.session_state.get("uedm_extra_boms", [])
    if isinstance(extra, list) and extra:
        sources["extra_bom_sources"] = len(extra)
        for src in extra:
            sdf = src.get("df")
            stype = src.get("type", "Other")
            sname = src.get("name", "uploaded")
            if not isinstance(sdf, pd.DataFrame) or sdf.empty:
                continue
            try:
                _df, colmap = map_bom_schema(sdf.copy())
            except Exception:
                _df, colmap = sdf.copy(), {}
            # resolve columns
            pcol = colmap.get("parent_part_number") or ("parent_part_number" if "parent_part_number" in _df.columns else None)
            ccol = colmap.get("part_number") or ("part_number" if "part_number" in _df.columns else None)
            qcol = colmap.get("quantity") or ("quantity" if "quantity" in _df.columns else None)
            if not pcol or not ccol:
                continue
            for _, r in _df.iterrows():
                p = str(r.get(pcol, "")).strip()
                c = str(r.get(ccol, "")).strip()
                if p:
                    items.append({"item_id": f"part:{p}", "item_type": "part", "name": p, "source": stype})
                if c:
                    items.append({"item_id": f"part:{c}", "item_type": "part", "name": c, "source": stype})
                if p and c:
                    rels.append({"src": f"part:{p}", "dst": f"part:{c}", "rel_type": f"contains_{stype.lower().replace(' ','_')}", "quantity": r.get(qcol, None) if qcol else None, "uom": r.get(colmap.get("uom","uom"), "")})
                # keep a minimal trace to source
                if c:
                    attrs.append({"item_id": f"part:{c}", "key": "source_bom", "value": sname, "value_type": "str"})

    # If no bom_df in session, fall back to sqlite bom links
    if (not isinstance(bom_df, pd.DataFrame)) or len(bom_df) == 0:
        try:
            con = sqlite3.connect(DB_PATH, check_same_thread=False)
            df_links = pd.read_sql_query("SELECT parent_part_id, child_part_id, quantity, uom, effectivity, notes FROM bom", con)
            con.close()
            sources["bom_links"] = int(len(df_links))
            for _, r in df_links.iterrows():
                p = str(r.get("parent_part_id", "")).strip()
                c = str(r.get("child_part_id", "")).strip()
                if p:
                    items.append({"item_id": f"part:{p}", "item_type": "part", "name": p, "source": "bom"})
                if c:
                    items.append({"item_id": f"part:{c}", "item_type": "part", "name": c, "source": "bom"})
                if p and c:
                    rels.append({"src": f"part:{p}", "dst": f"part:{c}", "rel_type": "contains", "quantity": r.get("quantity", None), "uom": r.get("uom", ""), "effectivity": r.get("effectivity", "")})
        except Exception:
            pass

    # CAD/BIM quick metrics (from documents table)
    # Store results as attributes on document items
    try:
        docs_full = pd.read_sql_query("SELECT doc_id, filename, stored_path FROM documents", sqlite3.connect(DB_PATH, check_same_thread=False))
    except Exception:
        docs_full = pd.DataFrame()

    cad_metrics_count = 0
    for _, r in docs_full.iterrows():
        doc_id = str(r.get("doc_id"))
        fn = str(r.get("filename"))
        pth = str(r.get("stored_path"))
        if not pth or not os.path.exists(pth):
            continue
        ext = fn.lower().split(".")[-1]
        if ext in THREED_EXTS or ext in BIM_EXTS:
            try:
                b = open(pth, "rb").read()
            except Exception:
                continue
            if ext in THREED_EXTS:
                m, _fig = analyze_3d_model_bytes(b, fn)
            else:
                m, _fig = analyze_ifc_model_bytes(b, fn)
            cad_metrics_count += 1
            for k, v in (m or {}).items():
                if k in {"filename"}:
                    continue
                vtype = "num" if isinstance(v, (int,float,np.integer,np.floating)) else "str"
                attrs.append({"item_id": f"doc:{doc_id}", "key": f"model_{k}", "value": v, "value_type": vtype})

    sources["cad_bim_models_analyzed"] = cad_metrics_count

    # Deduplicate items by item_id
    if items:
        df_items = pd.DataFrame(items).drop_duplicates(subset=["item_id"], keep="last")
    else:
        df_items = pd.DataFrame(columns=["item_id","item_type","name","source","ref_id"])
    df_rels = pd.DataFrame(rels) if rels else pd.DataFrame(columns=["src","dst","rel_type","quantity","uom","effectivity"])
    df_attrs = pd.DataFrame(attrs) if attrs else pd.DataFrame(columns=["item_id","key","value","value_type"])
    df_events = pd.DataFrame(events) if events else pd.DataFrame(columns=["event_id","event_type","ref","created_at"])

    uedm = {"items": df_items, "relationships": df_rels, "attributes": df_attrs, "events": df_events, "sources": sources}
    st.session_state["uedm"] = uedm
    return uedm


def uedm_kpi_summary(uedm: Dict[str, Any]) -> Dict[str, Any]:
    items = uedm.get("items", pd.DataFrame())
    rels = uedm.get("relationships", pd.DataFrame())
    attrs = uedm.get("attributes", pd.DataFrame())

    def _count_where(df: pd.DataFrame, col: str, val: str) -> int:
        if col not in df.columns:
            return 0
        return int((df[col] == val).sum())

    summary = {
        "items_total": int(len(items)),
        "parts_total": _count_where(items, "item_type", "part"),
        "documents_total": _count_where(items, "item_type", "document"),
        "relationships_total": int(len(rels)),
        "attributes_total": int(len(attrs)),
    }

    # numeric attributes rollups
    if not attrs.empty and "value_type" in attrs.columns:
        num = attrs[attrs["value_type"] == "num"].copy()
        # try to coerce
        try:
            num["value_num"] = pd.to_numeric(num["value"], errors="coerce")
        except Exception:
            num["value_num"] = np.nan
        summary["numeric_attr_count"] = int(num["value_num"].notna().sum())

    return summary


def compare_uedm_snapshots(a: Dict[str, Any], b: Dict[str, Any]) -> Dict[str, Any]:
    """Compute deltas between two UEDM snapshots."""
    ai = a.get("items", pd.DataFrame())
    bi = b.get("items", pd.DataFrame())
    ar = a.get("relationships", pd.DataFrame())
    br = b.get("relationships", pd.DataFrame())
    aa = a.get("attributes", pd.DataFrame())
    ba = b.get("attributes", pd.DataFrame())

    added_items = []
    removed_items = []
    if "item_id" in ai.columns and "item_id" in bi.columns:
        aset = set(ai["item_id"].astype(str).tolist())
        bset = set(bi["item_id"].astype(str).tolist())
        added_items = sorted(list(bset - aset))
        removed_items = sorted(list(aset - bset))

    # relationships delta by signature
    def sig_rel(df: pd.DataFrame) -> set:
        if df.empty:
            return set()
        cols = [c for c in ["src","dst","rel_type","quantity","uom","effectivity"] if c in df.columns]
        return set(df[cols].astype(str).agg("|".join, axis=1).tolist())

    ar_s = sig_rel(ar)
    br_s = sig_rel(br)
    added_rels = sorted(list(br_s - ar_s))
    removed_rels = sorted(list(ar_s - br_s))

    # attribute numeric deltas (same item_id+key)
    delta_rows = []
    if (not aa.empty) and (not ba.empty) and ("item_id" in aa.columns) and ("key" in aa.columns) and ("item_id" in ba.columns) and ("key" in ba.columns):
        aa2 = aa.copy()
        ba2 = ba.copy()
        aa2["value_num"] = pd.to_numeric(aa2.get("value"), errors="coerce")
        ba2["value_num"] = pd.to_numeric(ba2.get("value"), errors="coerce")
        m = aa2.merge(ba2, on=["item_id","key"], how="inner", suffixes=("_a","_b"))
        # numeric where both numeric
        m = m[m["value_num_a"].notna() & m["value_num_b"].notna()]
        if not m.empty:
            m["delta"] = m["value_num_b"] - m["value_num_a"]
            # top magnitude
            m2 = m.reindex(m["delta"].abs().sort_values(ascending=False).index)
            for _, r in m2.head(50).iterrows():
                delta_rows.append({
                    "item_id": r["item_id"],
                    "key": r["key"],
                    "from": float(r["value_num_a"]),
                    "to": float(r["value_num_b"]),
                    "delta": float(r["delta"]),
                })

    return {
        "items_added": added_items,
        "items_removed": removed_items,
        "rels_added": added_rels,
        "rels_removed": removed_rels,
        "numeric_attribute_deltas": delta_rows,
        "kpi_a": uedm_kpi_summary(a),
        "kpi_b": uedm_kpi_summary(b),
    }
    if ext in {"stp","step","igs","iges"}:
        metrics["status"] = "cad_kernel_required_for_step_iges"
        return metrics, None

    try:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix="."+ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        mesh = trimesh.load(tmp_path, force="mesh")
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

        if mesh is None:
            metrics["status"] = "load_failed"
            return metrics, None

        # Ensure single mesh
        if isinstance(mesh, trimesh.Scene):
            geoms = [g for g in mesh.geometry.values()]
            if geoms:
                mesh = trimesh.util.concatenate(geoms)
            else:
                metrics["status"] = "empty_scene"
                return metrics, None

        metrics.update({
            "status": "ok",
            "vertices": int(getattr(mesh, "vertices", np.zeros((0,3))).shape[0]),
            "faces": int(getattr(mesh, "faces", np.zeros((0,3))).shape[0]),
        })

        bounds = mesh.bounds if hasattr(mesh, "bounds") else None
        if bounds is not None:
            (minv, maxv) = bounds
            dims = (maxv - minv).tolist()
            metrics["bbox_dims"] = {"x": float(dims[0]), "y": float(dims[1]), "z": float(dims[2])}

        if hasattr(mesh, "volume"):
            try:
                metrics["volume"] = float(mesh.volume)
            except Exception:
                pass
        if hasattr(mesh, "area"):
            try:
                metrics["surface_area"] = float(mesh.area)
            except Exception:
                pass

        # Plotly mesh
        v = np.asarray(mesh.vertices)
        f = np.asarray(mesh.faces)
        fig = go.Figure(data=[go.Mesh3d(
            x=v[:,0], y=v[:,1], z=v[:,2],
            i=f[:,0], j=f[:,1], k=f[:,2],
            opacity=0.7
        )])
        fig.update_layout(
            margin=dict(l=0, r=0, t=30, b=0),
            scene=dict(aspectmode="data"),
            height=520
        )
        return metrics, fig

    except Exception as e:
        metrics["status"] = "error"
        metrics["error"] = str(e)
        return metrics, None



def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def json_dumps_pretty(obj: Any) -> str:
    return json.dumps(obj, indent=2, ensure_ascii=False)


def load_cynthai_logo_bytes() -> Optional[bytes]:
    """Load CynthAI logo bytes from common locations; never raises."""
    candidates: List[str] = []

    # Environment override
    envp = os.environ.get("CYNTHAI_LOGO_PATH", "").strip()
    if envp:
        candidates.append(envp)

    # User-provided default path (Windows)
    candidates.append(r"C:\Users\d.adewunmi\OneDrive - Crowe Foederer BV\Documenten\EngineeringChange\CynthAI_logo.png")

    # Local repo paths
    base = os.path.dirname(__file__)
    candidates.extend([
        os.path.join(base, "CynthAI_logo.png"),
        os.path.join(base, "assets", "CynthAI_logo.png"),
        os.path.join(base, "assets", "cynthai_logo.png"),
    ])

    for pth in candidates:
        try:
            if pth and os.path.exists(pth):
                with open(pth, "rb") as f:
                    return f.read()
        except Exception:
            continue
    return None


def page_title(title: str, subtitle: Optional[str] = None) -> None:
    """Standard page header with optional logo aligned to the far right (all pages)."""
    logo = load_cynthai_logo_bytes()
    if logo:
        c1, c2 = st.columns([0.86, 0.14])
        with c1:
            st.markdown(f"## {title}")
            if subtitle:
                st.caption(subtitle)
        with c2:
            st.image(logo, use_container_width=True)
    else:
        st.markdown(f"## {title}")
        if subtitle:
            st.caption(subtitle)


def inject_theme() -> None:
    """Dark mode + refined red accent theme optimized for engineering workflows."""
    st.markdown(
        """
        <style>
        :root{
            --bg:#0b0f14;
            --panel:#111827;
            --panel2:#0f172a;
            --text:#e5e7eb;
            --muted:#9ca3af;
            --accent:#e11d48;
            --accent2:#fb7185;
            --border:#263041;
        }
        html, body, [class*="stApp"] { background: var(--bg) !important; color: var(--text) !important; }
        [data-testid="stHeader"] { background: rgba(0,0,0,0) !important; }
        [data-testid="stSidebar"] { background: linear-gradient(180deg, var(--panel2), var(--panel)) !important; }
        .stMarkdown, .stText, .stCaption { color: var(--text) !important; }
        .stCaption { color: var(--muted) !important; }

        .stButton>button {
            border: 1px solid var(--border) !important;
            background: #0f172a !important;
            color: var(--text) !important;
            border-radius: 12px !important;
            padding: 0.55rem 0.9rem !important;
        }
        .stButton>button[kind="primary"], .stDownloadButton>button {
            background: var(--accent) !important;
            border: 1px solid var(--accent) !important;
            color: white !important;
        }
        .stButton>button:hover, .stDownloadButton>button:hover { filter: brightness(1.05); }

        [data-baseweb="base-input"] input, [data-baseweb="textarea"] textarea{
            background: #0f172a !important;
            border: 1px solid var(--border) !important;
            border-radius: 12px !important;
            color: var(--text) !important;
        }
        [data-baseweb="select"]>div{
            background: #0f172a !important;
            border: 1px solid var(--border) !important;
            border-radius: 12px !important;
            color: var(--text) !important;
        }
        [data-testid="stDataFrame"] { border: 1px solid var(--border) !important; border-radius: 12px !important; overflow: hidden; }
        
        /* KPI cards (icons + glow hover) */
        .kpi-card{
            background: linear-gradient(180deg, rgba(17,24,39,0.95), rgba(15,23,42,0.95));
            border: 1px solid var(--border);
            border-radius: 16px;
            padding: 14px 14px 12px 14px;
            margin: 6px 0;
            transition: transform .18s ease, box-shadow .18s ease, border-color .18s ease, filter .18s ease;
            position: relative;
            overflow: hidden;
        }
        .kpi-card:before{
            content:"";
            position:absolute;
            inset:-40%;
            background: radial-gradient(circle at 20% 20%, rgba(225,29,72,0.18), transparent 60%),
                        radial-gradient(circle at 80% 30%, rgba(251,113,133,0.12), transparent 55%);
            transform: rotate(8deg);
            opacity: .8;
            pointer-events:none;
        }
        .kpi-card:hover{
            transform: translateY(-2px);
            border-color: rgba(225,29,72,0.55);
            box-shadow: 0 0 0 1px rgba(225,29,72,0.2), 0 12px 30px rgba(0,0,0,0.45), 0 0 24px rgba(225,29,72,0.18);
            filter: brightness(1.05);
        }
        .kpi-top{
            display:flex;
            align-items:center;
            gap:10px;
            position:relative;
            z-index:1;
        }
        .kpi-ic{
            width:34px;height:34px;
            border-radius: 12px;
            display:flex;
            align-items:center;
            justify-content:center;
            background: rgba(225,29,72,0.12);
            border: 1px solid rgba(225,29,72,0.22);
            font-size: 18px;
        }
        .kpi-label{
            font-size: 12.5px;
            color: var(--muted);
            letter-spacing: .2px;
        }
        .kpi-value{
            margin-top: 8px;
            font-size: 26px;
            font-weight: 800;
            line-height: 1.05;
            color: var(--text);
            position:relative;
            z-index:1;
        }
        .kpi-help{
            margin-top: 6px;
            font-size: 11.5px;
            color: var(--muted);
            position:relative;
            z-index:1;
        }
</style>
        """,
        unsafe_allow_html=True,
    )


# ----------------------------
# DB Layer (SQLite)
# ----------------------------
def db() -> sqlite3.Connection:
    ensure_dirs()
    con = sqlite3.connect(DB_PATH, check_same_thread=False, timeout=30)
    con.execute("PRAGMA journal_mode=WAL;")
    con.execute("PRAGMA busy_timeout=30000;")
    con.execute("PRAGMA synchronous=NORMAL;")
    con.execute("PRAGMA foreign_keys=ON;")
    return con


def _db_write_retry(fn, retries: int = 8, base_sleep: float = 0.15):
    """
    Execute a DB write function under a process lock and retry if sqlite reports 'database is locked'.
    The callable must perform its own commit/close (or use a context inside).
    """
    last_err = None
    for i in range(retries):
        try:
            with DB_WRITE_LOCK:
                return fn()
        except sqlite3.OperationalError as e:
            last_err = e
            if "locked" in str(e).lower():
                time.sleep(base_sleep * (i + 1))
                continue
            raise
    if last_err:
        raise last_err


def init_db() -> None:
    con = db()
    cur = con.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS documents (
        doc_id TEXT PRIMARY KEY,
        filename TEXT,
        file_ext TEXT,
        mime TEXT,
        sha256 TEXT,
        stored_path TEXT,
        uploaded_at TEXT,
        metadata_json TEXT,
        extracted_text TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS chunks (
        chunk_id TEXT PRIMARY KEY,
        doc_id TEXT,
        chunk_index INTEGER,
        text TEXT,
        meta_json TEXT,
        FOREIGN KEY(doc_id) REFERENCES documents(doc_id) ON DELETE CASCADE
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS parts (
        part_id TEXT PRIMARY KEY,
        part_number TEXT,
        part_name TEXT,
        revision TEXT,
        attributes_json TEXT,
        created_at TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS bom (
        bom_id TEXT PRIMARY KEY,
        parent_part_id TEXT,
        child_part_id TEXT,
        quantity REAL,
        uom TEXT,
        effectivity TEXT,
        notes TEXT,
        created_at TEXT,
        FOREIGN KEY(parent_part_id) REFERENCES parts(part_id),
        FOREIGN KEY(child_part_id) REFERENCES parts(part_id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS changes (
        change_id TEXT PRIMARY KEY,
        change_type TEXT, -- ECR or ECO
        title TEXT,
        status TEXT,
        requester TEXT,
        owner TEXT,
        created_at TEXT,
        updated_at TEXT,
        intent_json TEXT,
        impacted_json TEXT,
        risk_json TEXT,
        workflow_json TEXT,
        evidence_json TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS change_actions (
        action_id TEXT PRIMARY KEY,
        change_id TEXT,
        action_type TEXT,
        actor TEXT,
        notes TEXT,
        created_at TEXT,
        payload_json TEXT,
        FOREIGN KEY(change_id) REFERENCES changes(change_id) ON DELETE CASCADE
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS change_history_training (
        row_id INTEGER PRIMARY KEY AUTOINCREMENT,
        change_id TEXT,
        text TEXT,
        outcome_delay INTEGER,
        outcome_quality_issue INTEGER,
        created_at TEXT
    )
    """)

    # --- External integration logs / caches (SAP / Teamcenter) ---
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS external_sync_log (
            id TEXT PRIMARY KEY,
            system TEXT,
            action TEXT,
            ref_id TEXT,
            status TEXT,
            payload_json TEXT,
            created_at TEXT
        );
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS external_bom_cache (
            id TEXT PRIMARY KEY,
            system TEXT,
            key TEXT,
            payload_json TEXT,
            created_at TEXT
        );
        """
    )

    # --- Session archives (UI state snapshots) ---
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS session_archives (
            archive_id TEXT PRIMARY KEY,
            name TEXT,
            payload_json TEXT,
            created_at TEXT
        );
        """
    )

    con.commit()
    con.close()


# ----------------------------
# Entity extraction (lightweight, extensible)
# ----------------------------
PART_PATTERNS = [
    # Common part number patterns like "ABC-12345", "123-456-789", "PN: 12345-AB"
    re.compile(r"\b(?:PN[:\s]*)?([A-Z0-9]{2,}(?:[-_][A-Z0-9]{2,}){1,5})\b"),
    re.compile(r"\b([0-9]{2,}(?:[-_][0-9]{2,}){1,5})\b"),
]
REV_PATTERN = re.compile(r"\b(?:REV|Rev|Revision)\s*[:\-]?\s*([A-Z0-9]{1,3})\b")
DIM_PATTERN = re.compile(r"\b(\d+(?:\.\d+)?)\s*(mm|MM|cm|CM|m|M|in|IN)\b")
REQ_PATTERN = re.compile(r"\b(shall|must|required to|requirement)\b", re.IGNORECASE)


def extract_entities(text: str) -> Dict[str, Any]:
    parts = set()
    for pat in PART_PATTERNS:
        for m in pat.finditer(text):
            pn = m.group(1).strip().upper()
            # reduce false positives: ignore pure dates like 2024-01-01, ignore too-short
            if len(pn) < 5:
                continue
            if re.match(r"^\d{4}[-_]\d{2}[-_]\d{2}$", pn):
                continue
            parts.add(pn)

    revs = set(m.group(1).strip().upper() for m in REV_PATTERN.finditer(text))
    dims = [{"value": float(m.group(1)), "uom": m.group(2)} for m in DIM_PATTERN.finditer(text)]

    req_like = bool(REQ_PATTERN.search(text))

    return {
        "part_numbers": sorted(parts),
        "revisions": sorted(revs),
        "dimensions": dims[:200],
        "requirements_like": req_like,
    }


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 150) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + chunk_size)
        chunks.append(text[i:j])
        i = max(j - overlap, j)
    return chunks


# ----------------------------
# Document ingestion
# ----------------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDFs.
    Strategy:
    1) pdfplumber (pdfminer) for embedded text
    2) PyMuPDF (fitz) for embedded/vector text (often better for drawings/title blocks)
    3) OCR fallback (requires pytesseract + local Tesseract install) using rendered page images
    """
    # Read OCR preferences from Streamlit session if available
    force_ocr = False
    ocr_pages = 1
    try:
        force_ocr = bool(st.session_state.get("force_pdf_ocr", False))
        ocr_pages = int(st.session_state.get("pdf_ocr_pages", 1))
    except Exception:
        pass

    def _ocr_with_tesseract(max_pages: int = 1) -> str:
        if not (PYTESSERACT_AVAILABLE and pytesseract is not None and fitz is not None and Image is not None):
            return ""
        try:
            # Allow overriding Tesseract path via env var if needed (Windows)
            tcmd = os.environ.get("TESSERACT_CMD", "").strip()
            if tcmd:
                pytesseract.pytesseract.tesseract_cmd = tcmd
        except Exception:
            pass

        out_text: List[str] = []
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            n = min(max_pages, doc.page_count)
            for i in range(n):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=fitz.Matrix(3.0, 3.0), alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                txt = pytesseract.image_to_string(img) or ""
                if txt.strip():
                    out_text.append(txt.strip())
            doc.close()
        except Exception:
            return ""
        return "\n\n".join(out_text).strip()

    if force_ocr:
        return _ocr_with_tesseract(max_pages=max(1, min(6, ocr_pages)))

    # 1) pdfplumber
    text_parts: List[str] = []
    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    try:
                        text_parts.append(page.extract_text() or "")
                    except Exception:
                        continue
        except Exception:
            pass
    text1 = "\n".join(text_parts).strip()

    # 2) PyMuPDF text extraction
    text2 = ""
    if len(text1) < 30 and fitz is not None:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages: List[str] = []
            for i in range(doc.page_count):
                try:
                    pages.append(doc.load_page(i).get_text("text") or "")
                except Exception:
                    continue
            doc.close()
            text2 = "\n".join(pages).strip()
        except Exception:
            text2 = ""

    best = text1 if len(text1) >= len(text2) else text2

    # 3) OCR fallback if still empty (common for scanned drawings)
    if len(best) < 30:
        best = _ocr_with_tesseract(max_pages=max(1, min(6, ocr_pages)))

    return best.strip()



def extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        return ""
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs).strip()



def extract_text_from_csv(file_bytes: bytes, max_rows: int = 40) -> str:
    """
    Extract a readable text representation from CSV content.
    Produces a compact preview (schema + head rows) to support LLM Q&A.
    """
    try:
        import pandas as pd
        import io as _io
        # Try common delimiters automatically
        bio = _io.BytesIO(file_bytes)
        try:
            df = pd.read_csv(bio)
        except Exception:
            bio = _io.BytesIO(file_bytes)
            df = pd.read_csv(bio, sep=";")
        # Build preview text
        lines = []
        lines.append(f"CSV rows: {len(df)} | columns: {len(df.columns)}")
        lines.append("Columns: " + ", ".join([str(c) for c in df.columns.tolist()]))
        head = df.head(max_rows)
        lines.append("\nPreview (head):")
        # Use to_string for stable formatting
        lines.append(head.to_string(index=False))
        return "\n".join(lines).strip()
    except Exception:
        # Fallback to raw text
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def extract_text_from_xlsx(file_bytes: bytes, max_rows: int = 30) -> str:
    """
    Extract a readable text representation from Excel content (XLSX/XLS).
    Produces per-sheet schema + head rows.
    """
    try:
        import pandas as pd
        import io as _io
        bio = _io.BytesIO(file_bytes)
        sheets = pd.read_excel(bio, sheet_name=None)
        out_lines = []
        for sheet_name, df in sheets.items():
            out_lines.append(f"Sheet: {sheet_name} | rows: {len(df)} | cols: {len(df.columns)}")
            out_lines.append("Columns: " + ", ".join([str(c) for c in df.columns.tolist()]))
            out_lines.append(df.head(max_rows).to_string(index=False))
            out_lines.append("")
        return "\n".join(out_lines).strip()
    except Exception:
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def render_pdf_page_to_png_bytes(file_bytes: bytes, pageno: int = 0, zoom: float = 2.0) -> Optional[bytes]:
    if fitz is None:
        return None
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pageno = max(0, min(pageno, doc.page_count - 1))
    page = doc.load_page(pageno)
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return pix.tobytes("png")


def read_image(file_bytes: bytes) -> Optional["Image.Image"]:
    if Image is None:
        return None
    try:
        return Image.open(io.BytesIO(file_bytes)).convert("RGB")
    except Exception:
        return None



def zoomable_image_viewer(img_bytes: bytes, height: int = 560, label: str = "") -> None:
    """High-fidelity zoom + pan viewer for technical drawings (wheel zoom, drag pan)."""
    if not img_bytes:
        st.info("No image available.")
        return
    uid = uuid.uuid4().hex
    b64 = base64.b64encode(img_bytes).decode("utf-8")
    title = f"<div style='color:#bbb;font-size:13px;margin:0 0 6px 0;'>{label}</div>" if label else ""
    html = f"""
    <div style="background:#0f0f10;border:1px solid #2a2a2a;border-radius:10px;padding:10px;">
      {title}
      <div id="wrap_{uid}" style="height:{height}px;overflow:hidden;position:relative;border-radius:8px;border:1px solid #1f1f1f;">
        <img id="img_{uid}" src="data:image/png;base64,{b64}"
             style="transform-origin:0 0;cursor:grab;user-select:none;max-width:none;will-change:transform;" />
      </div>
      <div style="display:flex;gap:10px;align-items:center;margin-top:8px;color:#999;font-size:12px;">
        <span>Wheel to zoom • Drag to pan</span>
        <span id="z_{uid}"></span>
      </div>
    </div>
    <script>
      (function() {{
        const wrap = document.getElementById("wrap_{uid}");
        const img = document.getElementById("img_{uid}");
        const z = document.getElementById("z_{uid}");
        let scale = 1.0, tx = 0, ty = 0;
        let dragging = false, sx = 0, sy = 0;

        function clamp(v, lo, hi) {{ return Math.min(Math.max(v, lo), hi); }}
        function apply() {{
          scale = clamp(scale, 0.2, 12.0);
          img.style.transform = `translate(${{tx}}px, ${{ty}}px) scale(${{scale}})`;
          z.textContent = `Zoom: ${{scale.toFixed(2)}}×`;
        }}
        wrap.addEventListener("wheel", (e) => {{
          e.preventDefault();
          const rect = wrap.getBoundingClientRect();
          const mx = e.clientX - rect.left;
          const my = e.clientY - rect.top;
          const delta = e.deltaY < 0 ? 1.12 : 0.88;
          tx = mx - (mx - tx) * delta;
          ty = my - (my - ty) * delta;
          scale *= delta;
          apply();
        }}, {{ passive: false }});

        wrap.addEventListener("mousedown", (e) => {{
          dragging = true;
          sx = e.clientX; sy = e.clientY;
          img.style.cursor = "grabbing";
        }});
        window.addEventListener("mouseup", () => {{
          dragging = false;
          img.style.cursor = "grab";
        }});
        window.addEventListener("mousemove", (e) => {{
          if (!dragging) return;
          tx += (e.clientX - sx);
          ty += (e.clientY - sy);
          sx = e.clientX; sy = e.clientY;
          apply();
        }});
        apply();
      }})();
    </script>
    """
    components.html(html, height=height + 120)



def safe_st_image(img, caption: str = "", use_container_width: bool = True, **kwargs):
    """Render an image safely (bytes/path/PIL). If Streamlit fails to parse it, show a friendly warning."""
    if img is None:
        st.warning(f"{caption or 'Image'} unavailable.")
        return
    try:
        st.image(img, caption=caption, use_container_width=use_container_width, **kwargs)
    except Exception as e:
        st.warning(f"Could not render {caption or 'image'} (invalid/unsupported image data).")
        with st.expander("Details", expanded=False):
            st.code(str(e))
            try:
                if isinstance(img, (bytes, bytearray)):
                    st.caption(f"Bytes length: {len(img):,}")
                    sig = bytes(img[:16])
                    st.caption(f"Header bytes: {sig!r}")
            except Exception:
                pass


def blend_png_bytes(a_png: bytes, b_png: bytes, alpha: float = 0.5) -> Optional[bytes]:
    """Blend two raster images (bytes) into a single overlay image."""
    if Image is None:
        return None
    ia = read_image(a_png) if isinstance(a_png, (bytes, bytearray)) else None
    ib = read_image(b_png) if isinstance(b_png, (bytes, bytearray)) else None
    if ia is None or ib is None:
        return None
    ia = ia.convert("RGBA"); ib = ib.convert("RGBA")
    w = max(ia.size[0], ib.size[0]); h = max(ia.size[1], ib.size[1])
    ia = ia.resize((w, h)); ib = ib.resize((w, h))
    out = Image.blend(ia, ib, float(alpha))
    buf = io.BytesIO(); out.convert("RGB").save(buf, format="PNG")
    return buf.getvalue()


def pdf_bytes_from_text(title: str, text: str) -> bytes:
    """Create a simple, legible PDF from plain/markdown-ish text."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
    except Exception:
        return (text or "").encode("utf-8", errors="ignore")

    raw = (text or "").replace("\r\n", "\n")
    raw = re.sub(r"^#\s*", "", raw, flags=re.M)
    raw = re.sub(r"[*_]{1,3}", "", raw)
    lines_out = []
    for ln in raw.split("\n"):
        ln = ln.strip("\n")
        if not ln:
            lines_out.append("")
        else:
            lines_out.extend(textwrap.wrap(ln, width=100) or [""])
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4
    x = 42
    y = h - 54
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title[:120])
    y -= 22
    c.setFont("Helvetica", 10.5)
    for ln in lines_out:
        if y < 60:
            c.showPage()
            y = h - 54
            c.setFont("Helvetica", 10.5)
        c.drawString(x, y, ln[:140])
        y -= 14
    c.showPage()
    c.save()
    return buf.getvalue()


def pdf_bytes_procedure_pack(
    title: str,
    procedure_text: str,
    change_request: str = "",
    change_plan: Optional[Dict[str, Any]] = None,
    images: Optional[Dict[str, bytes]] = None,
    user: str = "",
) -> bytes:
    """Create a procedure pack PDF with change details + embedded images (baseline/updated/diff)."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import ImageReader
    except Exception:
        return (procedure_text or "").encode("utf-8", errors="ignore")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4

    def header():
        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, H - 50, title or "Procedure Pack")
        c.setFont("Helvetica", 10)
        meta = []
        if user:
            meta.append(f"User: {user}")
        meta.append(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        c.drawString(40, H - 66, " | ".join(meta))
        c.line(40, H - 72, W - 40, H - 72)

    def draw_paragraph(text: str, x: float, y: float, width_chars: int = 110, line_h: int = 12):
        c.setFont("Helvetica", 10)
        for ln in (text or "").splitlines():
            ln = ln.rstrip()
            if not ln.strip():
                y -= line_h
                continue
            wrapped = textwrap.wrap(ln, width=width_chars) or [ln]
            for wln in wrapped:
                if y < 60:
                    c.showPage()
                    header()
                    y = H - 90
                    c.setFont("Helvetica", 10)
                c.drawString(x, y, wln)
                y -= line_h
        return y

    header()
    y = H - 95

    # Section: Change request + plan
    if (change_request or "").strip():
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "Engineering Change Request")
        y -= 18
        y = draw_paragraph(change_request, 40, y)
        y -= 10

    if isinstance(change_plan, dict) and change_plan:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "Change Details (Plan)")
        y -= 18
        plan_txt = json.dumps(change_plan, indent=2, ensure_ascii=False)
        y = draw_paragraph(plan_txt, 40, y, width_chars=120, line_h=10)
        y -= 10

    # Section: Images
    imgs = images or {}
    if imgs:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "Visuals (Baseline / Updated / Diff)")
        y -= 18

        for key in ["baseline", "updated", "diff"]:
            if key not in imgs or not imgs[key]:
                continue
            try:
                img = ImageReader(io.BytesIO(imgs[key]))
                iw, ih = img.getSize()
                max_w = W - 80
                max_h = (H - 160) / 2.0
                scale = min(max_w / float(iw), max_h / float(ih), 1.0)
                dw, dh = iw * scale, ih * scale
                if y - dh < 60:
                    c.showPage()
                    header()
                    y = H - 90
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(40, y, "Visuals (continued)")
                    y -= 18
                c.setFont("Helvetica-Bold", 11)
                c.drawString(40, y, key.capitalize())
                y -= 14
                c.drawImage(img, 40, y - dh, width=dw, height=dh, preserveAspectRatio=True, mask='auto')
                y -= (dh + 14)
            except Exception:
                continue

    # Section: Procedure
    c.showPage()
    header()
    y = H - 95
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Revised Procedure")
    y -= 18
    _ = draw_paragraph(procedure_text or "", 40, y, width_chars=110)

    c.save()
    buf.seek(0)
    return buf.read()



def dcu_register_artifact(path: str) -> None:
    st.session_state.setdefault("dcu_artifacts", [])
    if path and path not in st.session_state["dcu_artifacts"]:
        st.session_state["dcu_artifacts"].append(path)


def dcu_clear_outputs(delete_files: bool = False) -> None:
    paths = st.session_state.get("dcu_artifacts", []) or []
    if delete_files:
        for p in paths:
            try:
                if p and os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass
    for k in [
        "dcu_plan", "dcu_updated_diagram", "dcu_revised_procedure",
        "dcu_updated_rep", "dcu_artifacts", "dcu_manual_pdf_path"
    ]:
        if k in st.session_state:
            del st.session_state[k]


def dcu_archive_outputs() -> Optional[str]:
    ensure_dirs()
    paths = st.session_state.get("dcu_artifacts", []) or []
    if not paths:
        return None
    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    folder = os.path.join(ARCHIVES_DIR, f"dcu_{stamp}")
    os.makedirs(folder, exist_ok=True)
    for p in paths:
        try:
            if p and os.path.exists(p):
                dst = os.path.join(folder, os.path.basename(p))
                with open(p, "rb") as fsrc, open(dst, "wb") as fdst:
                    fdst.write(fsrc.read())
        except Exception:
            pass
    return folder



def upsert_document(filename: str, file_bytes: bytes, mime: str = "", metadata: Optional[Dict[str, Any]] = None) -> str:
    """
    Ingest a file into the document library and generate deterministic chunks.
    Robust against concurrent writes (uses a process lock + retry).
    """
    ensure_dirs()
    doc_id = f"doc_{sha256_bytes(file_bytes)[:16]}"
    ext = os.path.splitext(filename)[1].lower().strip(".")
    sha = sha256_bytes(file_bytes)

    stored_path = os.path.join(FILES_DIR, f"{doc_id}.{ext or 'bin'}")
    if not os.path.exists(stored_path):
        try:
            with open(stored_path, "wb") as f:
                f.write(file_bytes)
        except Exception:
            pass

    extracted_text = ""
    try:
        if ext == "pdf":
            extracted_text = extract_text_from_pdf(file_bytes)
        elif ext == "docx":
            extracted_text = extract_text_from_docx(file_bytes)
        elif ext == "csv":
            extracted_text = extract_text_from_csv(file_bytes)
        elif ext in ("xlsx", "xls"):
            extracted_text = extract_text_from_xlsx(file_bytes)
        elif ext in ("txt", "md", "json"):
            extracted_text = extract_text_from_txt(file_bytes)
    except Exception:
        extracted_text = ""

    meta = metadata or {}
    try:
        meta["entities"] = extract_entities(extracted_text) if extracted_text else []
    except Exception:
        meta["entities"] = []
    meta.setdefault("summary", "")
    meta.setdefault("revisions", [])
    meta.setdefault("dimensions", [])
    meta.setdefault("requirements_like", False)
    meta_json = json.dumps(meta, ensure_ascii=False)

    chunks_list = list(chunk_text(extracted_text, chunk_size=1000, overlap=150)) if extracted_text else []

    def _write():
        con = db()
        con.execute(
            """
            INSERT OR REPLACE INTO documents(doc_id, filename, file_ext, mime, sha256, stored_path, uploaded_at, metadata_json, extracted_text)
            VALUES(?,?,?,?,?,?,?,?,?)
            """,
            (doc_id, filename, ext, mime, sha, stored_path, now_iso(), meta_json, extracted_text),
        )
        con.execute("DELETE FROM chunks WHERE doc_id=?", (doc_id,))
        for idx, ch in enumerate(chunks_list):
            chunk_id = f"chk_{doc_id}_{idx:04d}"
            con.execute(
                """
                INSERT OR REPLACE INTO chunks(chunk_id, doc_id, chunk_index, text, meta_json)
                VALUES(?,?,?,?,?)
                """,
                (chunk_id, doc_id, idx, ch, json.dumps({"start": idx}, ensure_ascii=False)),
            )
        con.commit()
        con.close()

    _db_write_retry(_write)
    return doc_id

def list_documents() -> pd.DataFrame:
    con = db()
    df = pd.read_sql_query("SELECT doc_id, filename, file_ext, uploaded_at, sha256 FROM documents ORDER BY uploaded_at DESC", con)
    con.close()
    return df


def get_document(doc_id: str) -> Optional[Dict[str, Any]]:
    con = db()
    cur = con.cursor()
    cur.execute("SELECT doc_id, filename, file_ext, mime, sha256, stored_path, uploaded_at, metadata_json, extracted_text FROM documents WHERE doc_id=?", (doc_id,))
    row = cur.fetchone()
    con.close()
    if not row:
        return None
    return {
        "doc_id": row[0],
        "filename": row[1],
        "file_ext": row[2],
        "mime": row[3],
        "sha256": row[4],
        "stored_path": row[5],
        "uploaded_at": row[6],
        "metadata": json.loads(row[7] or "{}"),
        "extracted_text": row[8] or "",
    }



def get_document_chunks(doc_id: str, chunk_chars: int = 1200, overlap_chars: int = 150) -> List[Dict[str, Any]]:
    """
    Return text chunks for a document.
    - If the 'chunks' table contains entries, use them.
    - Otherwise, chunk the document's extracted_text deterministically.
    """
    # Try DB chunks first (if populated)
    try:
        con = db()
        cur = con.cursor()
        cur.execute(
            "SELECT chunk_id, chunk_index, text, meta_json FROM chunks WHERE doc_id=? ORDER BY chunk_index ASC",
            (doc_id,),
        )
        rows = cur.fetchall()
        con.close()
        if rows:
            out: List[Dict[str, Any]] = []
            for r in rows:
                out.append({
                    "chunk_id": r[0],
                    "doc_id": doc_id,
                    "chunk_index": int(r[1]),
                    "text": r[2] or "",
                    "meta": json.loads(r[3] or "{}"),
                })
            return out
    except Exception:
        # fall back to deterministic chunking
        pass

    doc = get_document(doc_id)
    if not doc:
        return []
    text = (doc.get("extracted_text") or "").strip()
    if not text:
        return []

    # Deterministic sliding-window chunking by characters
    chunks: List[Dict[str, Any]] = []
    i = 0
    idx = 0
    n = len(text)
    while i < n:
        j = min(n, i + chunk_chars)
        chunk = text[i:j].strip()
        if chunk:
            chunks.append({
                "chunk_id": f"{doc_id}-ch{idx:04d}",
                "doc_id": doc_id,
                "chunk_index": idx,
                "text": chunk,
                "meta": {"filename": doc.get("filename", ""), "start": i, "end": j},
            })
            idx += 1
        if j >= n:
            break
        i = max(0, j - overlap_chars)
    return chunks


@st.cache_data(show_spinner=False)
@st.cache_data(show_spinner=False, ttl=3600)
def get_document_chunks_cached(doc_id: str, chunk_chars: int = 1200, overlap_chars: int = 150) -> List[Dict[str, Any]]:
    """Cached wrapper around get_document_chunks for faster reruns."""
    return get_document_chunks(doc_id, chunk_chars=chunk_chars, overlap_chars=overlap_chars)


def delete_document(doc_id: str) -> None:
    con = db()
    cur = con.cursor()
    cur.execute("SELECT stored_path FROM documents WHERE doc_id=?", (doc_id,))
    row = cur.fetchone()
    if row and row[0] and os.path.exists(row[0]):
        try:
            os.remove(row[0])
        except Exception:
            pass
    cur.execute("DELETE FROM documents WHERE doc_id=?", (doc_id,))
    con.commit()
    con.close()

def reprocess_document(doc_id: str, force_ocr: bool = False) -> None:
    """Re-run extraction/chunking on an existing stored document."""
    doc = get_document(doc_id)
    if not doc:
        return
    stored_path = doc.get("stored_path")
    if not stored_path or not os.path.exists(stored_path):
        return
    try:
        with open(stored_path, "rb") as f:
            b = f.read()
        # Temporarily override OCR setting for this run
        if force_ocr:
            try:
                st.session_state["force_pdf_ocr"] = True
            except Exception:
                pass
        upsert_document(doc.get("filename", os.path.basename(stored_path)), b, mime=str(doc.get("mime") or ""), metadata={"reprocessed": True, "forced_ocr": bool(force_ocr)})
    finally:
        if force_ocr:
            try:
                st.session_state["force_pdf_ocr"] = False
            except Exception:
                pass



# ----------------------------
# Parts/BOM / graph
# ----------------------------

def upsert_part(part_number: str, part_name: str = "", revision: str = "", attributes: Optional[Dict[str, Any]] = None) -> str:
    part_number = (part_number or "").strip().upper()
    if not part_number:
        raise ValueError("part_number required")
    part_id = f"part_{hashlib.sha256(part_number.encode('utf-8')).hexdigest()[:16]}"

    def _write():
        con = db()
        con.execute(
            """
            INSERT OR REPLACE INTO parts(part_id, part_number, part_name, revision, attributes_json, created_at)
            VALUES(?,?,?,?,?,?)
            """,
            (part_id, part_number, part_name, revision, json.dumps(attributes or {}, ensure_ascii=False), now_iso()),
        )
        con.commit()
        con.close()

    _db_write_retry(_write)
    return part_id

def list_parts() -> pd.DataFrame:
    con = db()
    df = pd.read_sql_query("SELECT part_id, part_number, part_name, revision, created_at FROM parts ORDER BY part_number", con)
    con.close()
    return df




def list_changes() -> pd.DataFrame:
    """
    List all ECR/ECO records in the local change register.
    Returns a DataFrame for consistent UI handling (len(df) works for metrics).
    """
    con = db()
    try:
        df = pd.read_sql_query(
            """
            SELECT change_id, change_type, title, status, requester, owner, created_at, updated_at
            FROM changes
            ORDER BY updated_at DESC
            """,
            con,
        )
    except Exception:
        df = pd.DataFrame(columns=["change_id","change_type","title","status","requester","owner","created_at","updated_at"])
    con.close()
    return df

def get_part_by_number(part_number: str) -> Optional[str]:
    con = db()
    cur = con.cursor()
    cur.execute("SELECT part_id FROM parts WHERE part_number=?", ((part_number or "").strip().upper(),))
    row = cur.fetchone()
    con.close()
    return row[0] if row else None



def import_bom_csv(df: pd.DataFrame) -> Tuple[int, int]:
    """
    Expected columns (case-insensitive):
      parent_part_number, child_part_number, quantity
    Optional:
      uom, effectivity, notes

    Robust implementation:
    - Avoids opening multiple sqlite connections per row (prevents 'database is locked')
    - Inserts parts + bom edges in a single write transaction with retry.
    """
    cols = {c.lower().strip(): c for c in df.columns}
    required = ["parent_part_number", "child_part_number", "quantity"]
    for r in required:
        if r not in cols:
            raise ValueError(f"Missing required column: {r}")

    edges: List[Tuple[str, str, float, str, str, str]] = []
    part_numbers: set = set()

    for _, row in df.iterrows():
        parent_pn = str(row[cols["parent_part_number"]]).strip().upper()
        child_pn = str(row[cols["child_part_number"]]).strip().upper()
        if not parent_pn or not child_pn or parent_pn.lower() == "nan" or child_pn.lower() == "nan":
            continue

        qty = safe_float(row[cols["quantity"]], 1.0)

        uom = ""
        if "uom" in cols:
            try:
                uom = str(row[cols["uom"]]).strip()
            except Exception:
                uom = ""
        uom = uom or "EA"

        eff = ""
        if "effectivity" in cols:
            try:
                eff = str(row[cols["effectivity"]]).strip()
            except Exception:
                eff = ""

        notes = ""
        if "notes" in cols:
            try:
                notes = str(row[cols["notes"]]).strip()
            except Exception:
                notes = ""

        edges.append((parent_pn, child_pn, float(qty), uom, eff, notes))
        part_numbers.add(parent_pn)
        part_numbers.add(child_pn)

    inserted_links = len(edges)
    if inserted_links == 0:
        return len(list_parts()), 0

    def _write():
        con = db()
        cur = con.cursor()

        # Fetch existing part_ids for these part_numbers
        pn_list = list(part_numbers)
        pn_to_id: Dict[str, str] = {}

        # sqlite has a parameter limit; chunk IN queries
        for i in range(0, len(pn_list), 800):
            chunk = pn_list[i:i+800]
            qmarks = ",".join(["?"] * len(chunk))
            for pn, pid in cur.execute(f"SELECT part_number, part_id FROM parts WHERE part_number IN ({qmarks})", chunk).fetchall():
                pn_to_id[str(pn).upper()] = pid

        # Upsert any missing parts inside THIS transaction (no extra connections)
        for pn in pn_list:
            if pn not in pn_to_id:
                part_id = f"part_{hashlib.sha256(pn.encode('utf-8')).hexdigest()[:16]}"
                pn_to_id[pn] = part_id
                cur.execute(
                    """
                    INSERT OR REPLACE INTO parts(part_id, part_number, part_name, revision, attributes_json, created_at)
                    VALUES(?,?,?,?,?,?)
                    """,
                    (part_id, pn, "", "", json.dumps({}, ensure_ascii=False), now_iso()),
                )

        # Insert bom edges
        for parent_pn, child_pn, qty, uom, eff, notes in edges:
            parent_id = pn_to_id.get(parent_pn)
            child_id = pn_to_id.get(child_pn)
            if not parent_id or not child_id:
                continue
            bom_key = f"{parent_id}|{child_id}|{qty}|{uom}|{eff}"
            bom_id = f"bom_{hashlib.sha256(bom_key.encode('utf-8')).hexdigest()[:16]}"
            cur.execute(
                """
                INSERT OR REPLACE INTO bom(bom_id, parent_part_id, child_part_id, quantity, uom, effectivity, notes, created_at)
                VALUES(?,?,?,?,?,?,?,?)
                """,
                (bom_id, parent_id, child_id, qty, uom, eff, notes, now_iso()),
            )

        con.commit()
        # total parts after import
        total_parts = cur.execute("SELECT COUNT(*) FROM parts").fetchone()[0]
        con.close()
        return int(total_parts)

    total_parts = _db_write_retry(_write)
    return int(total_parts), int(inserted_links)

def build_graph() -> nx.DiGraph:
    con = db()
    parts_df = pd.read_sql_query("SELECT part_id, part_number, part_name, revision FROM parts", con)
    bom_df = pd.read_sql_query("SELECT parent_part_id, child_part_id, quantity, uom, effectivity, notes FROM bom", con)
    docs_df = pd.read_sql_query("SELECT doc_id, filename, metadata_json FROM documents", con)
    con.close()

    G = nx.DiGraph()

    # Part nodes
    for _, r in parts_df.iterrows():
        G.add_node(r["part_id"], type="part", label=r["part_number"], part_number=r["part_number"], part_name=r.get("part_name", ""), revision=r.get("revision", ""))

    # BOM edges
    for _, r in bom_df.iterrows():
        if r["parent_part_id"] in G and r["child_part_id"] in G:
            G.add_edge(r["parent_part_id"], r["child_part_id"], type="bom", quantity=float(r.get("quantity", 1.0) or 1.0), uom=r.get("uom", "EA"), effectivity=r.get("effectivity", ""), notes=r.get("notes", ""))

    # Document nodes and doc->part edges
    for _, r in docs_df.iterrows():
        doc_id = r["doc_id"]
        meta = json.loads(r.get("metadata_json") or "{}")
        entities = meta.get("entities", {})
        pns = entities.get("part_numbers", []) or []
        G.add_node(doc_id, type="doc", label=r["filename"], filename=r["filename"])
        for pn in pns[:200]:
            part_id = get_part_by_number(pn)
            if part_id and part_id in G:
                G.add_edge(doc_id, part_id, type="mentions")

    return G


def plot_graph(G: nx.DiGraph, focus_nodes: Optional[List[str]] = None) -> None:
    # Simple force layout
    if len(G.nodes) == 0:
        st.info("Graph is empty. Import BOM and/or ingest documents to populate it.")
        return

    # build subgraph around focus
    H = G
    if focus_nodes:
        nodes = set()
        for n in focus_nodes:
            if n in G:
                nodes.add(n)
                nodes |= set(nx.single_source_shortest_path_length(G.to_undirected(), n, cutoff=2).keys())
        H = G.subgraph(list(nodes)).copy()

    pos = nx.spring_layout(H, seed=42, k=0.9 / max(1, math.sqrt(len(H.nodes))))

    # edges
    edge_x, edge_y = [], []
    for u, v in H.edges():
        x0, y0 = pos[u]
        x1, y1 = pos[v]
        edge_x += [x0, x1, None]
        edge_y += [y0, y1, None]

    edge_trace = go.Scatter(
        x=edge_x, y=edge_y,
        line=dict(width=0.8),
        hoverinfo='none',
        mode='lines'
    )

    # nodes
    node_x, node_y, node_text, node_size = [], [], [], []
    for n in H.nodes():
        x, y = pos[n]
        node_x.append(x); node_y.append(y)
        t = H.nodes[n].get("type", "")
        lbl = H.nodes[n].get("label", n)
        node_text.append(f"{t}: {lbl}")
        node_size.append(14 if t == "part" else 10)

    node_trace = go.Scatter(
        x=node_x, y=node_y,
        mode='markers+text',
        text=[H.nodes[n].get("label", "") for n in H.nodes()],
        textposition="top center",
        hovertext=node_text,
        hoverinfo="text",
        marker=dict(size=node_size, line=dict(width=1)),
    )

    fig = go.Figure(data=[edge_trace, node_trace])
    fig.update_layout(
        showlegend=False,
        margin=dict(l=10, r=10, t=10, b=10),
        height=540
    )
    st.plotly_chart(fig, use_container_width=True)


# ----------------------------
# Retrieval (local TF-IDF over chunks)
# ----------------------------
@dataclass
class RetrievalResult:
    chunk_id: str
    doc_id: str
    score: float
    text: str


def retrieve(query: str, top_k: int = 6) -> List[RetrievalResult]:
    if TfidfVectorizer is None:
        return []
    con = db()
    chunks_df = pd.read_sql_query("SELECT chunk_id, doc_id, text FROM chunks", con)
    con.close()
    if chunks_df.empty or not (query or "").strip():
        return []
    corpus = chunks_df["text"].fillna("").tolist()
    vec = TfidfVectorizer(stop_words="english", max_features=6000)
    X = vec.fit_transform(corpus + [query])
    q = X[-1]
    sims = (X[:-1] @ q.T).toarray().ravel()
    idx = np.argsort(-sims)[:top_k]
    out = []
    for i in idx:
        out.append(RetrievalResult(
            chunk_id=str(chunks_df.iloc[i]["chunk_id"]),
            doc_id=str(chunks_df.iloc[i]["doc_id"]),
            score=float(sims[i]),
            text=str(chunks_df.iloc[i]["text"]),
        ))
    return out


@st.cache_data(show_spinner=False)
def _tfidf_rank_small(query: str, texts: List[str], top_k: int = 8) -> List[int]:
    """Fast, small-corpus TF-IDF ranking for up to ~50 chunks."""
    if TfidfVectorizer is None or not texts:
        return list(range(min(top_k, len(texts))))
    try:
        vec = TfidfVectorizer(stop_words="english", max_features=6000)
        X = vec.fit_transform(texts + [query])
        qv = X[-1]
        sims = (X[:-1] @ qv.T).toarray().ravel()
        idx = np.argsort(-sims)[:top_k]
        return [int(i) for i in idx]
    except Exception:
        return list(range(min(top_k, len(texts))))


def rank_context_blocks(question: str, blocks: List[Dict[str, str]], top_k: int = 8, max_chars_per_block: int = 1200) -> List[Dict[str, str]]:
    """
    Rank provided context blocks by relevance to the question (fast local TF-IDF).
    Returns the top_k blocks, with each block truncated to max_chars_per_block to reduce latency/cost.
    """
    if not blocks:
        return []
    texts = [(b.get("text") or "")[:max_chars_per_block] for b in blocks]
    order = _tfidf_rank_small(question or "", texts, top_k=top_k)
    out: List[Dict[str, str]] = []
    for i in order:
        if 0 <= i < len(blocks):
            b = dict(blocks[i])
            b["text"] = (b.get("text") or "")[:max_chars_per_block]
            out.append(b)
    return out


# ----------------------------
# Change objects + workflow
# ----------------------------
CHANGE_STATUSES = {
    "ECR": ["Draft", "Submitted", "In Review", "Approved", "Rejected", "Archived"],
    "ECO": ["Draft", "Submitted", "In Review", "Approved", "Released", "Rejected", "Archived"],
}

DEFAULT_WORKFLOW = {
    "roles": {
        "Requester": ["Draft", "Submitted"],
        "Change Owner": ["Draft", "Submitted", "In Review"],
        "Approver": ["In Review", "Approved", "Rejected", "Released"],
    },
    "approval_steps": [
        {"name": "Engineering Review", "role": "Approver", "required": True},
        {"name": "Quality Review", "role": "Approver", "required": True},
        {"name": "Manufacturing Review", "role": "Approver", "required": False},
    ]
}


def new_id(prefix: str) -> str:
    return f"{prefix}_{hashlib.sha256((prefix+now_iso()+str(time.time())).encode('utf-8')).hexdigest()[:12]}"



def create_change(change_type: str, title: str, requester: str, owner: str, intent: Dict[str, Any]) -> str:
    change_type = change_type.upper().strip()
    if change_type not in ("ECR", "ECO"):
        raise ValueError("change_type must be ECR or ECO")
    change_id = new_id(change_type.lower())

    def _write():
        con = db()
        con.execute(
            """
            INSERT INTO changes(change_id, change_type, title, status, requester, owner, created_at, updated_at,
                               intent_json, impacted_json, risk_json, workflow_json, evidence_json)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                change_id, change_type, title, "Draft", requester, owner, now_iso(), now_iso(),
                json.dumps(intent, ensure_ascii=False),
                json.dumps({}, ensure_ascii=False),
                json.dumps({}, ensure_ascii=False),
                json.dumps(DEFAULT_WORKFLOW, ensure_ascii=False),
                json.dumps({"retrieval": [], "docs": []}, ensure_ascii=False),
            ),
        )
        con.commit()
        con.close()

    _db_write_retry(_write)
    log_action(change_id, "CREATE", requester, notes=f"Created {change_type} in Draft", payload={"title": title})
    return change_id


def log_action(change_id: str, action_type: str, actor: str, notes: str = "", payload: Optional[Dict[str, Any]] = None) -> None:
    action_id = new_id("act")

    def _write():
        con = db()
        con.execute(
            """
            INSERT INTO actions(action_id, change_id, action_type, actor, notes, payload_json, created_at)
            VALUES(?,?,?,?,?,?,?)
            """,
            (
                action_id,
                change_id,
                action_type,
                actor,
                notes,
                json.dumps(payload or {}, ensure_ascii=False),
                now_iso(),
            ),
        )
        con.commit()
        con.close()

    _db_write_retry(_write)

    vals.append(now_iso())
    vals.append(change_id)
    con = db()
    con.execute(f"UPDATE changes SET {', '.join(sets)}, updated_at=? WHERE change_id=?", tuple(vals))
    con.commit()
    con.close()


def get_actions(change_id: str) -> pd.DataFrame:
    con = db()
    df = pd.read_sql_query("""
        SELECT created_at, actor, action_type, notes, payload_json
        FROM change_actions
        WHERE change_id=?
        ORDER BY created_at DESC
    """, con, params=(change_id,))
    con.close()
    return df


def can_transition(change_type: str, new_status: str) -> bool:
    return new_status in CHANGE_STATUSES.get(change_type, [])


# ----------------------------
# “Autonomous” generation (rule-based + optional OpenAI)
# ----------------------------


def _llm_provider_choice() -> str:
    """Returns one of: 'openai', 'mistral', 'auto'."""
    raw = (st.session_state.get("llm_provider") or "OpenAI").strip().lower()
    if raw.startswith("mistral"):
        return "mistral"
    if raw.startswith("auto"):
        return "auto"
    return "openai"


def mistral_api_key_from_session() -> str:
    return (st.session_state.get("mistral_api_key") or "").strip()


def mistral_is_configured() -> bool:
    return REQUESTS_AVAILABLE and bool(mistral_api_key_from_session())


def mistral_chat_complete(
    system: str,
    user: str,
    temperature: float = 0.2,
    model: Optional[str] = None,
    max_tokens: int = 900,
    stream: bool = False,
):
    """
    Minimal Mistral Chat Completions call via HTTPS.
    Returns (text, meta, error).
    """
    if not REQUESTS_AVAILABLE:
        return None, {"provider": "mistral"}, "requests_unavailable"
    key = mistral_api_key_from_session()
    if not key:
        return None, {"provider": "mistral"}, "no_mistral_key"
    model = model or (st.session_state.get("mistral_model") or DEFAULT_MISTRAL_MODEL)

    headers = {
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": float(temperature) if temperature is not None else 0.2,
        "max_tokens": int(max_tokens) if max_tokens is not None else 900,
        "stream": bool(stream),
    }

    try:
        if not stream:
            r = requests.post(MISTRAL_CHAT_URL, headers=headers, json=payload, timeout=90)
            if r.status_code >= 400:
                return None, {"provider": "mistral", "status": r.status_code, "body": r.text[:800]}, f"http_{r.status_code}"
            js = r.json()
            # OpenAI-like schema: choices[0].message.content
            content = ""
            try:
                choices = js.get("choices") or []
                if choices:
                    msg = (choices[0].get("message") or {})
                    content = (msg.get("content") or "").strip()
            except Exception:
                content = ""
            if not content:
                # last resort: stringify truncated response
                content = json.dumps(js, ensure_ascii=False)[:2000]
            meta = {"provider": "mistral", "model": js.get("model", model), "usage": js.get("usage", {})}
            return content, meta, None

        # Streaming (SSE): yield text deltas where present; if schema differs, yield raw lines.
        r = requests.post(MISTRAL_CHAT_URL, headers=headers, json=payload, timeout=90, stream=True)
        if r.status_code >= 400:
            return None, {"provider": "mistral", "status": r.status_code, "body": r.text[:800]}, f"http_{r.status_code}"

        def _iter():
            agg = ""
            for raw in r.iter_lines(decode_unicode=True):
                if not raw:
                    continue
                line = raw.strip()
                if line.startswith("data:"):
                    line = line[len("data:"):].strip()
                if line == "[DONE]":
                    break
                try:
                    ev = json.loads(line)
                    choices = ev.get("choices") or []
                    if choices:
                        delta = (choices[0].get("delta") or choices[0].get("message") or {})
                        # delta may contain {"content": "..."}
                        dtxt = (delta.get("content") or "")
                        if dtxt:
                            nonlocal_agg = None
                            # python closure safe
                            yield dtxt
                            agg += dtxt
                            continue
                except Exception:
                    # fallback: emit raw
                    yield line + "\n"
                    agg += line + "\n"
            # store aggregated
            return

        return _iter, {"provider": "mistral", "model": model}, None
    except Exception as e:
        return None, {"provider": "mistral", "model": model}, f"{type(e).__name__}: {e}"
class _OpenAIResponsesCompatResult:
    def __init__(self, text: str):
        self.output_text = text or ""


class _OpenAIResponsesCompatStream:
    def __init__(self, text: str):
        self._text = text or ""
    def __enter__(self):
        return self
    def __exit__(self, exc_type, exc, tb):
        return False
    def __iter__(self):
        # emit a single response-like delta event for compatibility
        if self._text:
            yield {"type": "response.output_text.delta", "delta": self._text}


def _responses_input_to_chat_messages(inp: Any) -> List[Dict[str, Any]]:
    msgs: List[Dict[str, Any]] = []
    if isinstance(inp, str):
        return [{"role": "user", "content": inp}]
    if not isinstance(inp, list):
        return [{"role": "user", "content": str(inp)}]
    for m in inp:
        if not isinstance(m, dict):
            continue
        role = m.get("role", "user")
        content = m.get("content", "")
        if isinstance(content, str):
            msgs.append({"role": role, "content": content})
            continue
        if isinstance(content, list):
            parts = []
            for c in content:
                if isinstance(c, dict):
                    if c.get("type") in ("input_text", "text", "output_text"):
                        parts.append(str(c.get("text", "")))
                    elif "content" in c:
                        parts.append(str(c.get("content", "")))
                elif c is not None:
                    parts.append(str(c))
            msgs.append({"role": role, "content": "\n".join([x for x in parts if x])})
        else:
            msgs.append({"role": role, "content": str(content)})
    return msgs or [{"role": "user", "content": ""}]


class _OpenAIResponsesCompat:
    def __init__(self, client: Any):
        self._client = client

    def create(self, **kwargs):
        chat = getattr(self._client, "chat", None)
        comps = getattr(chat, "completions", None) if chat is not None else None
        if comps is None or not hasattr(comps, "create"):
            raise AttributeError("OpenAI client has neither responses nor chat.completions API")
        model = kwargs.get("model")
        temperature = kwargs.get("temperature", None)
        max_tokens = kwargs.get("max_output_tokens") or kwargs.get("max_tokens")
        messages = _responses_input_to_chat_messages(kwargs.get("input"))
        payload = {"model": model, "messages": messages}
        if temperature is not None:
            payload["temperature"] = temperature
        if max_tokens is not None:
            payload["max_tokens"] = max_tokens
        resp = comps.create(**payload)
        text = ""
        try:
            choices = getattr(resp, "choices", None) or []
            if choices:
                msg = getattr(choices[0], "message", None)
                if msg is not None:
                    c = getattr(msg, "content", None)
                    if isinstance(c, str):
                        text = c
                    elif isinstance(c, list):
                        acc = []
                        for item in c:
                            t = getattr(item, "text", None) if not isinstance(item, dict) else item.get("text")
                            if t:
                                acc.append(str(t))
                        text = "".join(acc)
        except Exception:
            text = ""
        if not text:
            text = str(resp)
        return _OpenAIResponsesCompatResult(text)

    def stream(self, **kwargs):
        # Compatibility fallback: non-streaming create wrapped as iterable stream
        r = self.create(**kwargs)
        return _OpenAIResponsesCompatStream(getattr(r, "output_text", ""))


def openai_client_from_session() -> Optional["OpenAI"]:
    if not OPENAI_AVAILABLE:
        return None
    key = (st.session_state.get("openai_api_key") or "").strip()
    if not key:
        return None
    try:
        client = OpenAI(api_key=key)
        if not hasattr(client, "responses"):
            try:
                setattr(client, "responses", _OpenAIResponsesCompat(client))
            except Exception:
                pass
        return client
    except Exception:
        return None


def llm_generate_json(system: str, user: str, schema_hint: str, temperature: float = 0.2, model: str = DEFAULT_MODEL) -> Optional[Dict[str, Any]]:
    """
    Uses OpenAI Responses API to request JSON-ish output.
    We keep this robust: if parsing fails, we return None.
    """
    client = openai_client_from_session()
    if client is None:
        return None

    prompt = f"""{user}

Return ONLY valid JSON matching this schema (best effort):
{schema_hint}
"""
    try:
        resp = client.responses.create(
            model=model,
            input=[
                {
                    "role": "system",
                    "content": [{"type": "input_text", "text": system}],
                },
                {
                    "role": "user",
                    "content": [{"type": "input_text", "text": prompt}],
                },
            ],
            temperature=temperature,
        )
        txt = (getattr(resp, "output_text", "") or "").strip()
        # tolerate leading/trailing text by extracting first json object/array
        m = re.search(r"(\{.*\}|\[.*\])", txt, flags=re.S)
        if not m:
            return None
        return json.loads(m.group(1))
    except Exception:
        return None



def _robust_extract_json(text: str) -> Optional[Dict[str, Any]]:
    """
    Best-effort JSON extraction from LLM output.
    Returns dict if possible, else None.
    """
    if not text:
        return None
    t = text.strip()
    # Fast path: pure JSON
    try:
        obj = json.loads(t)
        return obj if isinstance(obj, dict) else {"_data": obj}
    except Exception:
        pass

    # Try to extract the first JSON object (first '{' to last '}')
    i = t.find("{")
    j = t.rfind("}")
    if i != -1 and j != -1 and j > i:
        candidate = t[i:j+1]
        try:
            obj = json.loads(candidate)
            return obj if isinstance(obj, dict) else {"_data": obj}
        except Exception:
            pass

    # Regex fallback
    m = re.search(r"(\\{.*\\})", t, flags=re.S)
    if m:
        try:
            obj = json.loads(m.group(1))
            return obj if isinstance(obj, dict) else {"_data": obj}
        except Exception:
            return None
    return None


def llm_generate_json_debug(
    system: str,
    user: str,
    schema_hint: str,
    temperature: float = 0.2,
    model: str = DEFAULT_MODEL,
) -> Tuple[Optional[Dict[str, Any]], Dict[str, Any]]:
    """
    Like llm_generate_json, but returns debug metadata explaining failures
    and tries a small model fallback list if the chosen model fails.
    """
    client = openai_client_from_session()
    if client is None:
        return None, {"status": "no_key_or_sdk", "error": "OpenAI client unavailable (missing key or SDK)."}

    prompt = f"""{user}

Return ONLY valid JSON matching this schema (best effort):
{schema_hint}
"""

    models_to_try = []
    chosen = (model or "").strip()
    if chosen:
        models_to_try.append(chosen)
    for cand in ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "gpt-4.1"]:
        if cand not in models_to_try:
            models_to_try.append(cand)

    last_err = None
    last_txt = ""
    used_model = None

    for mdl in models_to_try:
        try:
            resp = client.responses.create(
                model=mdl,
                input=[
                    {"role": "system", "content": [{"type": "input_text", "text": system}]},
                    {"role": "user", "content": [{"type": "input_text", "text": prompt}]},
                ],
                temperature=temperature,
            )
            txt = (getattr(resp, "output_text", "") or "").strip()
            used_model = mdl
            last_txt = txt
            parsed = _robust_extract_json(txt)
            if parsed is not None:
                return parsed, {"status": "ok", "model": used_model, "raw_preview": txt[:2000]}
            last_err = "json_parse_failed"
        except Exception as e:
            last_err = str(e)

    return None, {
        "status": "error",
        "model": used_model or chosen,
        "error": last_err or "unknown_error",
        "raw_preview": (last_txt or "")[:2000],
    }


def heuristic_change_plan(rep: Dict[str, Any], change_request: str) -> Dict[str, Any]:
    """
    Deterministic fallback: builds a structured plan from OCR-extracted attributes + user request.
    Keeps the workflow moving when AI is disabled or parsing fails.
    """
    cr = (change_request or "").strip()
    attrs = rep.get("attributes", {}) if isinstance(rep, dict) else {}
    ops = []

    arrow = re.split(r"\\s*(?:->|→)\\s*", cr)
    from_part = arrow[0].strip() if arrow else cr
    to_part = arrow[1].strip() if len(arrow) > 1 else ""

    dia_re = re.compile(r"(?:Ø|⌀)\\s*(\\d+(?:\\.\\d+)?)")
    tol_re = re.compile(r"(?:±|\\+/-)\\s*(\\d+(?:\\.\\d+)?)")
    ang_re = re.compile(r"(\\d+(?:\\.\\d+)?)\\s*(?:°|deg)\\b", re.I)

    from_dia = dia_re.search(from_part)
    to_dia = dia_re.search(to_part) if to_part else None
    if from_dia and to_dia:
        ops.append({
            "type": "update_dimension",
            "target": f"Diameter Ø{from_dia.group(1)}",
            "from": f"Ø{from_dia.group(1)}",
            "to": f"Ø{to_dia.group(1)}",
            "confidence": 0.55,
            "evidence": ["user_request"],
        })

    from_tol = tol_re.search(from_part)
    to_tol = tol_re.search(to_part) if to_part else None
    if from_tol and to_tol:
        ops.append({
            "type": "update_tolerance",
            "target": "Tolerance",
            "from": f"±{from_tol.group(1)}",
            "to": f"±{to_tol.group(1)}",
            "confidence": 0.50,
            "evidence": ["user_request"],
        })

    from_ang = ang_re.search(from_part)
    to_ang = ang_re.search(to_part) if to_part else None
    if from_ang and to_ang:
        ops.append({
            "type": "update_angle",
            "target": "Angle",
            "from": f"{from_ang.group(1)}°",
            "to": f"{to_ang.group(1)}°",
            "confidence": 0.50,
            "evidence": ["user_request"],
        })

    if any(k in cr.lower() for k in ["note", "annotation", "label", "callout", "text"]):
        ops.append({
            "type": "update_note",
            "target": "Annotation/Note (unspecified)",
            "from": "",
            "to": cr,
            "confidence": 0.35,
            "evidence": ["user_request"],
        })

    if not ops:
        ops.append({
            "type": "other",
            "target": "Diagram",
            "from": "",
            "to": cr,
            "confidence": 0.30,
            "evidence": ["user_request"],
        })

    img_prompt = (
        "Update this engineering diagram to apply the following change request exactly, while preserving all "
        "unchanged geometry, dimensions, annotations, title block styling, line weights, and overall layout:\\n"
        f"{cr}\\n"
        "Do not introduce new elements unless required by the change request."
    )

    plan = {
        "change_plan": {
            "summary": cr[:240],
            "detected_elements_used": ["ocr_attributes" if attrs else "none"],
            "operations": ops,
        },
        "image_edit_prompt": img_prompt,
        "procedure_update_instructions": {
            "add_steps": ["Update work instructions to reflect the new drawing attributes where applicable."],
            "remove_steps": [],
            "modify_steps": ["Update inspection/measurement steps tied to changed dimensions/tolerances/angles."],
        },
        "model_report": {
            "overall_confidence": float(np.mean([o.get("confidence", 0.3) for o in ops])) if ops else 0.30,
            "limitations": [
                "AI planning was unavailable; this plan is heuristic.",
                "Exact geometric targets may require AI Vision interpretation to localize callouts/bounding boxes.",
            ],
            "key_parameters": ["heuristic_parser=v1", "uses_ocr_attributes=" + str(bool(attrs))],
        },
    }
    return plan



def llm_generate_text(system: str, user: str, temperature: float = 0.2, model: Optional[str] = None) -> Tuple[Optional[str], Optional[str]]:
    """
    Returns (text, error).
    Routes to the selected provider:
      - OpenAI (Responses API) if configured
      - Mistral (Chat Completions HTTPS) if configured
      - Auto: OpenAI → Mistral fallback
    """
    provider = _llm_provider_choice()

    # 1) OpenAI path
    if provider in ("openai", "auto"):
        client = openai_client_from_session()
        if client is not None:
            model_ = model or st.session_state.get("openai_model", DEFAULT_MODEL)
            try:
                resp = client.responses.create(
                    model=model_,
                    input=[
                        {"role": "system", "content": [{"type": "input_text", "text": system}]},
                        {"role": "user", "content": [{"type": "input_text", "text": user}]},
                    ],
                    temperature=temperature,
                )
                txt = (getattr(resp, "output_text", "") or "").strip()
                return txt, None
            except Exception as e:
                # If auto, fall through to Mistral
                if provider != "auto":
                    return None, f"{type(e).__name__}: {e}"

        if provider != "auto":
            return None, "openai_unavailable_or_no_key"

    # 2) Mistral path
    if provider in ("mistral", "auto"):
        txt, meta, err = mistral_chat_complete(
            system=system,
            user=user,
            temperature=temperature,
            model=(model or st.session_state.get("mistral_model") or DEFAULT_MISTRAL_MODEL),
            max_tokens=900,
            stream=False,
        )
        if txt is not None and not err:
            return txt, None
        return None, err or "mistral_unavailable_or_no_key"

    return None, "llm_unavailable_or_parse_failed"

class _NamedBytesIO(io.BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name





def llm_stream_text(
    system: str,
    user: str,
    temperature: float = 0.2,
    model: str = DEFAULT_MODEL,
    max_output_tokens: int = 700,
) -> Tuple[Optional[callable], Dict[str, Any], Optional[str]]:
    """
    Returns (generator_fn, meta, error).
    generator_fn yields incremental text chunks for streaming.
    meta will contain the final aggregated text under meta["text"] when streaming finishes.
    """
    provider = _llm_provider_choice()
    meta: Dict[str, Any] = {"text": "", "provider": provider}

    # Mistral path (we stream as a single chunk by default for robustness)
    if provider == "mistral":
        def _gen_mistral():
            txt, m, err = mistral_chat_complete(
                system=system,
                user=user,
                temperature=temperature,
                model=(st.session_state.get("mistral_model") or DEFAULT_MISTRAL_MODEL),
                max_tokens=max_output_tokens,
                stream=False,
            )
            if err or txt is None:
                meta["text"] = ""
                yield f"[Mistral error] {err or 'unknown_error'}"
                return
            meta.update(m or {})
            meta["text"] = txt
            yield txt

        if not mistral_is_configured():
            # deterministic fallback
            def _gen_no_llm():
                txt = deterministic_narrative([], question=user)
                meta["text"] = txt
                yield txt
            return _gen_no_llm, meta, None

        return _gen_mistral, meta, None

    # OpenAI path (and Auto default)
    client = openai_client_from_session()
    if client is None:
        if provider == "auto" and mistral_is_configured():
            # fallback to Mistral
            st.session_state["llm_provider"] = "Auto (OpenAI → Mistral)"  # keep label stable
            def _gen_auto_mistral():
                txt, m, err = mistral_chat_complete(
                    system=system,
                    user=user,
                    temperature=temperature,
                    model=(st.session_state.get("mistral_model") or DEFAULT_MISTRAL_MODEL),
                    max_tokens=max_output_tokens,
                    stream=False,
                )
                if err or txt is None:
                    meta["text"] = ""
                    yield f"[Mistral error] {err or 'unknown_error'}"
                    return
                meta.update(m or {})
                meta["text"] = txt
                yield txt
            return _gen_auto_mistral, meta, None

        # deterministic narrative as one chunk
        def _gen_no_llm():
            txt = deterministic_narrative([], question=user)
            meta["text"] = txt
            yield txt
        return _gen_no_llm, meta, None

    model = model or st.session_state.get("openai_model", DEFAULT_MODEL)

    payload = dict(
        model=model,
        input=[
            {"role": "system", "content": [{"type": "input_text", "text": system}]},
            {"role": "user", "content": [{"type": "input_text", "text": user}]},
        ],
        temperature=temperature,
        max_output_tokens=max_output_tokens,
    )

    def _extract_delta(ev: Any) -> Optional[str]:
        # Handle both object and dict-shaped events (SDK versions vary)
        t = None
        delta = None
        if isinstance(ev, dict):
            t = ev.get("type") or ev.get("event") or ""
            delta = ev.get("delta") or ev.get("text") or ev.get("content")
        else:
            t = getattr(ev, "type", "") or getattr(ev, "event", "") or ""
            delta = getattr(ev, "delta", None) or getattr(ev, "text", None) or getattr(ev, "content", None)
        if delta is None:
            return None
        # Filter: only text deltas
        if t and ("delta" in t or "output_text" in t or "message" in t):
            return str(delta)
        # If type missing, still accept short deltas
        return str(delta)

    def _gen():
        try:
            # Preferred: responses.stream context manager (new SDK)
            if hasattr(getattr(client, "responses", None), "stream"):
                with client.responses.stream(**payload) as stream:
                    for ev in stream:
                        d = _extract_delta(ev)
                        if d:
                            meta["text"] += d
                            yield d
                return
            # Fallback: non-streaming
            resp = client.responses.create(**payload)
            txt = (getattr(resp, "output_text", "") or "").strip()
            meta["text"] = txt
            yield txt
        except Exception as e:
            yield f"[OpenAI error] {type(e).__name__}: {e}"

    return _gen, meta, None

def openai_image_edit_or_generate(
    baseline_image_bytes: bytes,
    prompt: str,
    size: str = "1024x1024",
    prefer_edit: bool = True,
    model: str = "gpt-image-1"
) -> Tuple[Optional[bytes], Optional[str]]:
    """Attempts an image edit first, then falls back to image generation."""
    client = openai_client_from_session()
    if client is None:
        return None, "no_openai_client"

    images_api = getattr(client, "images", None)
    if images_api is None:
        return None, "openai_images_api_unavailable"

    edit_err = None
    if prefer_edit:
        try:
            img_file = _NamedBytesIO(baseline_image_bytes, "baseline.png")
            if hasattr(images_api, "edit"):
                resp = images_api.edit(model=model, image=img_file, prompt=prompt, size=size)
            elif hasattr(images_api, "edits"):
                resp = images_api.edits(model=model, image=img_file, prompt=prompt, size=size)
            else:
                resp = None
            if resp is not None:
                data = getattr(resp, "data", None) or []
                if data and getattr(data[0], "b64_json", None):
                    return base64.b64decode(data[0].b64_json), None
        except Exception as e:
            msg = str(e)
            if ("must be verified" in msg.lower() and "gpt-image" in msg.lower()) or ("verify organization" in msg.lower() and "gpt-image" in msg.lower()):
                return None, "org_verification_required"
            edit_err = f"{type(e).__name__}: {e}"

    # Fallback: generation
    try:
        if hasattr(images_api, "generate"):
            resp = images_api.generate(model=model, prompt=prompt, size=size)
        elif hasattr(images_api, "create"):
            resp = images_api.create(model=model, prompt=prompt, size=size)
        else:
            return None, "openai_images_generate_unavailable"

        data = getattr(resp, "data", None) or []
        if data and getattr(data[0], "b64_json", None):
            return base64.b64decode(data[0].b64_json), None
        return None, "image_generation_no_b64"
    except Exception as e:
        msg = str(e)
        if ("must be verified" in msg.lower() and "gpt-image" in msg.lower()) or ("verify organization" in msg.lower() and "gpt-image" in msg.lower()):
            return None, "org_verification_required"
        gen_err = f"{type(e).__name__}: {e}"
        if edit_err:
            return None, f"edit_failed={edit_err} | generate_failed={gen_err}"
        return None, gen_err

def sd_webui_img2img(
    init_png: bytes,
    prompt: str,
    url: str = "http://127.0.0.1:7860",
    negative_prompt: str = "",
    denoising_strength: float = 0.45,
    steps: int = 28,
    cfg_scale: float = 7.0,
    width: int = 1024,
    height: int = 1024,
    timeout_s: int = 180,
) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Local Stable Diffusion (AUTOMATIC1111) img2img backend.
    Requires the WebUI API enabled and reachable at {url}.
    Returns PNG bytes or error string.
    """
    try:
        import requests  # type: ignore
    except Exception:
        return None, "requests_not_installed"
    try:
        b64 = base64.b64encode(init_png).decode("utf-8")
        payload = {
            "init_images": [b64],
            "prompt": prompt,
            "negative_prompt": negative_prompt,
            "denoising_strength": float(denoising_strength),
            "steps": int(steps),
            "cfg_scale": float(cfg_scale),
            "width": int(width),
            "height": int(height),
        }
        api = url.rstrip("/") + "/sdapi/v1/img2img"
        r = requests.post(api, json=payload, timeout=timeout_s)
        if r.status_code != 200:
            return None, f"sd_webui_http_{r.status_code}: {r.text[:600]}"
        data = r.json()
        imgs = data.get("images") or []
        if not imgs:
            return None, "sd_webui_no_images_returned"
        out = base64.b64decode(imgs[0].split(",", 1)[-1])
        return out, None
    except Exception as e:
        return None, f"sd_webui_error: {type(e).__name__}: {e}"


def overlay_rule_based_edit(
    base_png: bytes,
    rep: Dict[str, Any],
    plan: Dict[str, Any],
    change_request: str,
) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Deterministic fallback that *modifies the raster* by replacing recognized dimension/note tokens
    using OCR bounding boxes and drawing overlays (no generative model).
    This preserves the baseline drawing and makes changes visible for comparison, but is limited
    by OCR quality and token localization.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
    except Exception:
        return None, "pillow_not_installed"

    try:
        img = Image.open(io.BytesIO(base_png)).convert("RGBA")
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default()

        # Prefer OCR words from representation
        ocr = rep.get("ocr") if isinstance(rep, dict) else None
        words = (ocr or {}).get("words") or []
        # Build a simple index by token
        token_map = {}
        for w in words:
            t = (w.get("text") or "").strip()
            if not t:
                continue
            token_map.setdefault(t, []).append(w)

        ops = (plan.get("change_plan", {}) if isinstance(plan, dict) else {}).get("operations", []) or []
        # If no ops, create a single callout overlay
        if not ops:
            ops = [{"type": "other", "from": "", "to": change_request, "target": "Diagram"}]

        # Utility to draw highlighted replacement over a bbox
        def _replace_bbox(b, new_text):
            x, y = int(b.get("left", 0)), int(b.get("top", 0))
            w, h = int(b.get("width", 0)), int(b.get("height", 0))
            if w <= 0 or h <= 0:
                return
            pad = 2
            # cover old text
            draw.rectangle([x - pad, y - pad, x + w + pad, y + h + pad], fill=(0, 0, 0, 255))
            # draw new text in red-ish accent
            draw.text((x, y), new_text, font=font, fill=(255, 70, 70, 255))
            # outline
            draw.rectangle([x - pad, y - pad, x + w + pad, y + h + pad], outline=(255, 70, 70, 255), width=2)

        # Heuristic token match: try exact 'from', else try numeric match
        changed = 0
        for op in ops:
            frm = (op.get("from") or "").strip()
            to = (op.get("to") or "").strip()
            if not to:
                continue

            matched = False
            if frm and frm in token_map:
                for b in token_map[frm][:3]:
                    _replace_bbox(b, to)
                    changed += 1
                    matched = True

            if not matched and frm:
                # numeric fuzzy match: compare digits only
                frm_digits = re.sub(r"[^0-9.]", "", frm)
                if frm_digits:
                    for t, bxs in token_map.items():
                        if re.sub(r"[^0-9.]", "", t) == frm_digits:
                            for b in bxs[:2]:
                                _replace_bbox(b, to)
                                changed += 1
                                matched = True
                            break

            if not matched:
                # Add a callout in the corner if we couldn't localize
                draw.rectangle([10, 10, 10 + 520, 10 + 60], fill=(0, 0, 0, 200))
                draw.text((20, 20), f"CHANGE: {to[:80]}", font=font, fill=(255, 70, 70, 255))
                changed += 1

        # Add a footer tag
        W, H = img.size
        draw.rectangle([0, H - 24, W, H], fill=(0, 0, 0, 220))
        draw.text((10, H - 20), f"Rule-based update (OCR-driven) | items_applied={changed}", font=font, fill=(255, 70, 70, 255))

        out = io.BytesIO()
        img.convert("RGB").save(out, format="PNG")
        return out.getvalue(), None
    except Exception as e:
        return None, f"overlay_edit_error: {type(e).__name__}: {e}"


def compute_impacts(G: nx.DiGraph, affected_parts: List[str], depth: int = 3) -> Dict[str, Any]:
    """
    Graph-propagate impacts around affected parts: upstream assemblies + downstream components + docs.
    affected_parts: list of part_numbers
    """
    affected_part_ids = []
    for pn in affected_parts:
        pid = get_part_by_number(pn)
        if pid:
            affected_part_ids.append(pid)

    impacts = {
        "affected_parts": affected_parts,
        "affected_part_ids": affected_part_ids,
        "impacted_parts": [],
        "impacted_docs": [],
        "upstream_assemblies": [],
        "downstream_components": [],
        "impact_summary": {},
    }

    if not affected_part_ids or len(G.nodes) == 0:
        return impacts

    UG = G.to_undirected()
    impacted_nodes = set()
    for pid in affected_part_ids:
        neighborhood = nx.single_source_shortest_path_length(UG, pid, cutoff=depth)
        impacted_nodes |= set(neighborhood.keys())

    # categorize
    parts = []
    docs = []
    for n in impacted_nodes:
        t = G.nodes[n].get("type", "")
        if t == "part":
            parts.append(n)
        elif t == "doc":
            docs.append(n)

    # upstream assemblies: nodes that have edges to affected nodes in BOM sense
    upstream = set()
    downstream = set()
    for pid in affected_part_ids:
        upstream |= set(G.predecessors(pid))  # parent assemblies + docs mentioning
        downstream |= set(G.successors(pid))  # child components

    # filter to parts
    upstream_parts = [u for u in upstream if G.nodes.get(u, {}).get("type") == "part"]
    downstream_parts = [d for d in downstream if G.nodes.get(d, {}).get("type") == "part"]

    # translate ids -> part_numbers, doc filenames
    con = db()
    part_map = dict(pd.read_sql_query("SELECT part_id, part_number FROM parts", con).values.tolist())
    doc_map = dict(pd.read_sql_query("SELECT doc_id, filename FROM documents", con).values.tolist())
    con.close()

    impacts["impacted_parts"] = sorted({part_map.get(p, p) for p in parts})
    impacts["upstream_assemblies"] = sorted({part_map.get(p, p) for p in upstream_parts})
    impacts["downstream_components"] = sorted({part_map.get(p, p) for p in downstream_parts})
    impacts["impacted_docs"] = sorted({doc_map.get(d, d) for d in docs})

    impacts["impact_summary"] = {
        "n_affected": len(affected_part_ids),
        "n_impacted_parts": len(impacts["impacted_parts"]),
        "n_upstream_assemblies": len(impacts["upstream_assemblies"]),
        "n_downstream_components": len(impacts["downstream_components"]),
        "n_impacted_docs": len(impacts["impacted_docs"]),
    }

    return impacts


def rule_based_risk_fmea(intent: Dict[str, Any], impacts: Dict[str, Any]) -> Dict[str, Any]:
    """
    Lightweight DFMEA/PFMEA starter:
    - Severity/Occurrence/Detection heuristics based on “change type” + impact breadth + compliance flags.
    """
    change_category = (intent.get("change_category") or "Design").strip()
    compliance = intent.get("compliance_context", [])
    breadth = safe_int(impacts.get("impact_summary", {}).get("n_impacted_parts", 0), 0)

    base_sev = 7 if change_category.lower() in ("design", "material", "interface") else 5
    if compliance:
        base_sev = min(10, base_sev + 2)
    if breadth >= 10:
        base_sev = min(10, base_sev + 1)

    base_occ = 5 if breadth >= 10 else 3
    base_det = 6 if change_category.lower() in ("process", "manufacturing") else 5

    fmea_rows = []
    affected = impacts.get("affected_parts", []) or ["(unspecified)"]
    for pn in affected[:25]:
        fmea_rows.append({
            "item": pn,
            "function": intent.get("objective", "Meets design intent"),
            "failure_mode": "Does not meet updated requirement / tolerance",
            "effect": "Performance degradation, rework, or nonconformance",
            "severity": base_sev,
            "cause": "Design change introduces new tolerance stack / interface mismatch",
            "occurrence": base_occ,
            "current_controls": "Peer review, drawing check, prototype build",
            "detection": base_det,
            "rpn": base_sev * base_occ * base_det,
            "recommended_actions": "Update validation plan; add inspection step; supplier communication; update work instructions",
        })

    # simple process readiness / industrialization checks
    readiness = {
        "manufacturing_feasibility": "Review required" if change_category.lower() in ("design", "material") else "Likely feasible",
        "tooling_impact": "Possible" if breadth >= 5 else "Low",
        "supplier_impact": "Possible" if intent.get("supplier_change") else "Unknown",
        "schedule_risk": "Medium" if breadth >= 10 else "Low",
        ".have_prototype_plan": bool(intent.get("validation_strategy")),
    }

    return {
        "summary": {
            "change_category": change_category,
            "baseline_severity": base_sev,
            "baseline_occurrence": base_occ,
            "baseline_detection": base_det,
        },
        "dfmea_pfmea": fmea_rows,
        "industrialization_readiness": readiness,
    }


def generate_acceptance_criteria(intent: Dict[str, Any], impacts: Dict[str, Any]) -> List[Dict[str, Any]]:
    criteria = []
    affected = impacts.get("affected_parts", []) or ["(unspecified)"]
    obj = intent.get("objective", "Meets updated design intent")
    for pn in affected[:20]:
        criteria.append({
            "part_number": pn,
            "criterion": f"{obj} for {pn} under defined operating conditions",
            "method": "Inspection + functional test",
            "evidence": "Test report + inspection record",
            "owner": intent.get("validation_owner", intent.get("requester", "Engineering")),
        })
    if intent.get("compliance_context"):
        criteria.append({
            "part_number": "(system)",
            "criterion": f"Compliance evidence updated: {', '.join(intent.get('compliance_context'))}",
            "method": "Document review",
            "evidence": "Updated compliance matrix / certificate",
            "owner": "Quality/Compliance",
        })
    return criteria


def build_change_package(intent: Dict[str, Any], impacts: Dict[str, Any], retrievals: List[RetrievalResult]) -> Dict[str, Any]:
    """
    Produces the “comprehensive change package” as structured JSON, with traceability.
    """
    # Optional: try LLM enrichment first
    schema_hint = """{
  "change_request_definition": { "problem_statement": "", "objective": "", "scope": "", "affected_parts": [], "assumptions": [], "constraints": [] },
  "impact_assessment": { "parts": [], "assemblies": [], "processes": [], "cost": { "estimate": 0, "drivers": [] }, "schedule": { "estimate_weeks": 0, "drivers": [] }, "compliance": [] },
  "risk_fmea": { "dfmea_pfmea": [ { "item": "", "function": "", "failure_mode": "", "effect": "", "severity": 1, "cause": "", "occurrence": 1, "current_controls": "", "detection": 1, "rpn": 1, "recommended_actions": "" } ] },
  "release_workflow": { "gates": [], "roles": {}, "artifacts": [] },
  "industrialization_readiness": { "manufacturing_feasibility": "", "tooling_impact": "", "supplier_impact": "", "schedule_risk": "" },
  "acceptance_validation": [ { "criterion": "", "method": "", "evidence": "" } ],
  "traceability": { "sources": [ { "doc_id": "", "chunk_id": "", "note": "" } ] }
}"""

    system = (
        "You are an expert engineering change control assistant. "
        "You must be conservative, traceable, and configuration-aware. "
        "Never invent part numbers; only use part numbers from the provided context."
    )

    context = {
        "intent": intent,
        "impacts": impacts,
        "top_evidence_chunks": [
            {"chunk_id": r.chunk_id, "doc_id": r.doc_id, "score": r.score, "text": r.text[:900]}
            for r in retrievals
        ]
    }

    llm = llm_generate_json(
        system=system,
        user=f"Context:\n{json_dumps_pretty(context)}\n\nGenerate a complete engineering change package.",
        schema_hint=schema_hint,
        temperature=0.2
    )

    if llm is None or not isinstance(llm, dict):
        # deterministic fallback
        risk = rule_based_risk_fmea(intent, impacts)
        acceptance = generate_acceptance_criteria(intent, impacts)
        pkg = {
            "change_request_definition": {
                "problem_statement": intent.get("problem_statement", ""),
                "objective": intent.get("objective", ""),
                "scope": intent.get("scope", ""),
                "affected_parts": impacts.get("affected_parts", []),
                "assumptions": intent.get("assumptions", []),
                "constraints": intent.get("constraints", []),
            },
            "impact_assessment": {
                "parts": impacts.get("impacted_parts", []),
                "assemblies": impacts.get("upstream_assemblies", []),
                "processes": intent.get("impacted_processes", []),
                "cost": {
                    "estimate": safe_float(intent.get("cost_estimate", 0.0), 0.0),
                    "drivers": ["BOM updates", "rework", "tooling changes", "supplier updates"],
                },
                "schedule": {
                    "estimate_weeks": safe_int(intent.get("schedule_weeks", 0), 0),
                    "drivers": ["review cycle", "prototype build", "test execution", "supplier lead time"],
                },
                "compliance": intent.get("compliance_context", []),
            },
            "risk_fmea": risk,
            "release_workflow": {
                "gates": DEFAULT_WORKFLOW.get("approval_steps", []),
                "roles": DEFAULT_WORKFLOW.get("roles", {}),
                "artifacts": [
                    "ECR/ECO record",
                    "Updated drawing/procedure pack",
                    "Validation plan + evidence references",
                    "Release checklist",
                ],
            },
            "industrialization_readiness": risk.get("industrialization_readiness", {}),
            "acceptance_validation": acceptance,
            "traceability": {
                "sources": [{"doc_id": r.doc_id, "chunk_id": r.chunk_id, "note": "Retrieved evidence"} for r in retrievals]
            }
        }
        return pkg

    # ensure traceability exists even if model omitted
    llm.setdefault("traceability", {})
    llm["traceability"].setdefault("sources", [{"doc_id": r.doc_id, "chunk_id": r.chunk_id, "note": "Retrieved evidence"} for r in retrievals])
    return llm


# ----------------------------
# Visual diff utilities
# ----------------------------
def image_diff(a_bytes: bytes, b_bytes: bytes) -> Optional[bytes]:
    if Image is None:
        return None
    ia = read_image(a_bytes)
    ib = read_image(b_bytes)
    if ia is None or ib is None:
        return None
    # resize to same
    w = max(ia.size[0], ib.size[0])
    h = max(ia.size[1], ib.size[1])
    ia2 = ia.resize((w, h))
    ib2 = ib.resize((w, h))
    diff = ImageChops.difference(ia2, ib2)
    # enhance diff for visibility
    diff = ImageEnhance.Contrast(diff).enhance(2.0)
    out = io.BytesIO()
    diff.save(out, format="PNG")
    return out.getvalue()


def _ensure_same_size(ia: "Image.Image", ib: "Image.Image") -> Tuple["Image.Image", "Image.Image"]:
    w = max(ia.size[0], ib.size[0])
    h = max(ia.size[1], ib.size[1])
    return ia.resize((w, h)), ib.resize((w, h))


def _to_gray_np(img: "Image.Image") -> np.ndarray:
    return np.array(img.convert("L"), dtype=np.int16)


def compute_visual_metrics(ia: Any, ib: Any, threshold: int = 25) -> Dict[str, Any]:
    """Transparent, data-driven metrics for visual change + similarity.

    Accepts PIL Images or raw image bytes.
    """
    if Image is None:
        return {"available": False, "error": "PIL (Pillow) is not available."}

    ia_img = read_image(bytes(ia)) if isinstance(ia, (bytes, bytearray)) else ia
    ib_img = read_image(bytes(ib)) if isinstance(ib, (bytes, bytearray)) else ib

    if ia_img is None or ib_img is None:
        return {"available": False, "error": "Could not decode one or both images for visual metrics."}

    ia2, ib2 = _ensure_same_size(ia_img, ib_img)
    a = _to_gray_np(ia2)
    b = _to_gray_np(ib2)
    diff = np.abs(a - b)
    changed = (diff > int(threshold)).astype(np.float32)
    change_ratio = float(changed.mean())
    mse = float(np.mean((a.astype(np.float32) - b.astype(np.float32)) ** 2))
    nmse = float(mse / (255.0 ** 2))

    ssim_val = None
    if SKIMAGE_AVAILABLE and skimage_ssim is not None:
        try:
            ssim_val = float(skimage_ssim(a.astype(np.uint8), b.astype(np.uint8)))
        except Exception:
            ssim_val = None

    return {
        "available": True,
        "resolution": {"width": ia2.size[0], "height": ia2.size[1]},
        "threshold": int(threshold),
        "pixel_change_ratio": change_ratio,
        "mse": mse,
        "nmse": nmse,
        "ssim": ssim_val,
    }


def make_overlay_png(ia: "Image.Image", ib: "Image.Image", alpha: float = 0.5) -> bytes:
    ia2, ib2 = _ensure_same_size(ia, ib)
    alpha = float(max(0.0, min(1.0, alpha)))
    out_img = Image.blend(ia2, ib2, alpha)
    out = io.BytesIO()
    out_img.save(out, format="PNG")
    return out.getvalue()


def compute_diff_png(ia: "Image.Image", ib: "Image.Image", contrast: float = 2.0) -> bytes:
    ia2, ib2 = _ensure_same_size(ia, ib)
    diff = ImageChops.difference(ia2, ib2)
    diff = ImageEnhance.Contrast(diff).enhance(float(contrast))
    out = io.BytesIO()
    diff.save(out, format="PNG")
    return out.getvalue()


def tile_change_analysis(ia: "Image.Image", ib: "Image.Image", tile: int = 96, threshold: int = 25, top_k: int = 25) -> Tuple[pd.DataFrame, bytes]:
    """Quantify % changed pixels per tile and highlight the highest-change regions."""
    ia2, ib2 = _ensure_same_size(ia, ib)
    a = _to_gray_np(ia2)
    b = _to_gray_np(ib2)
    diff = np.abs(a - b)
    mask = (diff > int(threshold)).astype(np.uint8)

    W, H = ia2.size
    tile = int(max(32, min(512, tile)))

    rows = []
    for y0 in range(0, H, tile):
        for x0 in range(0, W, tile):
            y1 = min(H, y0 + tile)
            x1 = min(W, x0 + tile)
            region = mask[y0:y1, x0:x1]
            ratio = float(region.mean()) if region.size else 0.0
            if ratio > 0:
                rows.append({"x0": x0, "y0": y0, "x1": x1, "y1": y1, "change_ratio": ratio})

    df = pd.DataFrame(rows).sort_values("change_ratio", ascending=False).head(int(top_k))
    ann = ia2.copy()
    draw = ImageDraw.Draw(ann)
    for _, r in df.iterrows():
        draw.rectangle([int(r.x0), int(r.y0), int(r.x1), int(r.y1)], outline=(225, 29, 72), width=3)
    out = io.BytesIO()
    ann.save(out, format="PNG")
    return df, out.getvalue()


def zoom_crop(img: "Image.Image", zoom: float = 1.0, cx: float = 0.5, cy: float = 0.5) -> "Image.Image":
    """Deep zoom via crop + resize (pan via cx/cy in [0..1])."""
    zoom = float(max(1.0, min(20.0, zoom)))
    cx = float(max(0.0, min(1.0, cx)))
    cy = float(max(0.0, min(1.0, cy)))

    W, H = img.size
    cw = max(1, int(W / zoom))
    ch = max(1, int(H / zoom))
    x0 = int(cx * W - cw / 2)
    y0 = int(cy * H - ch / 2)
    x0 = max(0, min(W - cw, x0))
    y0 = max(0, min(H - ch, y0))
    crop = img.crop((x0, y0, x0 + cw, y0 + ch))
    out_w = min(W, 1400)
    out_h = max(1, int(out_w * (ch / cw)))
    return crop.resize((out_w, out_h), resample=Image.BICUBIC)


DIM_RX = re.compile(r"(?P<val>\d+(?:\.\d+)?)\s*(?P<uom>mm|MM|cm|CM|m|M|in|IN|inch|INCH|\")\b")
ANG_RX = re.compile(r"(?P<val>\d+(?:\.\d+)?)\s*(?:deg|DEG|°)\b")
TOL_RX = re.compile(r"(±\s*\d+(?:\.\d+)?|\+\s*\d+(?:\.\d+)?\s*/\s*-\s*\d+(?:\.\d+)?)")
DIA_RX = re.compile(r"(?:Ø|⌀)\s*(\d+(?:\.\d+)?)")
THREAD_RX = re.compile(r"\b(M\d+(?:x\d+(?:\.\d+)?)?)\b", re.IGNORECASE)


def extract_engineering_attributes(text: str) -> Dict[str, List[str]]:
    t = (text or "")
    dims = [f"{m.group('val')}{m.group('uom')}" for m in DIM_RX.finditer(t)]
    angs = [m.group('val') for m in ANG_RX.finditer(t)]
    tols = [m.group(1).replace(" ", "") for m in TOL_RX.finditer(t)]
    dias = [m.group(1) for m in DIA_RX.finditer(t)]
    threads = [m.group(1).upper() for m in THREAD_RX.finditer(t)]

    def uniq(xs):
        seen=set(); out=[]
        for x in xs:
            if x not in seen:
                seen.add(x); out.append(x)
        return out

    return {
        "dimensions": uniq(dims)[:500],
        "angles": uniq(angs)[:500],
        "tolerances": uniq(tols)[:500],
        "diameters": uniq(dias)[:500],
        "threads": uniq(threads)[:500],
    }


def diff_attributes(base: Dict[str, List[str]], mod: Dict[str, List[str]], base_conf: float, mod_conf: float) -> pd.DataFrame:
    rows = []
    for k in ["dimensions", "angles", "tolerances", "diameters", "threads"]:
        b = set(base.get(k, []) or [])
        m = set(mod.get(k, []) or [])
        for x in sorted(m - b):
            rows.append({"type": k, "change": "added", "baseline": "", "modified": x, "confidence": round(min(base_conf, mod_conf), 3)})
        for x in sorted(b - m):
            rows.append({"type": k, "change": "removed", "baseline": x, "modified": "", "confidence": round(min(base_conf, mod_conf), 3)})
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    return df.sort_values(["type", "change"])


# ----------------------------
# Predictive model (optional) trained from imported change outcomes
# ----------------------------
def import_change_history_training(df: pd.DataFrame) -> int:
    """
    Expected columns:
      change_id, text, outcome_delay, outcome_quality_issue
    """
    cols = {c.lower().strip(): c for c in df.columns}
    req = ["change_id", "text", "outcome_delay", "outcome_quality_issue"]
    for r in req:
        if r not in cols:
            raise ValueError(f"Missing required column: {r}")
    con = db()
    n = 0
    for _, row in df.iterrows():
        con.execute("""
            INSERT INTO change_history_training(change_id, text, outcome_delay, outcome_quality_issue, created_at)
            VALUES(?,?,?,?,?)
        """, (
            str(row[cols["change_id"]]),
            str(row[cols["text"]]),
            safe_int(row[cols["outcome_delay"]], 0),
            safe_int(row[cols["outcome_quality_issue"]], 0),
            now_iso()
        ))
        n += 1
    con.commit()
    con.close()
    return n


def train_predictor() -> Optional[Dict[str, Any]]:
    if TfidfVectorizer is None or LogisticRegression is None or train_test_split is None:
        return None
    con = db()
    df = pd.read_sql_query("SELECT text, outcome_delay, outcome_quality_issue FROM change_history_training", con)
    con.close()
    if len(df) < 30:
        return None
    X_text = df["text"].fillna("").tolist()

    vec = TfidfVectorizer(stop_words="english", max_features=8000)
    X = vec.fit_transform(X_text)

    # two models: delay and quality
    out = {}
    for target in ["outcome_delay", "outcome_quality_issue"]:
        y = df[target].astype(int).values
        Xtr, Xte, ytr, yte = train_test_split(X, y, test_size=0.25, random_state=42, stratify=y if len(set(y)) > 1 else None)
        clf = LogisticRegression(max_iter=1000)
        clf.fit(Xtr, ytr)
        proba = clf.predict_proba(Xte)[:, 1] if len(clf.classes_) == 2 else np.zeros_like(yte, dtype=float)
        auc = None
        try:
            auc = roc_auc_score(yte, proba) if len(set(yte)) > 1 else None
        except Exception:
            auc = None
        out[target] = {"model": clf, "vectorizer": vec, "auc": auc}
    return out


def predict_risks(models: Dict[str, Any], text: str) -> Dict[str, float]:
    if not models:
        return {}
    out = {}
    for target, pack in models.items():
        vec = pack["vectorizer"]
        clf = pack["model"]
        X = vec.transform([text or ""])
        if hasattr(clf, "predict_proba") and len(getattr(clf, "classes_", [])) == 2:
            out[target] = float(clf.predict_proba(X)[0, 1])
        else:
            out[target] = 0.0
    return out


# ----------------------------
# Export artifacts (DOCX/PDF + ZIP)
# ----------------------------
def write_docx_change_package(path: str, change: Dict[str, Any], pkg: Dict[str, Any]) -> None:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    doc.add_heading(f"{change['change_type']} — {change['title']}", level=0)
    doc.add_paragraph(f"Change ID: {change['change_id']}")
    doc.add_paragraph(f"Status: {change['status']}")
    doc.add_paragraph(f"Requester: {change['requester']} | Owner: {change['owner']}")
    doc.add_paragraph(f"Created: {change['created_at']} | Updated: {change['updated_at']}")

    doc.add_heading("1. Change Request Definition", level=1)
    crd = pkg.get("change_request_definition", {})
    for k in ["problem_statement", "objective", "scope"]:
        doc.add_paragraph(f"{k.replace('_',' ').title()}: {crd.get(k,'')}")
    doc.add_paragraph(f"Affected parts: {', '.join(crd.get('affected_parts', []) or [])}")

    doc.add_heading("2. Impact Assessment", level=1)
    ia = pkg.get("impact_assessment", {})
    doc.add_paragraph(f"Impacted parts: {', '.join(ia.get('parts', []) or [])}")
    doc.add_paragraph(f"Upstream assemblies: {', '.join(ia.get('assemblies', []) or [])}")
    doc.add_paragraph(f"Processes: {', '.join(ia.get('processes', []) or [])}")
    doc.add_paragraph(f"Cost estimate: {ia.get('cost', {}).get('estimate', 0)}")
    doc.add_paragraph(f"Schedule estimate (weeks): {ia.get('schedule', {}).get('estimate_weeks', 0)}")
    comp = ia.get("compliance", []) or []
    if comp:
        doc.add_paragraph(f"Compliance: {', '.join(comp)}")

    doc.add_heading("3. Risk & FMEA", level=1)
    rf = pkg.get("risk_fmea", {})
    rows = rf.get("dfmea_pfmea", rf.get("dfmea_pfmea", [])) if isinstance(rf, dict) else []
    if isinstance(rf, dict) and "dfmea_pfmea" in rf:
        rows = rf["dfmea_pfmea"]
    # table
    if rows:
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Failure mode"
        hdr[2].text = "Effect"
        hdr[3].text = "S/O/D"
        hdr[4].text = "RPN"
        hdr[5].text = "Recommended actions"
        for r in rows[:25]:
            c = table.add_row().cells
            c[0].text = str(r.get("item", ""))
            c[1].text = str(r.get("failure_mode", ""))
            c[2].text = str(r.get("effect", ""))
            c[3].text = f"{r.get('severity','')}/{r.get('occurrence','')}/{r.get('detection','')}"
            c[4].text = str(r.get("rpn", ""))
            c[5].text = str(r.get("recommended_actions", ""))

    doc.add_heading("4. Industrialization & Manufacturing Readiness", level=1)
    ir = pkg.get("industrialization_readiness", {}) or rf.get("industrialization_readiness", {})
    if isinstance(ir, dict):
        for k, v in ir.items():
            doc.add_paragraph(f"{k.replace('_',' ').title()}: {v}")

    doc.add_heading("5. Acceptance & Validation Criteria", level=1)
    av = pkg.get("acceptance_validation", []) or []
    for r in av[:40]:
        doc.add_paragraph(f"- {r.get('criterion','')} | Method: {r.get('method','')} | Evidence: {r.get('evidence','')}")

    doc.add_heading("6. Traceability (Evidence Sources)", level=1)
    tr = pkg.get("traceability", {}).get("sources", []) if isinstance(pkg.get("traceability", {}), dict) else []
    for s in tr[:50]:
        doc.add_paragraph(f"- doc={s.get('doc_id','')} chunk={s.get('chunk_id','')} : {s.get('note','')}")

    doc.save(path)


def write_pdf_change_summary(path: str, change: Dict[str, Any], pkg: Dict[str, Any]) -> None:
    if canvas is None or A4 is None:
        raise RuntimeError("reportlab not installed")
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    x = 40
    y = h - 50

    def line(txt: str, dy: int = 14, bold: bool = False) -> None:
        nonlocal y
        if y < 60:
            c.showPage()
            y = h - 50
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 11 if bold else 9)
        c.drawString(x, y, (txt or "")[:120])
        y -= dy

    line(f"{change['change_type']} — {change['title']}", bold=True, dy=18)
    line(f"Change ID: {change['change_id']}   Status: {change['status']}")
    line(f"Requester: {change['requester']}   Owner: {change['owner']}")
    line(f"Created: {change['created_at']}   Updated: {change['updated_at']}")
    line("", dy=10)

    crd = pkg.get("change_request_definition", {})
    line("Change Request Definition", bold=True, dy=16)
    line(f"Problem: {crd.get('problem_statement','')}", dy=12)
    line(f"Objective: {crd.get('objective','')}", dy=12)
    line(f"Scope: {crd.get('scope','')}", dy=12)
    line(f"Affected parts: {', '.join(crd.get('affected_parts', []) or [])}", dy=12)

    ia = pkg.get("impact_assessment", {})
    line("", dy=10)
    line("Impact Summary", bold=True, dy=16)
    line(f"Impacted parts: {len(ia.get('parts', []) or [])}")
    line(f"Upstream assemblies: {len(ia.get('assemblies', []) or [])}")
    line(f"Processes: {len(ia.get('processes', []) or [])}")
    line(f"Cost estimate: {ia.get('cost', {}).get('estimate', 0)}")
    line(f"Schedule estimate (weeks): {ia.get('schedule', {}).get('estimate_weeks', 0)}")

    line("", dy=10)
    line("Acceptance & Validation (high level)", bold=True, dy=16)
    av = pkg.get("acceptance_validation", []) or []
    for r in av[:10]:
        line(f"- {r.get('criterion','')}", dy=11)

    line("", dy=10)
    line("Traceability", bold=True, dy=16)
    tr = pkg.get("traceability", {}).get("sources", []) if isinstance(pkg.get("traceability", {}), dict) else []
    for s in tr[:12]:
        line(f"- {s.get('doc_id','')} / {s.get('chunk_id','')}", dy=11)

    c.save()


def export_change_package_zip(change_id: str) -> str:
    change = get_change(change_id)
    if not change:
        raise ValueError("Change not found")

    pkg = change.get("impacted", {}) or {}
    risk = change.get("risk", {}) or {}
    evidence = change.get("evidence", {}) or {}

    # unify package
    unified = {
        "intent": change.get("intent", {}),
        "package": pkg,
        "risk": risk,
        "evidence": evidence,
    }

    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = os.path.join(EXPORTS_DIR, change_id)
    os.makedirs(out_dir, exist_ok=True)
    json_path = os.path.join(out_dir, f"{change_id}_{stamp}.json")
    with open(json_path, "w", encoding="utf-8") as f:
        f.write(json_dumps_pretty(unified))

    docx_path = os.path.join(out_dir, f"{change_id}_{stamp}.docx")
    pdf_path = os.path.join(out_dir, f"{change_id}_{stamp}.pdf")
    if Document is not None:
        write_docx_change_package(docx_path, change, pkg.get("package", pkg) if isinstance(pkg, dict) else {})
    if canvas is not None:
        write_pdf_change_summary(pdf_path, change, pkg.get("package", pkg) if isinstance(pkg, dict) else {})

    zip_path = os.path.join(out_dir, f"{change_id}_{stamp}_export.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.write(json_path, arcname=os.path.basename(json_path))
        if os.path.exists(docx_path):
            z.write(docx_path, arcname=os.path.basename(docx_path))
        if os.path.exists(pdf_path):
            z.write(pdf_path, arcname=os.path.basename(pdf_path))

        # attach evidence docs (optional)
        doc_ids = []
        if isinstance(evidence, dict):
            doc_ids = evidence.get("docs", []) or []
        for doc_id in doc_ids[:50]:
            doc = get_document(doc_id)
            if doc and doc.get("stored_path") and os.path.exists(doc["stored_path"]):
                z.write(doc["stored_path"], arcname=f"evidence/{os.path.basename(doc['stored_path'])}")

    return zip_path


# ----------------------------
# UI
# ----------------------------
def sidebar_identity() -> Dict[str, str]:
    st.sidebar.markdown("### Identity & Settings")
    user = st.sidebar.text_input("Your name / user", value=st.session_state.get("user", "David"))
    role = st.sidebar.selectbox("Role", ["Requester", "Change Owner", "Approver"], index=0)
    st.session_state["user"] = user

    st.sidebar.markdown("---")
    st.sidebar.markdown("### PDF Extraction")
    force_pdf_ocr = st.sidebar.checkbox(
        "Force OCR for PDFs (for scanned drawings)",
        value=bool(st.session_state.get("force_pdf_ocr", False))
    )
    st.session_state["force_pdf_ocr"] = bool(force_pdf_ocr)
    pdf_ocr_pages = st.sidebar.slider(
        "OCR pages (first N pages)",
        min_value=1, max_value=6,
        value=int(st.session_state.get("pdf_ocr_pages", 1))
    )
    st.session_state["pdf_ocr_pages"] = int(pdf_ocr_pages)
    if force_pdf_ocr:
        st.sidebar.caption("OCR requires Tesseract installed locally. On Windows, set env var TESSERACT_CMD if needed.")

    st.sidebar.markdown("---")
    
    st.sidebar.markdown("### LLM Routing (optional)")

    provider_label = st.sidebar.selectbox(
        "Provider",
        ["OpenAI", "Mistral", "Auto (OpenAI → Mistral)"],
        index=["OpenAI", "Mistral", "Auto (OpenAI → Mistral)"].index(st.session_state.get("llm_provider", "OpenAI")) if st.session_state.get("llm_provider", "OpenAI") in ["OpenAI", "Mistral", "Auto (OpenAI → Mistral)"] else 0,
        help="Choose which LLM powers text generation in the AI Assistant and procedure generation. Auto tries OpenAI first, then Mistral.",
    )
    st.session_state["llm_provider"] = provider_label

    # --- OpenAI ---
    with st.sidebar.expander("OpenAI settings", expanded=(provider_label.startswith("OpenAI") or provider_label.startswith("Auto"))):
        if OPENAI_AVAILABLE:
            key = st.text_input("OpenAI API key", type="password", value=st.session_state.get("openai_api_key", ""))
            st.session_state["openai_api_key"] = key
            model = st.text_input("OpenAI model", value=st.session_state.get("openai_model", DEFAULT_MODEL))
            st.session_state["openai_model"] = model
            st.caption("Uses OpenAI Responses API via `from openai import OpenAI`.")
        else:
            st.info("OpenAI SDK not installed. `pip install openai` to enable OpenAI routing.")

    # --- Mistral ---
    with st.sidebar.expander("Mistral settings", expanded=(provider_label.startswith("Mistral") or provider_label.startswith("Auto"))):
        if REQUESTS_AVAILABLE:
            mkey = st.text_input("Mistral API key", type="password", value=st.session_state.get("mistral_api_key", ""))
            st.session_state["mistral_api_key"] = mkey
            mmodel = st.text_input("Mistral model", value=st.session_state.get("mistral_model", DEFAULT_MISTRAL_MODEL))
            st.session_state["mistral_model"] = mmodel
            st.caption("Uses Mistral Chat Completions API (`/v1/chat/completions`).")
        else:
            st.info("`requests` is not available. Install it to enable Mistral routing: `pip install requests`.")

    return {"user": user, "role": role}




# ----------------------------
# Session persistence & archive (keeps user inputs across navigation + allows snapshots)
# ----------------------------

SESSION_KEYS_EXCLUDE = {
    "openai_api_key",  # never archive secrets
    "mistral_api_key",  # never archive secrets
}


def _is_internal_session_key(k: str) -> bool:
    if not k:
        return True
    if k.startswith("_"):
        return True
    # Streamlit internal widget bookkeeping can be noisy
    if k.startswith("FormSubmitter"):
        return True
    return False


def _archive_folder_for(archive_id: str) -> str:
    ensure_dirs()
    folder = os.path.join(ARCHIVE_DIR, "session_archives", archive_id)
    os.makedirs(folder, exist_ok=True)
    return folder


def _persist_bytes(folder: str, key: str, b: bytes) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_\-\.]", "_", key)[:80]
    path = os.path.join(folder, f"{safe}.bin")
    with open(path, "wb") as f:
        f.write(b)
    return path


def _session_snapshot_payload(archive_id: str) -> Dict[str, Any]:
    """
    Snapshot *all* user-facing session keys for robust persistence across navigation.
    Bytes are persisted to an archive folder on disk (no size restriction).
    """
    folder = _archive_folder_for(archive_id)
    payload: Dict[str, Any] = {"version": "2.0", "archive_id": archive_id, "created_at": dt.datetime.utcnow().isoformat(), "keys": {}}

    for k in list(st.session_state.keys()):
        if k in SESSION_KEYS_EXCLUDE or _is_internal_session_key(k):
            continue
        v = st.session_state.get(k)
        if isinstance(v, (bytes, bytearray)):
            try:
                path = _persist_bytes(folder, k, bytes(v))
                payload["keys"][k] = {"__type__": "bytes_file", "path": path}
            except Exception:
                payload["keys"][k] = {"__type__": "bytes_file", "path": ""}
            continue
        # DataFrame support
        if isinstance(v, pd.DataFrame):
            try:
                payload["keys"][k] = {"__type__": "df_json", "data": v.to_json(orient="split")}
            except Exception:
                payload["keys"][k] = {"__type__": "str", "data": str(v)}
            continue
        # generic json-serializable
        try:
            json.dumps(v, ensure_ascii=False)
            payload["keys"][k] = v
        except Exception:
            payload["keys"][k] = {"__type__": "str", "data": str(v)}

    # also persist a pointer to any DCU archives folder
    payload["keys"]["dcu_archive_folder"] = st.session_state.get("dcu_archive_folder", "")
    return payload


def _apply_session_snapshot(payload: Dict[str, Any]) -> None:
    keys = (payload or {}).get("keys", {})
    if not isinstance(keys, dict):
        return
    for k, v in keys.items():
        if k in SESSION_KEYS_EXCLUDE or _is_internal_session_key(k):
            continue
        try:
            if isinstance(v, dict) and v.get("__type__") == "bytes_file":
                path = v.get("path") or ""
                if path and os.path.exists(path):
                    with open(path, "rb") as f:
                        st.session_state[k] = f.read()
                continue
            if isinstance(v, dict) and v.get("__type__") == "df_json":
                st.session_state[k] = pd.read_json(v.get("data", "{}"), orient="split")
                continue
            if isinstance(v, dict) and v.get("__type__") == "str":
                st.session_state[k] = v.get("data", "")
                continue
            st.session_state[k] = v
        except Exception:
            pass


def _encode_payload_for_db(payload: Dict[str, Any]) -> str:
    raw = json.dumps(payload, ensure_ascii=False)
    # compress if very large (keeps DB stable)
    if len(raw.encode("utf-8")) > 3_000_000:
        gz = gzip.compress(raw.encode("utf-8"))
        return json.dumps({"__compressed__": True, "data_b64": base64.b64encode(gz).decode("utf-8")})
    return raw


def _decode_payload_from_db(s: str) -> Dict[str, Any]:
    try:
        js = json.loads(s)
        if isinstance(js, dict) and js.get("__compressed__") and js.get("data_b64"):
            gz = base64.b64decode(js["data_b64"])
            raw = gzip.decompress(gz).decode("utf-8")
            return json.loads(raw)
        return js if isinstance(js, dict) else {}
    except Exception:
        return {}


def archive_current_session(name: str) -> Optional[str]:
    try:
        archive_id = f"arc_{dt.datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
        # archive any generated artifacts to disk first
        folder = dcu_archive_outputs()
        if folder:
            st.session_state["dcu_archive_folder"] = folder

        payload = _session_snapshot_payload(archive_id)
        payload_json = _encode_payload_for_db(payload)

        def _write():
            con = db()
            con.execute(
                "INSERT INTO session_archives(archive_id, name, payload_json, created_at) VALUES (?,?,?,?)",
                (archive_id, (name or "Session").strip(), payload_json, dt.datetime.utcnow().isoformat()),
            )
            con.commit()
            con.close()
        _db_write_retry(_write)
        return archive_id
    except Exception:
        return None


def list_session_archives() -> pd.DataFrame:
    try:
        con = db()
        df = pd.read_sql_query("SELECT archive_id, name, created_at FROM session_archives ORDER BY created_at DESC", con)
        con.close()
        return df
    except Exception:
        return pd.DataFrame(columns=["archive_id", "name", "created_at"])


def load_session_archive(archive_id: str) -> bool:
    try:
        con = db()
        row = con.execute("SELECT payload_json FROM session_archives WHERE archive_id=?", (archive_id,)).fetchone()
        con.close()
        if not row or not row[0]:
            return False
        payload = _decode_payload_from_db(row[0])
        _apply_session_snapshot(payload)
        return True
    except Exception:
        return False


def delete_session_archive(archive_id: str) -> bool:
    try:
        def _write():
            con = db()
            con.execute("DELETE FROM session_archives WHERE archive_id=?", (archive_id,))
            con.commit()
            con.close()
        _db_write_retry(_write)
        # also delete archive folder
        folder = os.path.join(ARCHIVE_DIR, "session_archives", archive_id)
        if os.path.exists(folder):
            try:
                import shutil
                shutil.rmtree(folder)
            except Exception:
                pass
        return True
    except Exception:
        return False


def session_persistence_sidebar() -> None:
    st.session_state.setdefault("user", "David")
    # --- BI chart workspace (BOM chart export/arrange) ---
    st.session_state.setdefault("bi_chart_store", {})          # id -> record
    st.session_state.setdefault("bi_dashboard_order", [])      # list of ids in dashboard
    st.session_state.setdefault("bi_archived_ids", [])         # list of ids archived

    with st.sidebar.expander("Session (persist + archive)", expanded=False):
        st.caption("Inputs persist across navigation. You can snapshot (archive) the full working session and restore later.")
        name = st.text_input("Archive name", value=f"{st.session_state.get('user','User')} — working session", key="session_archive_name")
        c1, c2 = st.columns(2)
        if c1.button("Archive now", key="btn_archive_session"):
            arc_id = archive_current_session(name)
            if arc_id:
                st.success("Session archived.")
            else:
                st.error("Could not archive session.")

        if c2.button("Reset inputs", key="btn_reset_session"):
            keep = {"user", "openai_api_key", "openai_model"}
            for k in list(st.session_state.keys()):
                if k not in keep and not k.startswith("_"):
                    try:
                        del st.session_state[k]
                    except Exception:
                        pass
            st.success("Inputs reset (documents remain in library).")
            try:
                st.rerun()
            except Exception:
                try:
                    st.experimental_rerun()
                except Exception:
                    pass

        df = list_session_archives()
        if df.empty:
            st.caption("No archives yet.")
            return
        pick = st.selectbox("Restore an archived session", ["(select)"] + [f"{r.archive_id} — {r.name}" for _, r in df.iterrows()], key="session_restore_pick")
        if pick != "(select)":
            arc_id = pick.split("—")[0].strip()
            c3, c4 = st.columns(2)
            if c3.button("Restore", key="btn_restore_session"):
                ok = load_session_archive(arc_id)
                if ok:
                    st.success("Session restored.")
                    try:
                        st.rerun()
                    except Exception:
                        try:
                            st.experimental_rerun()
                        except Exception:
                            pass
                else:
                    st.error("Could not restore archive.")
            if c4.button("Delete", key="btn_delete_archive"):
                delete_session_archive(arc_id)
                st.success("Archive deleted.")
                try:
                    st.rerun()
                except Exception:
                    pass

def page_home():
    page_title(APP_NAME, "End-to-end engineering change intelligence: ingest → reconcile → ECR/ECO → impact/risk → workflow → release package.")
    user = (st.session_state.get("user") or "User").strip() or "User"
    st.markdown(f"### Welcome {user} 👋")
    c1, c2, c3, c4 = st.columns(4)
    docs = list_documents()
    changes = list_changes()
    parts = list_parts()
    c1.metric("Documents", len(docs))
    c2.metric("Parts", len(parts))
    c3.metric("Changes", len(changes))
    # basic health check
    c4.metric("AI enabled", "Yes" if openai_client_from_session() is not None else "No")

    st.markdown("---")
    st.markdown("### Quick start")
    st.markdown("1) **Ingest** PDFs/drawings/BOM CSVs → 2) **Build graph** → 3) **Create ECR/ECO** → 4) **Generate package** → 5) **Review & approve** → 6) **Export release pack**")

    st.markdown("### What “ALL OF THE ABOVE” means in this app")
    st.write("- Automated ECR/ECO packages with traceability\n- Impact propagation across parts/assemblies/docs (graph)\n- Risk/FMEA + industrialization readiness\n- Visual baseline vs modified diff (images + PDFs)\n- Predictive what-if risk simulation (optional training)")
    st.info("Enterprise PLM/CAD rendering is available as plug-in points (connectors & viewer hooks).")


def page_ingest():
    page_title("Ingestion & Document Library", "Upload engineering docs; extract text/entities; build a traceable evidence base.")
    up = st.file_uploader(
    "Upload files (PDF/DOCX/TXT/CSV/JSON/XLSX/PNG/JPG/CAD)",
    type=["pdf", "docx", "txt", "csv", "json", "xlsx", "xls", "png", "jpg", "jpeg"] + CAD_FILE_EXTS,
    accept_multiple_files=True,
)
    if up:
        for f in up:
            b = f.getvalue()
            doc_id = upsert_document(f.name, b, mime=f.type or "")
            st.success(f"Ingested {f.name} → {doc_id}")

    docs = list_documents()
    st.markdown("#### Library")
    if docs.empty:
        st.info("No documents yet. Upload a PDF or CSV to get started.")
        return

    sel = st.selectbox("Select a document", docs["doc_id"].tolist(), format_func=lambda x: f"{docs.loc[docs['doc_id']==x,'filename'].values[0]}  ({x})")
    doc = get_document(sel)
    if not doc:
        return

    colA, colB = st.columns([2, 1])
    with colA:
        st.write("**Extracted entities**")
        st.json(doc["metadata"].get("entities", {}))
        st.write("**Text preview**")
        st.text_area("Extracted text (read-only)", value=(doc["extracted_text"][:6000] + ("..." if len(doc["extracted_text"]) > 6000 else "")), height=220)

    with colB:
        st.write("**Actions**")
        if st.button("Delete document", type="secondary"):
            delete_document(sel)
            st.success("Deleted.")
            st.rerun()

        if st.button("Reprocess / Refresh extraction", type="secondary"):
            reprocess_document(sel, force_ocr=False)
            st.success("Reprocessed.")
            st.rerun()

        if doc.get("file_ext") == "pdf":
            if st.button("Force OCR reprocess (PDF)", type="secondary"):
                reprocess_document(sel, force_ocr=True)
                st.success("Reprocessed with OCR (if available).")
                st.rerun()

        # Render preview
        stored_path = doc.get("stored_path")
        if stored_path and os.path.exists(stored_path):
            ext = doc["file_ext"]
            with open(stored_path, "rb") as f:
                b = f.read()
            if ext in ("png", "jpg", "jpeg") and Image is not None:
                st.image(b, caption=doc["filename"], use_container_width=True)
            elif ext == "pdf" and fitz is not None:
                png = render_pdf_page_to_png_bytes(b, pageno=0)
                if png:
                    st.image(png, caption=f"{doc['filename']} (page 1)", use_container_width=True)
            elif is_3d_filename(doc["filename"]):
                metrics, fig = analyze_3d_model_bytes(b, doc["filename"])
                st.json(metrics)
                if fig is not None:
                    st.plotly_chart(fig, use_container_width=True)
                    # Chart actions
                    bi_render_chart_actions(
                        fig,
                        title=f"Q&A chart — {title or 'Chart'}",
                        source="BOM Q&A",
                        meta={"question": q, "rows": int(len(dfq))},
                        key_prefix="qa",
                    )
                else:
                    st.info("3D preview unavailable (missing trimesh/plotly-go or STEP/IGES requires CAD kernel). You can still download the file.")
                    st.download_button("Download 3D model", data=b, file_name=doc["filename"])
            else:
                st.download_button("Download original", data=b, file_name=doc["filename"])


def page_bom_graph():
    page_title("Configuration, BOM & Engineering Graph", "Import BOM, inspect parts, and visualize dependencies for impact propagation.")
    st.markdown("#### Import BOM (CSV)")
    st.caption("Columns: parent_part_number, child_part_number, quantity, uom (optional), effectivity (optional), notes (optional)")

    bom_file = st.file_uploader("Upload BOM CSV", type=["csv"], key="bom_csv")
    if bom_file is not None:
        df = read_tabular_file(bom_file)
        df, _colmap = map_bom_schema(df)
        # Normalize edge-list columns if needed
        if 'child_part_number' not in df.columns and 'part_number' in df.columns:
            df = df.rename(columns={'part_number':'child_part_number'})
        try:
            n_parts, n_links = import_bom_csv(df)
            st.success(f"Imported BOM. Parts in DB: {n_parts}. Links imported: {n_links}.")
        except Exception as e:
            st.error(str(e))

    parts = list_parts()
    st.markdown("#### Parts")
    st.dataframe(parts, use_container_width=True, height=220)

    st.markdown("#### Graph")
    G = build_graph()
    focus = st.multiselect("Focus parts (optional)", options=parts["part_number"].tolist() if not parts.empty else [])
    focus_nodes = []
    for pn in focus:
        pid = get_part_by_number(pn)
        if pid:
            focus_nodes.append(pid)
    plot_graph(G, focus_nodes=focus_nodes if focus_nodes else None)





def apply_filters(df: pd.DataFrame, available_cols: List[str]) -> pd.DataFrame:
    """
    UI-driven filtering. User selects which columns to filter on.
    Robust to NaN-only numeric columns (avoids slider RangeError).
    """
    if df.empty:
        return df

    with st.expander("Filters", expanded=True):
        cols = [c for c in available_cols if c in df.columns]
        chosen = st.multiselect(
            "Choose filter fields",
            cols,
            default=[c for c in ["supplier", "material", "process", "revision"] if c in cols],
        )
        fdf = df

        for c in chosen:
            if c not in fdf.columns:
                continue
            series = fdf[c]

            # Numeric filter
            if pd.api.types.is_numeric_dtype(series):
                vals = pd.to_numeric(series, errors="coerce").to_numpy()
                finite = vals[np.isfinite(vals)]
                if finite.size == 0:
                    st.caption(f"{c}: no numeric values to filter")
                    continue
                mn = float(finite.min())
                mx = float(finite.max())
                if mn == mx:
                    st.caption(f"{c}: {mn} (fixed)")
                    continue
                v = st.slider(f"{c}", mn, mx, (mn, mx))
                fdf = fdf[
                    (pd.to_numeric(fdf[c], errors="coerce") >= v[0])
                    & (pd.to_numeric(fdf[c], errors="coerce") <= v[1])
                ]
            else:
                uniq = series.dropna().astype(str).unique().tolist()
                uniq = sorted(uniq)[:500]
                sel = st.multiselect(f"{c}", uniq, default=[])
                if sel:
                    fdf = fdf[fdf[c].astype(str).isin(sel)]

        return fdf



def _infer_kpi_icon(label: str) -> str:
    """Infer a helpful icon from KPI label text (keeps call-sites unchanged)."""
    s = (label or "").lower()
    rules = [
        ("part", "🧩"),
        ("serial", "🔢"),
        ("bom level", "🧬"),
        ("level", "🧬"),
        ("quantity", "📦"),
        ("qty", "📦"),
        ("unit cost", "💶"),
        ("extended cost", "💰"),
        ("lead", "⏱️"),
        ("eco", "🧾"),
        ("weight", "⚖️"),
        ("date difference", "📅"),
        ("difference", "📅"),
        ("effectivity", "📈"),
        ("timeline", "📈"),
    ]
    for k, ic in rules:
        if k in s:
            return ic
    return "📌"


def kpi_card(label: str, value: Any, help_text: str = ""):
    """
    KPI card with icon + hover glow.
    Backward compatible with the previous st.metric signature.
    """
    try:
        import html as _html
        safe_label = _html.escape(str(label))
        safe_value = _html.escape(str(value))
        safe_help = _html.escape(str(help_text)) if help_text else ""
        icon = _infer_kpi_icon(str(label))
        title_attr = f' title="{safe_help}"' if safe_help else ""
        help_html = f"<div class='kpi-help'>{safe_help}</div>" if safe_help else ""
        st.markdown(
            f"<div class='kpi-card'{title_attr}>"
            f"  <div class='kpi-top'><div class='kpi-ic'>{icon}</div><div class='kpi-label'>{safe_label}</div></div>"
            f"  <div class='kpi-value'>{safe_value}</div>"
            f"  {help_html}"
            f"</div>",
            unsafe_allow_html=True,
        )
    except Exception:
        # Fallback to standard metric if anything goes wrong
        st.metric(label, value, help=help_text if help_text else None)


# -----------------------------
# BI Dashboard Builder (export / archive / delete charts)
# -----------------------------
import hashlib as _hashlib
import json as _json
from datetime import datetime as _dt

try:
    import plotly.io as pio  # type: ignore
except Exception:
    pio = None  # type: ignore


# Optional: drag & drop ordering for BI dashboard (local Streamlit)
# Install: pip install streamlit-sortables
try:
    from streamlit_sortables import sort_items  # type: ignore
except Exception:
    sort_items = None  # type: ignore


def _bi_try_drag_drop_order(dash: List[str], store: dict) -> Optional[List[str]]:
    '''
    Returns a reordered list of chart IDs when drag & drop is available.
    Safe: never throws; returns None if not available.
    '''
    if sort_items is None:
        return None

    # Only active charts that exist in store
    active_ids = [cid for cid in dash if cid in store and store[cid].get("status") == "active"]
    if not active_ids:
        return None

    # Variant A: strings (most compatible)
    labels: List[str] = []
    for cid in active_ids:
        title = (store.get(cid) or {}).get("title") or "Chart"
        labels.append(f"{cid} | {title}")

    try:
        res = sort_items(labels, direction="vertical", key="bi_dnd_sort")  # type: ignore
        if isinstance(res, list) and res and isinstance(res[0], str):
            new_ids = [x.split(" | ", 1)[0].strip() for x in res]
            return [cid for cid in new_ids if cid in active_ids]
    except Exception:
        pass

    # Variant B: dict items (some versions support richer cards)
    try:
        items = [
            {
                "id": cid,
                "header": (store.get(cid) or {}).get("title") or "Chart",
                "content": (store.get(cid) or {}).get("source") or "",
            }
            for cid in active_ids
        ]
        res = sort_items(items, direction="vertical", key="bi_dnd_sort2")  # type: ignore
        if isinstance(res, list) and res and isinstance(res[0], dict) and "id" in res[0]:
            new_ids = [d.get("id") for d in res if d.get("id") in active_ids]
            return [cid for cid in new_ids if cid]
    except Exception:
        return None

    return None


def _bi_make_chart_id(title: str, source: str, meta: Optional[dict] = None) -> str:
    """Stable-ish id to avoid duplicates when the same chart is exported repeatedly."""
    try:
        payload = {"title": title, "source": source, "meta": meta or {}}
        s = _json.dumps(payload, sort_keys=True, default=str).encode("utf-8")
    except Exception:
        s = f"{title}|{source}".encode("utf-8")
    return _hashlib.sha1(s).hexdigest()[:10]


def _bi_fig_to_json(fig) -> str:
    try:
        return fig.to_json()
    except Exception:
        try:
            return _json.dumps(fig, default=str)
        except Exception:
            return "{}"


def _bi_fig_from_json(fig_json: str):
    if fig_json is None:
        return None
    try:
        if pio is not None:
            return pio.from_json(fig_json)
    except Exception:
        pass
    try:
        if go is not None:
            return go.Figure(_json.loads(fig_json))
    except Exception:
        return None
    return None


def bi_upsert_chart(fig, title: str, source: str, meta: Optional[dict] = None) -> str:
    """Store chart into session chart store and return its id."""
    st.session_state.setdefault("bi_chart_store", {})
    st.session_state.setdefault("bi_dashboard_order", [])
    st.session_state.setdefault("bi_archived_ids", [])

    cid = _bi_make_chart_id(title, source, meta)
    store: dict = st.session_state["bi_chart_store"]
    rec = store.get(cid, {})
    store[cid] = {
        "id": cid,
        "title": title,
        "source": source,
        "meta": meta or {},
        "created_at": rec.get("created_at") or _dt.now().strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": _dt.now().strftime("%Y-%m-%d %H:%M:%S"),
        "status": rec.get("status") or "active",  # active | archived
        "fig_json": _bi_fig_to_json(fig),
    }
    st.session_state["bi_chart_store"] = store
    st.session_state["bi_last_chart_id"] = cid
    return cid


def bi_render_chart_actions(fig, title: str, source: str, meta: Optional[dict] = None, key_prefix: str = "bi") -> None:
    """
    Render chart action buttons:
      - Archive: move chart to archive list
      - Export: add to BI dashboard page (builder)
      - Delete: remove chart from store + dashboard
    Notes:
      - We only save the chart to store when user clicks Archive/Export (so no rerun duplicates).
    """
    st.session_state.setdefault("bi_chart_store", {})
    st.session_state.setdefault("bi_dashboard_order", [])
    st.session_state.setdefault("bi_archived_ids", [])

    col1, col2, col3, col4 = st.columns([0.22, 0.26, 0.22, 0.30])
    with col1:
        exp = st.button("📤 Export to BI dashboard", key=f"{key_prefix}_export_{_bi_make_chart_id(title, source, meta)}")
    with col2:
        arc = st.button("📦 Archive chart", key=f"{key_prefix}_archive_{_bi_make_chart_id(title, source, meta)}")
    with col3:
        dele = st.button("🗑️ Delete chart", key=f"{key_prefix}_delete_{_bi_make_chart_id(title, source, meta)}")
    with col4:
        st.caption("Exported charts appear in **BI Dashboard Builder** (Navigation).")

    if exp or arc or dele:
        cid = bi_upsert_chart(fig, title=title, source=source, meta=meta)

        # Ensure status and dashboard membership
        if exp:
            if cid not in st.session_state["bi_dashboard_order"]:
                st.session_state["bi_dashboard_order"].append(cid)
            # if it was archived, restore
            store = st.session_state["bi_chart_store"]
            if cid in store:
                store[cid]["status"] = "active"
                st.session_state["bi_chart_store"] = store
            if cid in st.session_state["bi_archived_ids"]:
                st.session_state["bi_archived_ids"] = [x for x in st.session_state["bi_archived_ids"] if x != cid]
            st.success("Chart exported to BI Dashboard Builder.")
            st.experimental_rerun()

        if arc:
            store = st.session_state["bi_chart_store"]
            if cid in store:
                store[cid]["status"] = "archived"
                st.session_state["bi_chart_store"] = store
            if cid not in st.session_state["bi_archived_ids"]:
                st.session_state["bi_archived_ids"].append(cid)
            # remove from dashboard
            st.session_state["bi_dashboard_order"] = [x for x in st.session_state["bi_dashboard_order"] if x != cid]
            st.success("Chart archived.")
            st.experimental_rerun()

        if dele:
            store = st.session_state["bi_chart_store"]
            if cid in store:
                try:
                    del store[cid]
                except Exception:
                    pass
            st.session_state["bi_chart_store"] = store
            st.session_state["bi_dashboard_order"] = [x for x in st.session_state["bi_dashboard_order"] if x != cid]
            st.session_state["bi_archived_ids"] = [x for x in st.session_state["bi_archived_ids"] if x != cid]
            st.success("Chart deleted.")
            st.experimental_rerun()


def page_bi_dashboard_builder():
    page_title("BI Dashboard Builder", "Arrange exported charts into a mini BI dashboard (no-code).")

    st.session_state.setdefault("bi_chart_store", {})
    st.session_state.setdefault("bi_dashboard_order", [])
    st.session_state.setdefault("bi_archived_ids", [])

    store: dict = st.session_state["bi_chart_store"]
    dash: List[str] = st.session_state["bi_dashboard_order"]
    archived: List[str] = st.session_state["bi_archived_ids"]

    st.info("Tip: Export charts from **BOM Intelligence Dashboard → Chart Builder / Q&A**. Then arrange them here.")

    tabA, tabB = st.tabs(["📊 Dashboard", "📦 Archived"])

    with tabA:
        if not dash:
            st.warning("No charts exported yet. Go to **BOM Intelligence Dashboard** and click **Export to BI dashboard** under a chart.")
        else:
            # Order editor (Streamlit doesn't have native drag&drop without extra components)
            rows = []
            for i, cid in enumerate(dash, start=1):
                rec = store.get(cid)
                if not rec:
                    continue
                if rec.get("status") == "archived":
                    continue
                rows.append({
                    "order": i,
                    "id": cid,
                    "title": rec.get("title"),
                    "source": rec.get("source"),
                    "updated_at": rec.get("updated_at"),
                })

            df_order = pd.DataFrame(rows)
            st.markdown("#### Arrange charts")
            st.caption("Reorder your dashboard. On local Streamlit you can enable drag & drop (optional), or edit the order numbers below.")

            with st.expander("🧲 Drag & drop arrange (local)", expanded=True):
                if sort_items is None:
                    st.info("Drag & drop is optional. Install it locally with:  pip install streamlit-sortables")
                else:
                    st.write("Drag the items to reorder your dashboard charts.")
                    new_order = _bi_try_drag_drop_order(dash, store)
                    if new_order and new_order != dash:
                        st.session_state["bi_dashboard_order"] = new_order
                        dash = new_order
                        st.success("Dashboard order updated via drag & drop.")


            edited = st.data_editor(
                df_order,
                use_container_width=True,
                hide_index=True,
                num_rows="fixed",
                column_config={
                    "order": st.column_config.NumberColumn("Order", min_value=1, step=1),
                    "id": st.column_config.TextColumn("ID", disabled=True),
                    "title": st.column_config.TextColumn("Title", disabled=True),
                    "source": st.column_config.TextColumn("Source", disabled=True),
                    "updated_at": st.column_config.TextColumn("Updated", disabled=True),
                },
                key="bi_order_editor",
            )

            if edited is not None and not edited.empty:
                try:
                    edited_sorted = edited.sort_values("order")
                    new_order = edited_sorted["id"].tolist()
                    if new_order and new_order != dash:
                        st.session_state["bi_dashboard_order"] = new_order
                        dash = new_order
                        st.success("Dashboard order updated.")
                except Exception:
                    pass

            cols = st.columns([0.25, 0.25, 0.25, 0.25])
            with cols[0]:
                ncols = st.selectbox("Grid columns", [1, 2, 3], index=1, key="bi_grid_cols")
            with cols[1]:
                if st.button("⬇️ Download dashboard config (JSON)", key="bi_dl_cfg"):
                    cfg = {
                        "dashboard_order": st.session_state["bi_dashboard_order"],
                        "charts": {cid: store.get(cid) for cid in st.session_state["bi_dashboard_order"] if cid in store},
                    }
                    st.download_button(
                        "Download JSON",
                        data=_json.dumps(cfg, indent=2, default=str).encode("utf-8"),
                        file_name="bi_dashboard_config.json",
                        mime="application/json",
                        key="bi_dl_cfg_btn",
                    )
            with cols[2]:
                if st.button("🧹 Clear dashboard (keeps library)", key="bi_clear_dash"):
                    st.session_state["bi_dashboard_order"] = []
                    st.experimental_rerun()
            with cols[3]:
                st.caption("Optional: for true drag & drop, consider adding a Streamlit sortable component in production.")

            st.markdown("---")
            st.markdown("#### Dashboard preview")
            # render in grid
            active_recs = [store[cid] for cid in dash if cid in store and store[cid].get("status") == "active"]
            if not active_recs:
                st.info("No active charts in the dashboard yet.")
            else:
                for i in range(0, len(active_recs), ncols):
                    row = active_recs[i:i+ncols]
                    cols_row = st.columns(ncols)
                    for j in range(ncols):
                        if j >= len(row):
                            continue
                        rec = row[j]
                        with cols_row[j]:
                            st.markdown(f"**{rec.get('title','Chart')}**  \n<span class='small-muted'>{rec.get('source','')}</span>", unsafe_allow_html=True)
                            fig = _bi_fig_from_json(rec.get("fig_json", ""))
                            if fig is not None:
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.warning("Could not render chart (missing Plotly).")
                            b1, b2, b3 = st.columns(3)
                            if b1.button("Remove", key=f"bi_rm_{rec['id']}"):
                                st.session_state["bi_dashboard_order"] = [x for x in st.session_state["bi_dashboard_order"] if x != rec["id"]]
                                st.experimental_rerun()
                            if b2.button("Archive", key=f"bi_arc_{rec['id']}"):
                                store = st.session_state["bi_chart_store"]
                                if rec["id"] in store:
                                    store[rec["id"]]["status"] = "archived"
                                    st.session_state["bi_chart_store"] = store
                                if rec["id"] not in st.session_state["bi_archived_ids"]:
                                    st.session_state["bi_archived_ids"].append(rec["id"])
                                st.session_state["bi_dashboard_order"] = [x for x in st.session_state["bi_dashboard_order"] if x != rec["id"]]
                                st.experimental_rerun()
                            if b3.button("Delete", key=f"bi_del_{rec['id']}"):
                                store = st.session_state["bi_chart_store"]
                                if rec["id"] in store:
                                    del store[rec["id"]]
                                    st.session_state["bi_chart_store"] = store
                                st.session_state["bi_dashboard_order"] = [x for x in st.session_state["bi_dashboard_order"] if x != rec["id"]]
                                st.session_state["bi_archived_ids"] = [x for x in st.session_state["bi_archived_ids"] if x != rec["id"]]
                                st.experimental_rerun()

    with tabB:
        st.markdown("#### Archived charts")
        if not archived:
            st.info("No archived charts.")
        else:
            for cid in list(archived):
                rec = store.get(cid)
                if not rec:
                    continue
                with st.expander(f"{rec.get('title','Chart')}  •  {rec.get('source','')}", expanded=False):
                    fig = _bi_fig_from_json(rec.get("fig_json", ""))
                    if fig is not None:
                        st.plotly_chart(fig, use_container_width=True)
                    c1, c2 = st.columns(2)
                    if c1.button("Restore to library", key=f"bi_restore_{cid}"):
                        store = st.session_state["bi_chart_store"]
                        if cid in store:
                            store[cid]["status"] = "active"
                            st.session_state["bi_chart_store"] = store
                        st.session_state["bi_archived_ids"] = [x for x in st.session_state["bi_archived_ids"] if x != cid]
                        st.success("Restored.")
                        st.experimental_rerun()
                    if c2.button("Delete permanently", key=f"bi_del_arch_{cid}"):
                        store = st.session_state["bi_chart_store"]
                        if cid in store:
                            del store[cid]
                            st.session_state["bi_chart_store"] = store
                        st.session_state["bi_archived_ids"] = [x for x in st.session_state["bi_archived_ids"] if x != cid]
                        st.session_state["bi_dashboard_order"] = [x for x in st.session_state["bi_dashboard_order"] if x != cid]
                        st.success("Deleted.")
                        st.experimental_rerun()


def page_bom_dashboard():
    page_title("BOM Intelligence Dashboard", "Filterable KPIs, AI Q&A, aggregations, and real-time charting on BOM data.")
    st.caption("Supports rich BOM schemas and edge-list BOM import (parent_part_number, child_part_number, quantity...).")

    left, right = st.columns([1.05, 0.95], gap="large")

    with left:
        st.subheader("1) Load BOM data")
        bom_file = st.file_uploader("Upload BOM (CSV/XLSX)", type=["csv","xlsx","xls"], key="bom_dash_upload")
        if bom_file is not None:
            df0 = read_tabular_file(bom_file)
            df0, colmap = map_bom_schema(df0)
            st.session_state["bom_dashboard_df"] = df0
            st.session_state["bom_dashboard_colmap"] = colmap
            st.success(f"Loaded {df0.shape[0]:,} rows × {df0.shape[1]} columns.")
        else:
            df0 = st.session_state.get("bom_dashboard_df", pd.DataFrame())
            colmap = st.session_state.get("bom_dashboard_colmap", {})

        if df0.empty:
            st.info("Upload a BOM file to enable KPIs, filters, AI Q&A, and charts.")
            return

        # Ensure common numeric columns are numeric (honor schema mapping)
        for canon in ["quantity","unit_cost_eur","extended_cost_eur","lead_time_days","weight_kg","bom_level"]:
            c = colmap.get(canon, canon) if isinstance(colmap, dict) else canon
            if c in df0.columns:
                df0[c] = coerce_numeric_series(df0[c])

        # Date parsing
        for dc in ["effectivity_date","effective_from","effective_to","creation_time"]:
            if dc in df0.columns:
                df0[dc] = pd.to_datetime(df0[dc], errors="coerce")

        # Filters
        candidate_cols = [
            "serial_number","bom_description","material","process","supplier","supplier_part_no","revision",
            "effectivity_date","lifecycle_state","change_reason","compliance_tags","notes",
            "parent_rev","parent_revision","line_no","make_buy","supplier_name","manufacturer_name","manufacturer_part_no",
            "created_by","creation_time","material_spec","finish_spec","process_route","work_center","dimensions_mm",
            "tolerance_mm","criticality","effectivity_code","lifecycle_status","compliance_reach"
        ]
        available_cols = [c for c in candidate_cols if c in df0.columns]
        fdf = apply_filters(df0, available_cols)

        
        st.subheader("2) KPIs")

        # Resolve common fields
        part_col = colmap.get("part_number") or ("part_number" if "part_number" in fdf.columns else None) or (colmap.get("child_part_number") if colmap.get("child_part_number") in fdf.columns else None)
        serial_col = colmap.get("serial_number") or ("serial_number" if "serial_number" in fdf.columns else None)

        qty_col = (colmap.get("quantity") if isinstance(colmap, dict) else None) or ("quantity" if "quantity" in fdf.columns else None)
        unit_cost_col = (colmap.get("unit_cost_eur") if isinstance(colmap, dict) else None) or ("unit_cost_eur" if "unit_cost_eur" in fdf.columns else None)
        ext_cost_col = (colmap.get("extended_cost_eur") if isinstance(colmap, dict) else None) or ("extended_cost_eur" if "extended_cost_eur" in fdf.columns else None)
        lead_col = (colmap.get("lead_time_days") if isinstance(colmap, dict) else None) or ("lead_time_days" if "lead_time_days" in fdf.columns else None)
        weight_col = (colmap.get("weight_kg") if isinstance(colmap, dict) else None) or ("weight_kg" if "weight_kg" in fdf.columns else None)
        bom_level_col = "bom_level" if "bom_level" in fdf.columns else None
        eco_col = (colmap.get("eco_id") if isinstance(colmap, dict) else None) or ("eco_id" if "eco_id" in fdf.columns else None)

        qty_per_col = None
        for cand in ["total_qty_per", "qty_per", "qtyper", "total_qty_per_"]:
            if cand in fdf.columns:
                qty_per_col = cand
                break

        eff_date_col = None
        for cand in ["effectivity_date", "effectivity", "effective_from"]:
            if cand in fdf.columns:
                eff_date_col = cand
                break

        eff_from_col = "effective_from" if "effective_from" in fdf.columns else None
        eff_to_col = "effective_to" if "effective_to" in fdf.columns else None
        # Coerce mapped numeric columns to numeric for safe aggregations
        for _canon in ["quantity","unit_cost_eur","extended_cost_eur","lead_time_days","weight_kg","bom_level","bom_quantity"]:
            _c = colmap.get(_canon, _canon) if isinstance(colmap, dict) else _canon
            if _c in fdf.columns:
                fdf[_c] = coerce_numeric_series(fdf[_c])


        # KPIs (requested)
        r1 = st.columns(4)
        with r1[0]:
            kpi_card("Count Unique Part Number", int(fdf[part_col].nunique(dropna=True)) if part_col else "—")
            kpi_card("Count Unique Serial Number", int(fdf[serial_col].nunique(dropna=True)) if serial_col else 0, help_text="Serial column not found; showing 0" if not serial_col else "")
        with r1[1]:
            if bom_level_col:
                kpi_card("Bom Level (max)", int(np.nanmax(fdf[bom_level_col])) if fdf[bom_level_col].notna().any() else "—")
                kpi_card("Bom Level (avg)", f"{float(np.nanmean(fdf[bom_level_col])):,.2f}" if fdf[bom_level_col].notna().any() else "—")
            else:
                kpi_card("Bom Level", "—")
                kpi_card("Bom Level (avg)", "—")
        with r1[2]:
            kpi_card("Bom Quantity (avg)", f"{float(np.nanmean(fdf[qty_col])):,.2f}" if qty_col and fdf[qty_col].notna().any() else "—")
            kpi_card("Sum of Quantity", f"{float(np.nansum(fdf[qty_col])):,.2f}" if qty_col else "—")
        with r1[3]:
            if qty_per_col:
                kpi_card("Total Qty_per", f"{float(np.nansum(pd.to_numeric(fdf[qty_per_col], errors='coerce'))):,.2f}")
            elif qty_col:
                kpi_card("Total Qty_per", f"{float(np.nansum(fdf[qty_col])):,.2f}", help_text="Fallback: qty_per not found; using Sum of Quantity")
            else:
                kpi_card("Total Qty_per", "—")

        r2 = st.columns(4)
        with r2[0]:
            kpi_card("Sum Unit Cost (€)", f"{float(np.nansum(fdf[unit_cost_col])):,.2f}" if unit_cost_col else "0.00", help_text="Unit cost column not found; showing 0.00" if not unit_cost_col else "")
            kpi_card("Sum Extended Cost (€)", f"{float(np.nansum(fdf[ext_cost_col])):,.2f}" if ext_cost_col else (f"{float(np.nansum(pd.to_numeric(fdf[unit_cost_col], errors='coerce') * pd.to_numeric(fdf[qty_col], errors='coerce'))):,.2f}" if (unit_cost_col and qty_col) else "0.00"), help_text=("Derived from Unit Cost × Quantity" if (not ext_cost_col and unit_cost_col and qty_col) else ("Extended cost column not found; showing 0.00" if not ext_cost_col else "")))
        with r2[1]:
            kpi_card("Average Lead_time (days)", f"{float(np.nanmean(fdf[lead_col])):,.1f}" if lead_col and fdf[lead_col].notna().any() else "—")
            kpi_card("Total Lead Time (days)", f"{float(np.nansum(fdf[lead_col])):,.0f}" if lead_col else "—")
        with r2[2]:
            kpi_card("ECO_ID (unique)", int(fdf[eco_col].nunique(dropna=True)) if eco_col else 0, help_text="ECO_ID column not found; showing 0" if not eco_col else "")
            kpi_card("Weight KG (sum)", f"{float(np.nansum(fdf[weight_col])):,.2f}" if weight_col else "0.00", help_text="Weight column not found; showing 0.00" if not weight_col else "")
        with r2[3]:
            if eff_from_col and eff_to_col:
                dif = (pd.to_datetime(fdf[eff_from_col], errors="coerce") - pd.to_datetime(fdf[eff_to_col], errors="coerce")).dt.days
                kpi_card("Avg Date Difference (days)", f"{float(np.nanmean(dif)):.1f}" if dif.notna().any() else "—",
                         help_text="Computed as Effective_From - Effective_to")
            else:
                kpi_card("Avg Date Difference (days)", "0.0", help_text="Effective_From/Effective_To not found; showing 0.0")
        st.subheader("4) Data preview")
        st.dataframe(fdf.head(2000), use_container_width=True, height=340)
        st.download_button("Download filtered data (CSV)", data=fdf.to_csv(index=False).encode("utf-8"), file_name="bom_filtered.csv")


    with right:
        st.subheader("AI-powered BOM Q&A + Real-time Charts")

        # Quick chart builder
        with st.expander("Chart builder", expanded=True):
            cols_numeric = [c for c in fdf.columns if pd.api.types.is_numeric_dtype(fdf[c])]
            cols_any = list(fdf.columns)

            x = st.selectbox("X-axis", options=[c for c in cols_any if c in fdf.columns], index=cols_any.index("effectivity_date") if "effectivity_date" in cols_any else 0)
            y = st.selectbox("Y-axis (metric)", options=cols_numeric, index=cols_numeric.index("extended_cost_eur") if "extended_cost_eur" in cols_numeric else 0)
            agg = st.selectbox("Aggregation", options=["sum","mean","count"], index=0)
            chart_type = st.selectbox("Chart type", options=["line","bar","area"], index=0)

            try:
                g = fdf.copy()
                if pd.api.types.is_datetime64_any_dtype(g[x]):
                    g["_x"] = g[x].dt.date
                else:
                    g["_x"] = g[x].astype(str)

                if agg == "count":
                    grp = g.groupby("_x").size().reset_index(name="value")
                elif agg == "mean":
                    grp = g.groupby("_x")[y].mean().reset_index(name="value")
                else:
                    grp = g.groupby("_x")[y].sum().reset_index(name="value")

                if chart_type == "bar":
                    fig = px.bar(grp, x="_x", y="value", title=f"{agg}({y}) by {x}")
                elif chart_type == "area":
                    fig = px.area(grp, x="_x", y="value", title=f"{agg}({y}) by {x}")
                else:
                    fig = px.line(grp, x="_x", y="value", title=f"{agg}({y}) by {x}")
                st.plotly_chart(fig, use_container_width=True)
                # Chart actions
                bi_render_chart_actions(
                    fig,
                    title=f"Chart builder — {chart_type} {agg}({y}) by {x}",
                    source="BOM Chart Builder",
                    meta={"x": x, "y": y, "agg": agg, "chart_type": chart_type, "rows": int(len(fdf))},
                    key_prefix="cb",
                )
            except Exception as e:
                st.warning(f"Chart builder error: {e}")

        # AI / heuristic Q&A
        st.markdown("**Ask a question about the BOM (aggregations + filters + charts)**")
        q = st.text_area("Example: 'Sum extended cost by supplier' or 'Top 10 parts by quantity' or 'Plot sum extended cost over effectivity date'", height=90, key="bom_q")
        colA, colB = st.columns([0.35, 0.65])
        run = colA.button("Run", key="bom_run")
        colB.caption("If an LLM key is configured, the assistant can interpret more complex questions. Otherwise a fast heuristic engine is used.")

        def _heuristic_bom_answer(df: pd.DataFrame, q: str):
            qq = (q or "").strip().lower()
            if not qq:
                return ("", None, None)

            def _col(candidates):
                for cand in candidates:
                    lcand = str(cand).lower()
                    for c in df.columns:
                        cc = str(c).lower()
                        if cc == lcand or lcand in cc:
                            return c
                return None


            # Simple metric intents
            if "unique" in qq and "part" in qq and "count" in qq:
                return ("Unique part numbers", df[colmap.get("part_number","part_number")].nunique() if (colmap.get("part_number","part_number") in df.columns) else None, None)
            if "sum" in qq and "extended" in qq:
                if "extended_cost_eur" in df.columns:
                    return ("Sum extended cost (€)", float(np.nansum(df["extended_cost_eur"])), None)

            # Lead time / weight / dates / IDs (tabular BOM)
            if (("average" in qq) or ("avg" in qq)) and ("lead" in qq):
                col = _col(["lead_time", "leadtime", "lead_time_days", "lead_time_day", "lead_time_weeks", "lead_time_week"])
                if col is not None:
                    v = pd.to_numeric(df[col], errors="coerce")
                    return ("Average lead time", float(np.nanmean(v)), None)
            if (("total" in qq) or ("sum" in qq)) and ("lead" in qq):
                col = _col(["lead_time", "leadtime", "lead_time_days", "lead_time_day", "lead_time_weeks", "lead_time_week"])
                if col is not None:
                    v = pd.to_numeric(df[col], errors="coerce")
                    return ("Total lead time", float(np.nansum(v)), None)
            if ("unique" in qq) and ("serial" in qq):
                col = _col(["serial_number", "serial_no", "serial", "sn"])
                if col is not None:
                    return ("Unique serial numbers", int(pd.Series(df[col]).nunique(dropna=True)), None)
            if ("bom level" in qq) or ("bom_level" in qq):
                col = _col(["bom_level", "level"])
                if col is not None:
                    v = pd.to_numeric(df[col], errors="coerce")
                    return ("Max BOM level", float(np.nanmax(v)), None)
            if ("eco" in qq) and (("id" in qq) or ("eco_id" in qq) or qq.strip()=="eco"):
                col = _col(["eco_id", "eco", "change_id", "ecr_id"])
                if col is not None:
                    return ("Distinct ECO/ECR IDs", int(pd.Series(df[col]).nunique(dropna=True)), None)
            if ("weight" in qq):
                col = _col(["weight_kg", "weight", "mass_kg", "mass"])
                if col is not None:
                    v = pd.to_numeric(df[col], errors="coerce")
                    return ("Sum weight (kg)", float(np.nansum(v)), None)
            if ("average" in qq or "avg" in qq) and ("date difference" in qq or "datediff" in qq or "effective_from" in qq):
                c_from = _col(["effective_from", "effectivity_from", "effectivity_start", "valid_from"])
                c_to = _col(["effective_to", "effectivity_to", "effectivity_end", "valid_to"])
                if c_from is not None and c_to is not None:
                    d1 = pd.to_datetime(df[c_from], errors="coerce")
                    d2 = pd.to_datetime(df[c_to], errors="coerce")
                    delta = (d1 - d2).dt.days.astype("float")
                    return ("Average date difference (days)", float(np.nanmean(delta)), None)


            # group by
            by = None
            m = re.search(r"by\s+([a-z0-9_ ]+)$", qq)
            if m:
                cand = m.group(1).strip().replace(" ", "_")
                # find closest column
                for c in df.columns:
                    if cand in c:
                        by = c
                        break

            topn = 10
            m2 = re.search(r"top\s+(\d+)", qq)
            if m2:
                try:
                    topn = int(m2.group(1))
                except Exception:
                    pass

            if by and "sum" in qq and "quantity" in qq and "quantity" in df.columns:
                res = df.groupby(by)["quantity"].sum().sort_values(ascending=False).head(topn).reset_index()
                return (f"Sum(quantity) by {by} (top {topn})", res, None)

            if by and "sum" in qq and "extended" in qq and "extended_cost_eur" in df.columns:
                res = df.groupby(by)["extended_cost_eur"].sum().sort_values(ascending=False).head(topn).reset_index()
                return (f"Sum(extended_cost_eur) by {by} (top {topn})", res, None)

            if ("plot" in qq) or ("chart" in qq) or ("line" in qq) or ("trend" in qq):
                # Robust plot intent: supports queries like:
                # "plot sum extended cost over effectivity date"
                if px is None:
                    return ("Charts unavailable (Plotly not installed). Install: pip install plotly", None, None)

                def _find_col(candidates: List[str]) -> Optional[str]:
                    for c in candidates:
                        if c in df.columns:
                            return c
                    return None

                # ---- Y (metric) detection ----
                y_col: Optional[str] = None
                if ("extended" in qq and "cost" in qq) or ("ext" in qq and "cost" in qq):
                    y_col = _find_col([
                        "extended_cost_eur", "extended_cost", "extended_cost_euro", "extended_cost_eur_total",
                        "extended_cost_euro_total",
                    ])
                if y_col is None and ("unit" in qq and "cost" in qq):
                    y_col = _find_col(["unit_cost_eur", "unit_cost", "unit_cost_euro"])
                if y_col is None and (("qty_per" in qq) or ("qty per" in qq) or ("qtyper" in qq)):
                    y_col = _find_col(["qty_per", "qtyper", "quantity_per"])
                if y_col is None and (("quantity" in qq) or re.search(r"\bqty\b", qq)):
                    y_col = _find_col(["quantity", "qty"])

                # Last-resort: infer a likely numeric column by name
                if y_col is None:
                    for c in df.columns:
                        lc = str(c).lower()
                        if any(k in lc for k in ["extended", "ext", "total cost", "cost", "eur", "euro", "amount"]):
                            y_col = c
                            break

                # ---- X (date) detection ----
                x_col: Optional[str] = None
                if "effectiv" in qq:
                    x_col = _find_col(["effectivity_date", "effectivity", "effective_from", "effective_to"])
                if x_col is None:
                    for c in df.columns:
                        lc = str(c).lower()
                        if any(k in lc for k in ["effectivity", "effectiv", "effective_from", "effective from", "eff_from", "eff date", "effective date"]):
                            x_col = c
                            break
                if x_col is None:
                    for c in df.columns:
                        lc = str(c).lower()
                        if "date" in lc or "time" in lc:
                            x_col = c
                            break

                if x_col and y_col:
                    g = df.copy()
                    # Coerce date
                    g["_x"] = pd.to_datetime(g[x_col], errors="coerce").dt.date
                    g = g.dropna(subset=["_x"])
                    if g.empty:
                        return (f"No parseable dates found in '{x_col}' to plot.", None, None)

                    # Aggregation intent
                    agg_mode = "sum"
                    if ("mean" in qq) or ("avg" in qq) or ("average" in qq):
                        agg_mode = "mean"
                    elif "median" in qq:
                        agg_mode = "median"
                    elif "count" in qq:
                        agg_mode = "count"

                    if agg_mode == "count":
                        grp = g.groupby("_x").size().reset_index(name="value")
                        title = f"Count over {x_col}"
                    elif agg_mode == "mean":
                        grp = g.groupby("_x")[y_col].mean().reset_index(name="value")
                        title = f"Average {y_col} over {x_col}"
                    elif agg_mode == "median":
                        grp = g.groupby("_x")[y_col].median().reset_index(name="value")
                        title = f"Median {y_col} over {x_col}"
                    else:
                        grp = g.groupby("_x")[y_col].sum().reset_index(name="value")
                        title = f"Sum {y_col} over {x_col}"

                    fig = px.line(grp, x="_x", y="value", title=title, markers=True)
                    return ("Chart generated", grp, fig)

                missing = []
                if not y_col:
                    missing.append("a numeric metric column (e.g., Extended Cost)")
                if not x_col:
                    missing.append("a date column (e.g., Effectivity / Effective_From)")
                return ("I can plot that, but I couldn't find " + " and ".join(missing) +
                        ". Map columns in the sidebar, or ask e.g. 'plot sum <metric> over <date>'.", None, None)

            return ("I couldn't parse that question. Try: 'Sum extended cost by supplier' or 'Top 10 parts by quantity' or 'Plot sum extended cost over effectivity date'.", None, None)

        if run:
            dfq = fdf.copy()
            title, table_or_value, fig = None, None, None

            # Use existing LLM router if available; otherwise heuristic
            try:
                # Prefer fast heuristic unless user has LLM and wants it
                use_llm = bool(st.session_state.get("llm_provider")) and st.session_state.get("llm_provider") in ("OpenAI","Mistral","Auto")
            except Exception:
                use_llm = False

            if use_llm:
                # Ask LLM to produce a JSON plan (safe, structured)
                schema = {
                    "task": "bom_analysis",
                    "fields": ["metric","groupby","filter","top_n","chart"],
                    "metric": "one of: sum_extended_cost_eur, sum_quantity, avg_lead_time_days, count_unique_part_number, count_unique_serial_number",
                    "groupby": "optional column name",
                    "top_n": "optional int",
                    "chart": "optional: line|bar|none"
                }
                prompt = f"""You are a data analyst. Convert the user question into a strict JSON plan.
JSON schema example: {json.dumps(schema)}

Available columns: {list(dfq.columns)}

User question: {q}
Return ONLY JSON."""
                try:
                    raw = llm_text(prompt, max_tokens=220)
                    plan = json.loads(raw.strip())
                except Exception:
                    plan = None

                if plan:
                    try:
                        metric = plan.get("metric","")
                        groupby = plan.get("groupby")
                        top_n = int(plan.get("top_n",10)) if str(plan.get("top_n","")).strip() else 10
                        chart = (plan.get("chart") or "none").lower()

                        if metric == "sum_extended_cost_eur" and "extended_cost_eur" in dfq.columns:
                            if groupby and groupby in dfq.columns:
                                res = dfq.groupby(groupby)["extended_cost_eur"].sum().sort_values(ascending=False).head(top_n).reset_index()
                                title = f"Sum extended_cost_eur by {groupby}"
                                table_or_value = res
                                if chart in ("bar","line"):
                                    fig = px.bar(res, x=groupby, y="extended_cost_eur", title=title) if chart=="bar" else px.line(res, x=groupby, y="extended_cost_eur", title=title)
                            else:
                                title = "Sum extended_cost_eur (€)"
                                table_or_value = float(np.nansum(dfq["extended_cost_eur"]))
                        elif metric == "sum_quantity" and "quantity" in dfq.columns:
                            if groupby and groupby in dfq.columns:
                                res = dfq.groupby(groupby)["quantity"].sum().sort_values(ascending=False).head(top_n).reset_index()
                                title = f"Sum quantity by {groupby}"
                                table_or_value = res
                                if chart in ("bar","line"):
                                    fig = px.bar(res, x=groupby, y="quantity", title=title) if chart=="bar" else px.line(res, x=groupby, y="quantity", title=title)
                            else:
                                title = "Sum quantity"
                                table_or_value = float(np.nansum(dfq["quantity"]))
                        else:
                            title, table_or_value, fig = _heuristic_bom_answer(dfq, q)
                    except Exception:
                        title, table_or_value, fig = _heuristic_bom_answer(dfq, q)
                else:
                    title, table_or_value, fig = _heuristic_bom_answer(dfq, q)
            else:
                title, table_or_value, fig = _heuristic_bom_answer(dfq, q)

            if title:
                st.success(title)
            if table_or_value is None:
                pass
            elif isinstance(table_or_value, (int,float,np.integer,np.floating)):
                st.metric("Result", f"{float(table_or_value):,.2f}")
            elif isinstance(table_or_value, pd.DataFrame):
                st.dataframe(table_or_value, use_container_width=True, height=260)
            else:
                st.write(table_or_value)

            if fig is not None:
                st.plotly_chart(fig, use_container_width=True)





# ----------------------------
# Change Intelligence Layer (Milestones 1-3)
# ----------------------------

def page_uedm_hub():
    page_title("Unified Data Hub (UEDM)", "Normalize CAD/BIM/BOM/Docs into a single engineering data model for KPI dashboards, comparison and AI querying.")

    st.caption("This tab builds a Unified Engineering Data Model (UEDM) from whatever you have already ingested in the app (documents + BOM). It does NOT delete or modify your existing flows.")

    with st.expander("Add EBOM/MBOM/BOM data sources (optional)", expanded=False):
        st.caption("Upload additional BOM sources (EBOM, MBOM, service BOM) to be included in the UEDM and KPI dashboards.")
        bom_type = st.selectbox("Source type", ["EBOM","MBOM","Service BOM","Other"], index=0, key="uedm_extra_bom_type")
        up = st.file_uploader("Upload BOM (CSV or Excel)", type=["csv","xlsx"], key="uedm_extra_bom_upload")
        if "uedm_extra_boms" not in st.session_state:
            st.session_state["uedm_extra_boms"] = []
        if up is not None:
            try:
                if up.name.lower().endswith('.csv'):
                    df = pd.read_csv(up)
                else:
                    df = pd.read_excel(up)
                st.dataframe(df.head(50), use_container_width=True, height=220)
                if st.button("Add this source", key="uedm_add_extra_bom"):
                    st.session_state["uedm_extra_boms"].append({"type": bom_type, "name": up.name, "df": df})
                    st.success("Added source to UEDM inputs. Rebuild UEDM to include it.")
            except Exception as e:
                st.error(f"Could not read file: {e}")

        if st.session_state.get("uedm_extra_boms"):
            st.markdown("**Loaded extra sources**")
            for i, s in enumerate(st.session_state["uedm_extra_boms"]):
                c1,c2,c3 = st.columns([3,2,1])
                c1.write(f"{i+1}. {s.get('name')} ({s.get('type')})")
                c2.write(f"Rows: {len(s.get('df',[]))}")
                if c3.button("Remove", key=f"uedm_rm_src_{i}"):
                    st.session_state["uedm_extra_boms"].pop(i)
                    st.rerun()


    colA, colB, colC = st.columns([1,1,1])
    with colA:
        if st.button("Build / Refresh UEDM from current sources", use_container_width=True, key="uedm_refresh"):
            uedm = build_uedm_from_current_sources()
            st.session_state["uedm_last_built_at"] = now_iso()
            st.success("UEDM built from current sources.")
    with colB:
        snap_name = st.text_input("Snapshot name", value=f"Snapshot {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}", key="uedm_snap_name")
        if st.button("Save snapshot", use_container_width=True, key="uedm_save_snap"):
            uedm = st.session_state.get("uedm") or build_uedm_from_current_sources()
            sid = save_uedm_snapshot(snap_name, uedm)
            st.success(f"Saved snapshot: {sid}")
    with colC:
        st.metric("Last build", st.session_state.get("uedm_last_built_at", "—"))

    uedm = st.session_state.get("uedm")
    if not isinstance(uedm, dict) or (uedm.get("items") is None):
        st.info("No UEDM yet. Click 'Build / Refresh UEDM'.")
        return

    summary = uedm_kpi_summary(uedm)
    c1,c2,c3,c4,c5 = st.columns(5)
    c1.metric("Total Items", summary.get("items_total",0))
    c2.metric("Parts", summary.get("parts_total",0))
    c3.metric("Documents", summary.get("documents_total",0))
    c4.metric("Relationships", summary.get("relationships_total",0))
    c5.metric("Attributes", summary.get("attributes_total",0))

    with st.expander("Sources & Coverage", expanded=False):
        st.json(uedm.get("sources", {}))

    tabs = st.tabs(["Items", "Relationships", "Attributes", "Events"])
    tabs[0].dataframe(uedm.get("items", pd.DataFrame()), use_container_width=True, height=320)
    tabs[1].dataframe(uedm.get("relationships", pd.DataFrame()), use_container_width=True, height=320)
    tabs[2].dataframe(uedm.get("attributes", pd.DataFrame()), use_container_width=True, height=320)
    tabs[3].dataframe(uedm.get("events", pd.DataFrame()), use_container_width=True, height=320)


def _dynamic_numeric_fields(df: pd.DataFrame) -> List[str]:
    if df is None or df.empty:
        return []
    out = []
    for c in df.columns:
        if c.lower().endswith("_date") or c.lower().endswith("_time"):
            continue
        if pd.api.types.is_numeric_dtype(df[c]):
            out.append(c)
    return out


def _as_datetime_series(s: pd.Series) -> pd.Series:
    try:
        return pd.to_datetime(s, errors="coerce")
    except Exception:
        return pd.Series([pd.NaT]*len(s))


def page_change_intelligence_dashboard(identity: Dict[str,str]):
    page_title("Change Intelligence Dashboard", "Dynamic KPIs across BOM + CAD/BIM + documents, plus queryable, explainable insights.")

    uedm = st.session_state.get("uedm")
    if not isinstance(uedm, dict) or uedm.get("items") is None or len(uedm.get("items", [])) == 0:
        st.info("Build the UEDM first (Unified Data Hub tab).")
        return

    items = uedm.get("items", pd.DataFrame()).copy()
    attrs = uedm.get("attributes", pd.DataFrame()).copy()

    # Pivot attributes to wide form for parts only (best-effort)
    part_ids = items[items.get("item_type","") == "part"]["item_id"].astype(str).tolist() if ("item_type" in items.columns and "item_id" in items.columns) else []
    a_part = attrs[attrs["item_id"].astype(str).isin(part_ids)].copy() if (not attrs.empty and "item_id" in attrs.columns) else pd.DataFrame()

    wide = pd.DataFrame()
    if not a_part.empty and {"item_id","key","value"}.issubset(a_part.columns):
        try:
            wide = a_part.pivot_table(index="item_id", columns="key", values="value", aggfunc="first").reset_index()
        except Exception:
            wide = pd.DataFrame()

    st.subheader("Filter")
    # Filter on a few common columns when available
    filter_cols = [c for c in ["supplier","supplier_name","material","process","revision","lifecycle_status","compliance_tags","effectivity_date"] if (wide is not None and c in wide.columns)]
    fdf = wide.copy() if isinstance(wide, pd.DataFrame) else pd.DataFrame()
    if not fdf.empty and filter_cols:
        f = st.columns(min(4, len(filter_cols)))
        for i, c in enumerate(filter_cols[:4]):
            vals = sorted([v for v in fdf[c].dropna().astype(str).unique().tolist()])
            sel = f[i].multiselect(c.replace("_"," ").title(), vals, default=[], key=f"ci_filter_{c}")
            if sel:
                fdf = fdf[fdf[c].astype(str).isin(sel)]

    # KPI cards (parts) - dynamic but with engineering focus
    st.subheader("Core KPIs (current selection)")
    kcols = st.columns(6)
    kcols[0].metric("Parts (unique)", int(len(fdf)) if not fdf.empty else 0)

    def _sum(col):
        if fdf.empty or col not in fdf.columns:
            return None
        try:
            return float(pd.to_numeric(fdf[col], errors="coerce").fillna(0).sum())
        except Exception:
            return None

    def _avg(col):
        if fdf.empty or col not in fdf.columns:
            return None
        try:
            s = pd.to_numeric(fdf[col], errors="coerce")
            return float(s.dropna().mean()) if s.dropna().shape[0] else None
        except Exception:
            return None

    total_cost = _sum("extended_cost_eur") or _sum("extended_cost")
    unit_cost_sum = _sum("unit_cost_eur") or _sum("unit_cost")
    qty_sum = _sum("quantity") or _sum("bom_quantity")
    weight_sum = _sum("weight_kg")
    lead_avg = _avg("lead_time_days")

    kcols[1].metric("Sum Qty", f"{qty_sum:,.0f}" if qty_sum is not None else "—")
    kcols[2].metric("Sum Unit Cost (€)", f"{unit_cost_sum:,.2f}" if unit_cost_sum is not None else "—")
    kcols[3].metric("Sum Extended Cost (€)", f"{total_cost:,.2f}" if total_cost is not None else "—")
    kcols[4].metric("Avg Lead Time (days)", f"{lead_avg:,.1f}" if lead_avg is not None else "—")
    kcols[5].metric("Sum Weight (kg)", f"{weight_sum:,.2f}" if weight_sum is not None else "—")

    st.markdown("---")

    # Timeline chart if effectivity exists
    eff_col = None
    for c in ["effectivity_date","effectivity","effective_from","created_at","creation_time"]:
        if not fdf.empty and c in fdf.columns:
            eff_col = c
            break

    if eff_col and total_cost is not None:
        st.subheader("Effectivity timeline")
        try:
            tmp = fdf.copy()
            tmp[eff_col] = _as_datetime_series(tmp[eff_col])
            tmp = tmp[tmp[eff_col].notna()].copy()
            if not tmp.empty:
                # build cost over time as sum per date
                val_col = "extended_cost_eur" if "extended_cost_eur" in tmp.columns else ("extended_cost" if "extended_cost" in tmp.columns else None)
                if val_col:
                    tmp[val_col] = pd.to_numeric(tmp[val_col], errors="coerce")
                    grp = tmp.groupby(tmp[eff_col].dt.date)[val_col].sum().reset_index()
                    grp.columns = ["date","sum_extended_cost"]
                    fig = px.line(grp, x="date", y="sum_extended_cost")
                    st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.info(f"Timeline unavailable: {e}")

    # AI Query over UEDM (fast heuristics first; LLM optional)
    st.subheader("Ask questions across BOM + CAD/BIM + Docs")
    q = st.text_area("Examples: 'Top 10 parts by extended cost', 'Average lead time by supplier', 'What CAD/BIM models were ingested and what changed?', 'Summarize the latest procedure and its key parameters'", height=90, key="ci_q")
    col1, col2 = st.columns([1,1])
    run = col1.button("Run (fast)", key="ci_run_fast")
    run_llm = col2.button("Run with AI (explainable)", key="ci_run_ai")

    # Assemble a query dataframe from available fdf (parts) and doc attributes
    doc_attrs = attrs[attrs["item_id"].astype(str).str.startswith("doc:")].copy() if (not attrs.empty and "item_id" in attrs.columns) else pd.DataFrame()

    def fast_query(df: pd.DataFrame, q: str):
        qq = (q or "").lower()
        if df is None or df.empty:
            return {"type":"text","text":"No structured data loaded. Import BOM/EBOM/MBOM and build UEDM."}
        # simple groupby/sum/avg/top parsing
        if "top" in qq and ("cost" in qq or "extended" in qq):
            val = "extended_cost_eur" if "extended_cost_eur" in df.columns else ("extended_cost" if "extended_cost" in df.columns else None)
            if not val:
                return {"type":"text","text":"No extended cost column found."}
            tmp=df.copy(); tmp[val]=pd.to_numeric(tmp[val],errors='coerce')
            key = "item_id"
            out=tmp.groupby(key)[val].sum().sort_values(ascending=False).head(10).reset_index()
            out.columns=["part","sum_extended_cost"]
            return {"type":"table","df":out}
        if "average" in qq and "lead" in qq and "supplier" in qq:
            col_sup = "supplier" if "supplier" in df.columns else ("supplier_name" if "supplier_name" in df.columns else None)
            col_lt = "lead_time_days" if "lead_time_days" in df.columns else None
            if col_sup and col_lt:
                tmp=df.copy(); tmp[col_lt]=pd.to_numeric(tmp[col_lt],errors='coerce')
                out=tmp.groupby(col_sup)[col_lt].mean().sort_values(ascending=False).reset_index().head(20)
                out.columns=["supplier","avg_lead_time_days"]
                return {"type":"table","df":out}
        if "sum" in qq and "quantity" in qq:
            col_q = "quantity" if "quantity" in df.columns else ("bom_quantity" if "bom_quantity" in df.columns else None)
            if col_q:
                s=float(pd.to_numeric(df[col_q],errors='coerce').fillna(0).sum())
                return {"type":"text","text":f"Sum of quantity = {s:,.0f}"}
        return {"type":"text","text":"Try a more specific query (top, sum, avg, by supplier/material/revision/effectivity)."}

    if run:
        ans = fast_query(fdf if not fdf.empty else wide, q)
        if ans["type"] == "table":
            st.dataframe(ans["df"], use_container_width=True)
        else:
            st.write(ans["text"])

    if run_llm:
        # Build a compact context: KPI summary + sample rows + CAD/BIM doc metrics
        ctx = {
            "user": identity.get("user"),
            "role": identity.get("role"),
            "kpis": {
                "parts": int(len(fdf)) if not fdf.empty else 0,
                "sum_qty": qty_sum,
                "sum_ext_cost": total_cost,
                "avg_lead_time": lead_avg,
                "sum_weight": weight_sum,
            },
            "sample_parts": (fdf.head(25).to_dict("records") if isinstance(fdf, pd.DataFrame) and not fdf.empty else []),
            "cad_bim_models": doc_attrs[doc_attrs["key"].astype(str).str.startswith("model_")].head(120).to_dict("records") if not doc_attrs.empty else [],
        }
        prompt = (
            "You are an engineering change intelligence assistant. Answer the user's question using the provided structured context. "
            "Be traceable: cite which fields/rows you used and highlight assumptions. Provide an actionable output (decision + evidence + next actions).\n\n"
            f"Question: {q}\n\nContext(JSON):\n{json.dumps(ctx, ensure_ascii=False)[:15000]}"
        )
        try:
            out = llm_generate_text(prompt, temperature=0.2, max_tokens=900)
            st.markdown(out)
        except Exception as e:
            st.error(f"AI planning failed: {e}")


def page_snapshot_compare(identity: Dict[str,str]):
    page_title("Snapshot Compare (Prev vs Current)", "Detect KPI and data changes between engineering snapshots: BOM deltas, CAD/BIM metric deltas, coverage deltas.")

    st.caption("Create snapshots in the Unified Data Hub tab. Compare any two snapshots here to detect KPI drift and structural changes.")

    snaps = list_uedm_snapshots()
    if snaps.empty:
        st.info("No snapshots yet. Build UEDM and save a snapshot first.")
        return

    sids = snaps["snapshot_id"].tolist()
    labels = [f"{r['snapshot_id']} — {r['name']} ({r['created_at']})" for _, r in snaps.iterrows()]
    map_lbl = dict(zip(labels, sids))

    col1, col2 = st.columns(2)
    a_lbl = col1.selectbox("Snapshot A (baseline)", labels, index=0, key="snap_a")
    b_lbl = col2.selectbox("Snapshot B (current)", labels, index=0 if len(labels)==1 else 1, key="snap_b")

    if st.button("Compare", key="snap_compare"):
        a = load_uedm_snapshot(map_lbl[a_lbl])
        b = load_uedm_snapshot(map_lbl[b_lbl])
        if not a or not b:
            st.error("Could not load snapshots.")
            return
        res = compare_uedm_snapshots(a, b)
        st.session_state["snap_compare_res"] = res

    res = st.session_state.get("snap_compare_res")
    if not res:
        return

    k1,k2 = st.columns(2)
    kpi_a = res.get("kpi_a",{})
    kpi_b = res.get("kpi_b",{})
    with k1:
        st.subheader("KPIs — A")
        st.json(kpi_a)
    with k2:
        st.subheader("KPIs — B")
        st.json(kpi_b)

    st.subheader("Structural deltas")
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Items added", len(res.get("items_added",[])))
    c2.metric("Items removed", len(res.get("items_removed",[])))
    c3.metric("Rels added", len(res.get("rels_added",[])))
    c4.metric("Rels removed", len(res.get("rels_removed",[])))

    with st.expander("Items added", expanded=False):
        st.write(res.get("items_added",[])[:200])
    with st.expander("Items removed", expanded=False):
        st.write(res.get("items_removed",[])[:200])

    st.subheader("Top numeric KPI / attribute changes")
    deltas = pd.DataFrame(res.get("numeric_attribute_deltas", []))
    if not deltas.empty:
        st.dataframe(deltas, use_container_width=True, height=360)
    else:
        st.info("No numeric attribute deltas detected (or no overlapping numeric attributes).")

    # Export a Change Intelligence PDF report
    if st.button("Export Change Intelligence Report (PDF)", key="snap_export_pdf"):
        try:
            pdf_path = write_pdf_snapshot_report(res, identity)
            with open(pdf_path, "rb") as f:
                st.download_button("Download report", data=f.read(), file_name=os.path.basename(pdf_path), mime="application/pdf")
        except Exception as e:
            st.error(f"Could not export report: {e}")


def write_pdf_snapshot_report(res: Dict[str, Any], identity: Dict[str,str]) -> str:
    if canvas is None or A4 is None:
        raise RuntimeError("reportlab not installed")
    out_dir = os.path.join(EXPORTS_DIR, "snapshot_reports")
    os.makedirs(out_dir, exist_ok=True)
    stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(out_dir, f"change_intel_report_{stamp}.pdf")

    c = canvas.Canvas(path, pagesize=A4)
    w,h = A4
    x=40
    y=h-50

    def line(t, dy=12, bold=False):
        nonlocal y
        if y<60:
            c.showPage(); y=h-50
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 11 if bold else 9)
        c.drawString(x, y, str(t)[:140])
        y-=dy

    line("CynthAI© Engineering Change Copilot — Change Intelligence Report", bold=True, dy=18)
    line(f"Generated: {now_iso()}   User: {identity.get('user','')}   Role: {identity.get('role','')}")
    line("", dy=10)
    line("KPI Summary (Baseline vs Current)", bold=True, dy=16)
    line(f"A: {json.dumps(res.get('kpi_a',{}), ensure_ascii=False)[:160]}")
    line(f"B: {json.dumps(res.get('kpi_b',{}), ensure_ascii=False)[:160]}")
    line("", dy=10)
    line("Structural deltas", bold=True, dy=16)
    line(f"Items added: {len(res.get('items_added',[]))}   Items removed: {len(res.get('items_removed',[]))}")
    line(f"Relationships added: {len(res.get('rels_added',[]))}   removed: {len(res.get('rels_removed',[]))}")
    line("", dy=10)
    line("Top numeric deltas", bold=True, dy=16)
    for r in (res.get('numeric_attribute_deltas', []) or [])[:18]:
        line(f"{r.get('item_id','')} | {r.get('key','')} : {r.get('from','')} -> {r.get('to','')} (Δ {r.get('delta','')})", dy=11)

    c.save()
    return path



def page_geometry_bim_intelligence():
    page_title("Geometry & BIM Intelligence", "Analyze 3D CAD (STEP/STL/OBJ/GLB) and BIM (IFC) models, extract KPIs, and detect changes between revisions.")

    st.caption("This uses lightweight geometry proxies (bbox/volume/surface/vertex count) for fast change detection. For STEP/IGES exact analysis, configure FreeCADCmd below or use the CAD Drawing Studio tab.")

    with st.expander("STEP/IGES exact KPIs (FreeCADCmd fallback)", expanded=False):
        st.write("If trimesh cannot load STEP/IGES on your machine, the app can call FreeCADCmd to extract exact bounding box, volume, and surface area.")
        default_path = st.session_state.get("freecad_cmd_path", "C:/Users/d.adewunmi/AppData/Local/Programs/FreeCAD 1.0/bin/FreeCADCmd.exe")
        st.text_input("FreeCADCmd path", value=default_path, key="freecad_cmd_path", help="Example: C:/Users/.../FreeCAD 1.0/bin/FreeCADCmd.exe")
        colx, coly = st.columns(2)
        if colx.button("Auto-detect FreeCADCmd", key="freecad_autodetect"):
            cand = _resolve_freecadcmd()
            if cand:
                st.session_state["freecad_cmd_path"] = cand
                st.success(f"Detected: {cand}")
            else:
                st.warning("Could not detect FreeCADCmd automatically. Please set the path manually.")
        if coly.button("Test FreeCADCmd", key="freecad_test"):
            cand = _resolve_freecadcmd()
            if not cand:
                st.error("FreeCADCmd not resolved. Please set a valid path.")
            else:
                st.success(f"FreeCADCmd resolved as: {cand}")

    # pick two model files from the document library (CAD/BIM)
    docs = list_documents()
    if docs.empty:
        st.info("No documents in library. Ingest CAD/BIM files first (Connectors tab).")
        return

    cad_docs = docs[docs["file_ext"].astype(str).str.lower().isin(list(THREED_EXTS | BIM_EXTS))].copy() if "file_ext" in docs.columns else pd.DataFrame()
    if cad_docs.empty:
        st.info("No CAD/BIM model files detected in the library yet. Upload STEP/STL/OBJ/GLB/IFC.")
        return

    options = [f"{r['doc_id']} — {r['filename']}" for _, r in cad_docs.iterrows()]
    pickA = st.selectbox("Baseline model", options, index=0, key="geomA")
    pickB = st.selectbox("Modified model", options, index=0 if len(options)==1 else 1, key="geomB")

    def _load_doc_bytes(doc_id: str):
        d = get_document(doc_id)
        if not d:
            return "", b""
        pth = d.get("stored_path")
        if not pth or not os.path.exists(pth):
            return d.get("filename",""), b""
        return d.get("filename",""), open(pth,"rb").read()

    docA = pickA.split(" — ")[0].strip()
    docB = pickB.split(" — ")[0].strip()

    if st.button("Analyze & Compare", key="geom_compare"):
        fnA, bA = _load_doc_bytes(docA)
        fnB, bB = _load_doc_bytes(docB)
        if not bA or not bB:
            st.error("Could not load one of the model files.")
            return

        extA = fnA.lower().split(".")[-1]
        extB = fnB.lower().split(".")[-1]

        if extA in BIM_EXTS:
            mA, _ = analyze_ifc_model_bytes(bA, fnA)
        else:
            mA, figA = analyze_3d_model_bytes(bA, fnA)
            if figA is not None:
                st.plotly_chart(figA, use_container_width=True)

        if extB in BIM_EXTS:
            mB, _ = analyze_ifc_model_bytes(bB, fnB)
        else:
            mB, figB = analyze_3d_model_bytes(bB, fnB)
            if figB is not None:
                st.plotly_chart(figB, use_container_width=True)

        # compute delta
        delta = {}
        for k in set(list(mA.keys()) + list(mB.keys())):
            if k in {"filename"}:
                continue
            va, vb = mA.get(k), mB.get(k)
            if isinstance(va, (int,float)) and isinstance(vb, (int,float)):
                delta[k] = vb - va
        st.session_state["geom_compare"] = {"A": mA, "B": mB, "delta": delta}

    res = st.session_state.get("geom_compare")
    if res:
        st.subheader("Metrics")
        c1,c2 = st.columns(2)
        c1.json(res.get("A",{}))
        c2.json(res.get("B",{}))
        st.subheader("Delta")
        st.json(res.get("delta",{}))

def page_plm_sync(identity: Dict[str,str]):
    page_title("PLM Sync & APIs (Overlay)", "Lightweight, plug-in integration layer for Teamcenter/SAP/Siemens: pull snapshots, push approved change packs, keep systems consistent.")

    st.caption("This is an integration overlay. It does not replace PLM/ERP. For demo, it generates API payloads and simulates sync actions.")

    st.subheader("Connection profiles")
    col1,col2 = st.columns(2)
    with col1:
        tc_url = st.text_input("Teamcenter base URL", value=st.session_state.get("tc_url",""), key="tc_url")
        tc_user = st.text_input("Teamcenter user", value=st.session_state.get("tc_user",""), key="tc_user")
    with col2:
        sap_url = st.text_input("SAP OData/BAPI endpoint", value=st.session_state.get("sap_url",""), key="sap_url")
        sap_client = st.text_input("SAP client/tenant", value=st.session_state.get("sap_client",""), key="sap_client")

    st.markdown("---")
    st.subheader("Actions")
    colA,colB,colC = st.columns(3)
    if colA.button("Pull BOM snapshot (SAP)", use_container_width=True, key="plm_pull_bom"):
        # Demo: user already imports BOM via CSV; we just record an event
        _ensure_uedm_state()
        st.session_state.setdefault("plm_events", []).append({"ts": now_iso(), "action": "pull_bom", "system": "SAP", "status": "simulated"})
        st.success("Simulated SAP BOM pull. (For real: call SAP OData/BAPI and map into UEDM.)")

    if colB.button("Pull ECO metadata (Teamcenter)", use_container_width=True, key="plm_pull_eco"):
        st.session_state.setdefault("plm_events", []).append({"ts": now_iso(), "action": "pull_eco", "system": "Teamcenter", "status": "simulated"})
        st.success("Simulated Teamcenter ECO pull.")

    if colC.button("Push approved change pack (Teamcenter + SAP)", use_container_width=True, key="plm_push_pack"):
        st.session_state.setdefault("plm_events", []).append({"ts": now_iso(), "action": "push_pack", "system": "Teamcenter+SAP", "status": "simulated"})
        st.success("Simulated push of an approved change package: updated procedure + drawing pack + BOM updates.")

    st.subheader("Generated API payload (example)")
    sample_payload = {
        "change_id": st.session_state.get("current_change_id", "ECO-XXXX"),
        "user": identity.get("user"),
        "role": identity.get("role"),
        "inputs": {
            "bom_snapshot": "(reference to snapshot_id)",
            "drawing_pack": ["STEP","PDF","DXF","SVG"],
            "procedure_pdf": True,
        },
        "sync": {
            "teamcenter": {"endpoint": tc_url, "operation": "update_ECO", "mode": "bi-directional"},
            "sap": {"endpoint": sap_url, "operation": "update_BOM", "mode": "bi-directional"},
        },
        "audit": {"generated_at": now_iso(), "traceability": True}
    }
    st.code(json.dumps(sample_payload, indent=2), language="json")

    with st.expander("Sync event log", expanded=False):
        st.write(st.session_state.get("plm_events", []))

def page_change_workspace(identity: Dict[str, str]):
    page_title("Change Workspace (Create / Manage ECR & ECO)", "Capture intent → compute impacts → generate a full change package with traceability.")
    st.markdown("#### Create a new change")
    col1, col2 = st.columns([1, 1])
    with col1:
        change_type = st.selectbox("Type", ["ECR", "ECO"])
        title = st.text_input("Title", value="")
        owner = st.text_input("Change owner", value=identity["user"])
        requester = identity["user"]
    with col2:
        # capture design intent
        objective = st.text_area("Objective (design intent)", value="", height=80)
        problem_statement = st.text_area("Problem statement / reason", value="", height=80)

    parts = list_parts()
    affected_parts = st.multiselect("Affected part numbers", options=parts["part_number"].tolist() if not parts.empty else [])
    scope = st.text_area("Scope", value="")
    change_category = st.selectbox("Change category", ["Design", "Material", "Interface", "Process", "Documentation", "Supplier"])
    compliance_context = st.multiselect("Compliance context", ["AS9100", "IATF 16949", "ISO 9001", "EN 1090", "CE", "FAA/EASA", "GD&T", "BOM/Config control"])
    impacted_processes = st.text_input("Impacted processes (comma-separated)", value="Manufacturing, Quality, Supply Chain")
    schedule_weeks = st.number_input("Rough schedule impact (weeks)", min_value=0, max_value=52, value=4)
    cost_estimate = st.number_input("Rough cost impact estimate", min_value=0.0, value=0.0, step=1000.0)
    supplier_change = st.checkbox("Supplier impacted/changed", value=False)

    validation_strategy = st.text_area("Validation strategy", value="Prototype build + inspection + functional test")
    validation_owner = st.text_input("Validation owner", value="Engineering")

    if st.button("Create change", type="primary", disabled=(not title.strip())):
        intent = {
            "requester": requester,
            "objective": objective,
            "problem_statement": problem_statement,
            "scope": scope,
            "affected_parts": affected_parts,
            "change_category": change_category,
            "compliance_context": compliance_context,
            "impacted_processes": [p.strip() for p in impacted_processes.split(",") if p.strip()],
            "schedule_weeks": int(schedule_weeks),
            "cost_estimate": float(cost_estimate),
            "supplier_change": bool(supplier_change),
            "validation_strategy": validation_strategy,
            "validation_owner": validation_owner,
            "assumptions": [],
            "constraints": [],
        }
        cid = create_change(change_type, title, requester, owner, intent)
        st.success(f"Created {change_type}: {cid}")
        st.session_state["active_change_id"] = cid
        st.rerun()

    st.markdown("---")
    st.markdown("#### Existing changes")
    changes = list_changes()
    if changes.empty:
        st.info("No changes yet.")
        return

    sel = st.selectbox("Select a change", changes["change_id"].tolist(), index=0 if "active_change_id" not in st.session_state else max(0, changes.index[changes["change_id"] == st.session_state["active_change_id"]].tolist()[0] if (st.session_state["active_change_id"] in changes["change_id"].tolist()) else 0),
                     format_func=lambda x: f"{changes.loc[changes['change_id']==x,'change_type'].values[0]} | {changes.loc[changes['change_id']==x,'title'].values[0]} | {changes.loc[changes['change_id']==x,'status'].values[0]}")

    change = get_change(sel)
    if not change:
        return

    st.session_state["active_change_id"] = sel

    st.markdown("##### Change details")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Type", change["change_type"])
    c2.metric("Status", change["status"])
    c3.metric("Requester", change["requester"])
    c4.metric("Owner", change["owner"])

    with st.expander("Intent (editable)", expanded=False):
        intent = change["intent"] or {}
        intent["objective"] = st.text_area("Objective", value=intent.get("objective", ""), height=80, key=f"obj_{sel}")
        intent["problem_statement"] = st.text_area("Problem statement", value=intent.get("problem_statement", ""), height=80, key=f"ps_{sel}")
        intent["scope"] = st.text_area("Scope", value=intent.get("scope", ""), height=80, key=f"sc_{sel}")
        intent["affected_parts"] = st.multiselect("Affected parts", options=list_parts()["part_number"].tolist(), default=intent.get("affected_parts", []), key=f"ap_{sel}")
        if st.button("Save intent changes", key=f"save_intent_{sel}"):
            update_change_fields(sel, intent_json=json.dumps(intent, ensure_ascii=False))
            log_action(sel, "UPDATE_INTENT", identity["user"], notes="Updated intent")
            st.success("Saved.")
            st.rerun()

    st.markdown("##### Evidence retrieval (traceable)")
    q = st.text_input("Search your ingested documents (local TF-IDF)", value=f"{change['title']} {change['intent'].get('objective','')}"[:180])
    retrievals = retrieve(q, top_k=6)
    if retrievals:
        ev_df = pd.DataFrame([{"score": r.score, "doc_id": r.doc_id, "chunk_id": r.chunk_id, "text": r.text[:220] + ("..." if len(r.text) > 220 else "")} for r in retrievals])
        st.dataframe(ev_df, use_container_width=True, height=180)
    else:
        st.info("No retrieval results yet. Ingest documents first, or install scikit-learn.")

    st.markdown("##### Compute impacts")
    G = build_graph()
    affected = change["intent"].get("affected_parts", []) or []
    depth = st.slider("Propagation depth", min_value=1, max_value=5, value=3)
    if st.button("Compute impacts", type="secondary"):
        impacts = compute_impacts(G, affected_parts=affected, depth=depth)
        evidence = {
            "retrieval": [asdict(r) for r in retrievals],
            "docs": sorted(list({r.doc_id for r in retrievals})),
        }
        update_change_fields(sel, impacted_json=json.dumps({"impacts": impacts}, ensure_ascii=False), evidence_json=json.dumps(evidence, ensure_ascii=False))
        log_action(sel, "COMPUTE_IMPACTS", identity["user"], notes="Computed graph impacts")
        st.success("Impacts computed.")
        st.rerun()

    impacted = change.get("impacted", {}) or {}
    impacts = impacted.get("impacts", {}) if isinstance(impacted, dict) else {}
    if impacts:
        st.write("Impact summary")
        st.json(impacts.get("impact_summary", {}))
        with st.expander("Impacted lists", expanded=False):
            st.write("Impacted parts", impacts.get("impacted_parts", []))
            st.write("Upstream assemblies", impacts.get("upstream_assemblies", []))
            st.write("Downstream components", impacts.get("downstream_components", []))
            st.write("Impacted docs", impacts.get("impacted_docs", []))
        # Graph focus
        focus_nodes = []
        for pn in (impacts.get("affected_parts") or [])[:10]:
            pid = get_part_by_number(pn)
            if pid:
                focus_nodes.append(pid)
        plot_graph(G, focus_nodes=focus_nodes if focus_nodes else None)

    st.markdown("##### Generate full change package (ECR/ECO)")
    model_name = st.session_state.get("openai_model", DEFAULT_MODEL)
    st.caption(f"AI enrichment uses OpenAI Responses API when API key is provided; otherwise a deterministic rule-based package is produced. Model: {model_name}")
    if st.button("Generate package", type="primary"):
        # Ensure impacts exist
        if not impacts:
            st.warning("Compute impacts first.")
        else:
            pkg = build_change_package(change["intent"], impacts, retrievals)
            # store inside impacted_json as 'package' to keep everything together
            stored = {"impacts": impacts, "package": pkg}
            update_change_fields(sel, impacted_json=json.dumps(stored, ensure_ascii=False), risk_json=json.dumps(pkg.get("risk_fmea", {}), ensure_ascii=False))
            log_action(sel, "GENERATE_PACKAGE", identity["user"], notes="Generated change package", payload={"model": model_name, "ai_used": openai_client_from_session() is not None})
            st.success("Generated.")
            st.rerun()

    # Show generated package
    change = get_change(sel)
    impacted = change.get("impacted", {}) or {}
    pkg = impacted.get("package", {}) if isinstance(impacted, dict) else {}
    if pkg:
        with st.expander("Generated package (JSON)", expanded=False):
            st.code(json_dumps_pretty(pkg), language="json")

        st.markdown("##### Release & approval workflow")
        status = change["status"]
        allowed = CHANGE_STATUSES.get(change["change_type"], [])
        new_status = st.selectbox("Set status", allowed, index=allowed.index(status) if status in allowed else 0)
        if st.button("Apply status", type="secondary"):
            if can_transition(change["change_type"], new_status):
                update_change_fields(sel, status=new_status)
                log_action(sel, "STATUS", identity["user"], notes=f"Status set to {new_status}")
                st.success("Updated.")
                st.rerun()
            else:
                st.error("Invalid transition.")

        st.markdown("##### Audit log")
        st.dataframe(get_actions(sel), use_container_width=True, height=220)



def _bytes_from_doc_or_upload(upload, doc_id: Optional[str]) -> Tuple[Optional[bytes], Optional[str]]:
    """Return (bytes, filename) from either an uploader object or a doc_id from library."""
    if upload is not None:
        try:
            return upload.getvalue(), getattr(upload, "name", "upload")
        except Exception:
            return None, None
    if doc_id:
        doc = get_document(doc_id)
        if doc and doc.get("stored_path") and os.path.exists(doc["stored_path"]):
            with open(doc["stored_path"], "rb") as f:
                return f.read(), doc.get("filename", doc_id)
    return None, None


def _render_to_image(file_bytes: bytes, filename: str, pdf_page: int = 0, render_zoom: float = 3.0) -> Optional["Image.Image"]:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        if fitz is None:
            return None
        png = render_pdf_page_to_png_bytes(file_bytes, pageno=int(max(0, pdf_page)), zoom=float(render_zoom))
        if not png:
            return None
        return read_image(png)
    return read_image(file_bytes)


def _ocr_image_text(img: "Image.Image", lang: str = "eng") -> Tuple[str, float]:
    """OCR with mean confidence (0..1)."""
    if not PYTESSERACT_AVAILABLE or pytesseract is None:
        return "", 0.0
    try:
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, lang=lang)
        words = []
        confs = []
        for w, c in zip(data.get("text", []), data.get("conf", [])):
            w = (w or "").strip()
            try:
                c = float(c)
            except Exception:
                c = -1
            if w and c >= 0:
                words.append(w)
                confs.append(c)
        mean_conf = (float(np.mean(confs)) / 100.0) if confs else 0.0
        text_out = " ".join(words)
        return text_out, float(max(0.0, min(1.0, mean_conf)))
    except Exception:
        return "", 0.0


def ocr_image_to_data(png_bytes: bytes, lang: str = "eng") -> Dict[str, Any]:
    """
    OCR an image (PNG/JPG bytes) and return structured data.
    Returns keys:
      - available: bool
      - text: str (joined high-confidence tokens)
      - mean_conf: float (0..1)
      - words: list[ {text,left,top,width,height,conf} ]
      - engine: str
      - lang: str
      - error: str (optional)
    """
    if not PYTESSERACT_AVAILABLE or pytesseract is None:
        return {"available": False, "text": "", "mean_conf": 0.0, "words": [], "engine": "tesseract", "lang": lang,
                "error": "pytesseract not installed."}

    try:
        if Image is None:
            return {"available": False, "text": "", "mean_conf": 0.0, "words": [], "engine": "tesseract", "lang": lang,
                    "error": "PIL/pillow not installed."}

        img = Image.open(io.BytesIO(png_bytes)).convert("RGB")
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, lang=lang)

        words = []
        confs = []
        tokens = []
        for i in range(len(data.get("text", []))):
            w = (data["text"][i] or "").strip()
            try:
                c = float(data["conf"][i])
            except Exception:
                c = -1.0

            item = {
                "text": w,
                "left": int(data.get("left", [0])[i] or 0),
                "top": int(data.get("top", [0])[i] or 0),
                "width": int(data.get("width", [0])[i] or 0),
                "height": int(data.get("height", [0])[i] or 0),
                "conf": float(c),
            }
            if w:
                words.append(item)
            if w and c >= 0:
                confs.append(c)
                # Keep tokens above a small confidence threshold
                if c >= 30:
                    tokens.append(w)

        mean_conf = (float(np.mean(confs)) / 100.0) if confs else 0.0
        text_out = " ".join(tokens).strip()

        return {
            "available": True,
            "text": text_out,
            "mean_conf": float(max(0.0, min(1.0, mean_conf))),
            "words": words[:5000],
            "engine": "tesseract",
            "lang": lang,
        }
    except Exception as e:
        return {"available": False, "text": "", "mean_conf": 0.0, "words": [], "engine": "tesseract", "lang": lang,
                "error": str(e)}


def ocr_healthcheck() -> Tuple[bool, str]:
    """Check whether pytesseract + the Tesseract engine are available."""
    if not PYTESSERACT_AVAILABLE or pytesseract is None:
        return False, "OCR is disabled: pytesseract is not installed. Install with `pip install pytesseract`."
    try:
        v = pytesseract.get_tesseract_version()
        return True, f"Tesseract detected: {v}"
    except Exception:
        return False, (
            "OCR engine not found. Install **Tesseract OCR** on your machine and ensure it is on PATH. "
            "On Windows, you may need to set `pytesseract.pytesseract.tesseract_cmd` to your tesseract.exe path."
        )


def fallback_visual_narrative(diff_summary: Dict[str, Any], base_ocr: str, mod_ocr: str,
                             attr_deltas: Dict[str, Any], regions: List[Dict[str, Any]]) -> str:
    """Deterministic narrative when AI is not enabled or fails."""
    lines: List[str] = []
    lines.append("### Narrative summary (deterministic)")
    lines.append("")
    lines.append("**What was compared**: baseline vs modified engineering document view (current page/zoom).")
    if diff_summary:
        try:
            lines.append(f"- **Pixel change ratio**: {float(diff_summary.get('pixel_change_ratio', 0.0)):.3f}")
        except Exception:
            pass
        if diff_summary.get('ssim') is not None:
            lines.append(f"- **SSIM** (structural similarity): {diff_summary.get('ssim')}")
        lines.append(f"- **Changed regions detected**: {diff_summary.get('n_regions', 0)}")

    if regions:
        top = regions[0]
        lines.append(
            f"- **Most-changed region** (approx.): "
            f"x0={top.get('x0')} y0={top.get('y0')} x1={top.get('x1')} y1={top.get('y1')} "
            f"change_ratio={top.get('change_ratio')}"
        )

    added = (attr_deltas or {}).get("added", {})
    removed = (attr_deltas or {}).get("removed", {})
    def _count(d):
        return sum(len(v or []) for v in (d or {}).values()) if isinstance(d, dict) else 0
    n_added = _count(added); n_removed = _count(removed)

    if n_added or n_removed:
        lines.append(f"- **Attribute deltas (OCR-based)**: added={n_added}, removed={n_removed}")
        for k in ["dimensions", "angles", "tolerances", "diameters", "threads"]:
            a = (added or {}).get(k, []) or []
            r = (removed or {}).get(k, []) or []
            if a or r:
                lines.append(f"  - {k}: +{len(a)} / -{len(r)}")
    else:
        lines.append("- **Attribute deltas (OCR-based)**: none detected (or OCR unavailable for this view).")

    if not (base_ocr or "").strip() and not (mod_ocr or "").strip():
        lines.append("")
        lines.append("> Note: No readable text was extracted from the current view. If this is a drawing PDF, enable OCR and/or increase render zoom for better title-block capture.")

    return "\n".join(lines)


def fallback_visual_classification(diff_summary: Dict[str, Any], attr_deltas: Dict[str, Any]) -> Dict[str, Any]:
    """Rule-based classification as a fallback when AI is unavailable."""
    added = (attr_deltas or {}).get("added", {})
    removed = (attr_deltas or {}).get("removed", {})
    def has(k: str) -> bool:
        return bool((added or {}).get(k) or (removed or {}).get(k))

    cats: List[str] = []
    if has("dimensions"):
        cats.append("dimension")
    if has("tolerances"):
        cats.append("tolerance")
    if has("angles"):
        cats.append("angle")
    if has("diameters") or has("threads"):
        cats.append("geometry")

    pcr = 0.0
    try:
        pcr = float((diff_summary or {}).get("pixel_change_ratio", 0.0) or 0.0)
    except Exception:
        pcr = 0.0

    if not cats and pcr > 0.02:
        cats.append("layout")
    if not cats:
        cats = ["other"]

    return {
        "model_report": {
            "mode": "deterministic_fallback",
            "confidence_overall": 0.55 if cats != ["other"] else 0.35,
            "assumptions": [
                "Classification is based on OCR attribute deltas (if available) and pixel-level change metrics."
            ],
            "limitations": [
                "Without CAD-native geometry data, true geometric deltas depend on OCR readability.",
                "For image-only drawings, install/enable OCR and zoom into title blocks/dimension callouts.",
            ],
        },
        "changes": [
            {
                "category": c,
                "description": f"Detected potential {c} changes based on extracted deltas/metrics.",
                "baseline_value": None,
                "modified_value": None,
                "delta": None,
                "confidence": 0.65 if c != "other" else 0.35,
            } for c in cats
        ],
    }


def _downscale_for_llm(png_bytes: bytes, max_side: int = 1400) -> bytes:
    if Image is None:
        return png_bytes
    img = read_image(png_bytes)
    if img is None:
        return png_bytes
    W, H = img.size
    scale = min(1.0, float(max_side) / float(max(W, H)))
    if scale < 1.0:
        img = img.resize((int(W * scale), int(H * scale)), resample=Image.BICUBIC)
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def llm_vision_diff(b_png: bytes, m_png: bytes, base_ocr: str, mod_ocr: str, metrics: Dict[str, Any], model: str) -> Optional[Dict[str, Any]]:
    client = openai_client_from_session()
    if client is None:
        return {
            "error": "no_key",
            "message": "OpenAI key not provided. Enable OpenAI in the sidebar for AI vision classification.",
            "model_report": {"mode": "no_ai", "confidence_overall": 0.0, "assumptions": [], "limitations": ["No API key provided."]},
            "changes": [],
        }

    b_png = _downscale_for_llm(b_png)
    m_png = _downscale_for_llm(m_png)

    b64b = base64.b64encode(b_png).decode("utf-8")
    b64m = base64.b64encode(m_png).decode("utf-8")

    schema_hint = """{
  \"model_report\": { \"confidence_overall\": 0.0, \"assumptions\": [], \"limitations\": [] },
  \"changes\": [
    { \"category\": \"dimension|tolerance|angle|annotation|geometry|layout|other\",
      \"description\": \"\",
      \"baseline_value\": \"\",
      \"modified_value\": \"\",
      \"delta\": \"\",
      \"confidence\": 0.0
    }
  ],
  \"summary\": { \"what_changed\": \"\", \"engineering_implications\": \"\", \"recommended_actions\": [] }
}"""

    system = (
        "You are an engineering drawing comparison and change-control assistant. "
        "You must be conservative and traceable. "
        "Do NOT invent dimensions/tolerances; only report numeric values you can see or that exist in OCR text. "
        "When unsure, say 'uncertain' and lower confidence."
    )

    user = f"""
We are comparing a baseline and a modified engineering drawing/document.

Quantitative metrics (computed by the app):
{json_dumps_pretty(metrics)}

Baseline OCR text (may be incomplete):
{(base_ocr or '')[:4500]}

Modified OCR text (may be incomplete):
{(mod_ocr or '')[:4500]}

Task:
1) Identify and classify changes visible between baseline and modified (dimensions, tolerances, angles, annotations, geometries, layout).
2) Quantify changes when possible (baseline_value, modified_value, delta).
3) Provide confidence per change item (0..1).
4) Provide a short summary and recommended actions.

Return ONLY valid JSON matching the schema:
{schema_hint}
"""

    try:
        resp = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": [{"type": "input_text", "text": system}]},
                {"role": "user", "content": [
                    {"type": "input_text", "text": user},
                    {"type": "input_image", "image_url": f"data:image/png;base64,{b64b}"},
                    {"type": "input_image", "image_url": f"data:image/png;base64,{b64m}"},
                ]},
            ],
            temperature=0.2,
        )
        txt = (getattr(resp, "output_text", "") or "").strip()
        mm = re.search(r"(\{.*\})", txt, flags=re.S)
        if not mm:
            return None
        return json.loads(mm.group(1))
    except Exception as e:
        return {
            "error": "runtime_error",
            "message": str(e),
            "model_report": {"mode": "ai_error", "confidence_overall": 0.0, "assumptions": [], "limitations": ["OpenAI request failed."]},
            "changes": [],
        }


def generate_production_manual_struct(modified_text: str, diff_summary: Dict[str, Any], model: str) -> Dict[str, Any]:
    """AI-first (optional) production manual generation; deterministic fallback if no LLM."""
    client = openai_client_from_session()

    if client is None:
        return {
            "title": "Production Manual — Implementing Engineering Change",
            "sections": [
                {"heading": "1. Overview", "content": "This manual describes how to implement and industrialize the revised design using the modified engineering document as the source of truth."},
                {"heading": "2. Change Summary", "content": json_dumps_pretty(diff_summary)},
                {"heading": "3. Required Updates", "content": "Update relevant work instructions, inspection plans, tooling notes, and BOM/configuration records for the affected items."},
                {"heading": "4. Manufacturing Procedure", "content": "1) Review drawing notes and tolerances\n2) Prepare materials/tools\n3) Execute build/assembly steps per updated drawing\n4) Record in-process checks\n5) Perform final inspection"},
                {"heading": "5. Controls & Quality Gates", "content": "Define critical-to-quality characteristics; add incoming/in-process/final checks; ensure measurement system capability; update SPC where needed."},
                {"heading": "6. Validation & Acceptance", "content": "Prototype build, dimensional inspection, functional test; capture evidence (inspection records, test report) and obtain approvals."},
            ],
            "model_report": {"mode": "template", "confidence": 0.65}
        }

    schema_hint = """{
  \"title\": \"\",
  \"model_report\": { \"mode\": \"llm\", \"confidence\": 0.0, \"key_parameters\": [] },
  \"sections\": [ { \"heading\": \"\", \"content\": \"\" } ]
}"""

    system = (
        "You are an expert manufacturing engineer. "
        "Generate a production manual that is implementable, audit-friendly, and aligned with industrialization best practices. "
        "Be explicit about controls and validation. Do not invent part numbers or measurements; rely on provided context."
    )
    user = f"""
Modified design context (OCR/text extract):
{(modified_text or '')[:6000]}

Detected change summary (from the app):
{json_dumps_pretty(diff_summary)}

Generate a complete end-to-end production manual as structured sections suitable for PDF export.
Return ONLY valid JSON matching the schema:
{schema_hint}
"""

    try:
        resp = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": [{"type": "input_text", "text": system}]},
                {"role": "user", "content": [{"type": "input_text", "text": user}]},
            ],
            temperature=0.2,
        )
        txt = (getattr(resp, "output_text", "") or "").strip()
        mm = re.search(r"(\{.*\})", txt, flags=re.S)
        if not mm:
            raise ValueError("No JSON found")
        return json.loads(mm.group(1))
    except Exception:
        return {
            "title": "Production Manual — Implementing Engineering Change",
            "sections": [
                {"heading": "1. Overview", "content": "This manual describes how to implement and industrialize the revised design."},
                {"heading": "2. Change Summary", "content": json_dumps_pretty(diff_summary)},
                {"heading": "3. Procedure", "content": "Follow updated drawing and work instructions; implement required inspection and validation steps."},
            ],
            "model_report": {"mode": "llm_failed_fallback", "confidence": 0.5}
        }


def write_pdf_production_manual(path: str, manual: Dict[str, Any], images: Optional[Dict[str, bytes]] = None) -> None:
    if canvas is None or A4 is None:
        raise RuntimeError("reportlab not installed")
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    x = 40
    y = h - 55

    def wrap_lines(s: str, max_len: int = 95) -> List[str]:
        out = []
        for para in (s or "").splitlines() or [""]:
            p = para.strip()
            if not p:
                out.append("")
                continue
            while len(p) > max_len:
                out.append(p[:max_len])
                p = p[max_len:]
            out.append(p)
        return out

    def line(txt: str, dy: int = 12, bold: bool = False, size: int = 10) -> None:
        nonlocal y
        if y < 70:
            c.showPage()
            y = h - 55
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(x, y, (txt or "")[:120])
        y -= dy

    line(manual.get("title", "Production Manual"), bold=True, dy=18, size=14)

    mr = manual.get("model_report", {})
    if isinstance(mr, dict):
        line(f"Generated: {now_iso()} | Mode: {mr.get('mode','n/a')} | Confidence: {mr.get('confidence','n/a')}", dy=14, size=9)
    line("", dy=10)

    if images:
        try:
            if "modified" in images:
                img_path = path + "_modified.png"
                with open(img_path, "wb") as f:
                    f.write(images["modified"])
                c.drawImage(img_path, x, y-240, width=520, height=240, preserveAspectRatio=True, anchor='n')
                y -= 255
            if "diff" in images:
                img_path = path + "_diff.png"
                with open(img_path, "wb") as f:
                    f.write(images["diff"])
                c.drawImage(img_path, x, y-240, width=520, height=240, preserveAspectRatio=True, anchor='n')
                y -= 255
        except Exception:
            pass

    for sec in manual.get("sections", []) or []:
        if not isinstance(sec, dict):
            continue
        line(sec.get("heading", ""), bold=True, dy=16, size=12)
        for ln in wrap_lines(str(sec.get("content", "")), max_len=98):
            line(ln, dy=11, bold=False, size=9)
        line("", dy=8)

    c.save()


def page_visual_diff():
    inject_theme()
    page_title(
        "Visual Comparison (Baseline vs. Modified)",
        "Side-by-side, overlay, quantified diff, OCR-based attribute deltas, and AI-assisted change classification."
    )

    docs = list_documents()
    vis_docs = docs[docs["file_ext"].isin(["pdf", "png", "jpg", "jpeg"])] if not docs.empty else docs

    tab_up, tab_lib = st.tabs(["Upload", "Library"])

    base_bytes = mod_bytes = None
    base_name = mod_name = None

    with tab_up:
        st.markdown("#### Upload baseline and modified documents")
        col1, col2 = st.columns(2)
        with col1:
            base_up = st.file_uploader("Baseline (PDF/PNG/JPG)", type=["pdf", "png", "jpg", "jpeg"], key="base_vis_up")
        with col2:
            mod_up = st.file_uploader("Modified (PDF/PNG/JPG)", type=["pdf", "png", "jpg", "jpeg"], key="mod_vis_up")

        # Persist uploads across tab switching
        if base_up is not None:
            try:
                st.session_state["base_vis_up_name"] = getattr(base_up, "name", "baseline")
                st.session_state["base_vis_up_bytes"] = base_up.getvalue() if hasattr(base_up, "getvalue") else base_up.read()
            except Exception:
                pass
        if mod_up is not None:
            try:
                st.session_state["mod_vis_up_name"] = getattr(mod_up, "name", "modified")
                st.session_state["mod_vis_up_bytes"] = mod_up.getvalue() if hasattr(mod_up, "getvalue") else mod_up.read()
            except Exception:
                pass

        if base_up is None and st.session_state.get("base_vis_up_bytes"):
            base_bytes = st.session_state.get("base_vis_up_bytes")
            base_name = st.session_state.get("base_vis_up_name", "baseline")
        else:
            base_bytes, base_name = _bytes_from_doc_or_upload(base_up, None)

        if mod_up is None and st.session_state.get("mod_vis_up_bytes"):
            mod_bytes = st.session_state.get("mod_vis_up_bytes")
            mod_name = st.session_state.get("mod_vis_up_name", "modified")
        else:
            mod_bytes, mod_name = _bytes_from_doc_or_upload(mod_up, None)

    with tab_lib:
        st.markdown("#### Select from ingested library")
        if vis_docs.empty:
            st.info("No visual documents in library yet (PDF/PNG/JPG). Upload first in Ingestion & Library.")
        else:
            col1, col2 = st.columns(2)
            with col1:
                base_doc = st.selectbox(
                    "Baseline document",
                    vis_docs["doc_id"].tolist(),
                    key="base_doc_sel",
                    format_func=lambda x: f"{vis_docs.loc[vis_docs['doc_id']==x,'filename'].values[0]} ({x})",
                )
            with col2:
                mod_doc = st.selectbox(
                    "Modified document",
                    vis_docs["doc_id"].tolist(),
                    key="mod_doc_sel",
                    format_func=lambda x: f"{vis_docs.loc[vis_docs['doc_id']==x,'filename'].values[0]} ({x})",
                )
            if st.button("Load selected documents", type="secondary"):
                bb, bn = _bytes_from_doc_or_upload(None, base_doc)
                mb, mn = _bytes_from_doc_or_upload(None, mod_doc)
                st.session_state["vis_loaded"] = {"base_bytes": bb, "base_name": bn, "mod_bytes": mb, "mod_name": mn}
                st.rerun()

            loaded = st.session_state.get("vis_loaded", {})
            if loaded:
                base_bytes, base_name = loaded.get("base_bytes"), loaded.get("base_name")
                mod_bytes, mod_name = loaded.get("mod_bytes"), loaded.get("mod_name")

    if not base_bytes or not mod_bytes or not base_name or not mod_name:
        st.info("Provide both baseline and modified documents (Upload or Library) to start comparison.")
        return

    st.markdown("#### Rendering & inspection controls")
    r1, r2, r3, r4 = st.columns([1, 1, 1, 1])
    with r1:
        pdf_page = st.number_input("PDF page", min_value=0, value=0, step=1)
    with r2:
        render_zoom = st.slider("Render resolution (zoom)", min_value=1.0, max_value=6.0, value=3.0, step=0.5)
    with r3:
        view_zoom = st.select_slider("Deep zoom", options=[1, 2, 3, 4, 6, 8, 10, 12], value=2)
    with r4:
        overlay_alpha = st.slider("Overlay alpha", min_value=0.0, max_value=1.0, value=0.5, step=0.05)

    cx = st.slider("Pan X", 0.0, 1.0, 0.5, 0.01)
    cy = st.slider("Pan Y", 0.0, 1.0, 0.5, 0.01)

    base_img = _render_to_image(base_bytes, base_name, pdf_page=int(pdf_page), render_zoom=float(render_zoom))
    mod_img = _render_to_image(mod_bytes, mod_name, pdf_page=int(pdf_page), render_zoom=float(render_zoom))
    if base_img is None or mod_img is None:
        st.error("Unable to render. For PDFs, install PyMuPDF (`pip install pymupdf`).")
        return

    base_view = zoom_crop(base_img, zoom=float(view_zoom), cx=float(cx), cy=float(cy))
    mod_view = zoom_crop(mod_img, zoom=float(view_zoom), cx=float(cx), cy=float(cy))

    overlay_png = make_overlay_png(base_view, mod_view, alpha=float(overlay_alpha))
    diff_png = compute_diff_png(base_view, mod_view, contrast=2.2)

    st.markdown("#### Side-by-side and overlay")
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        st.image(base_view, caption=f"Baseline (zoomed): {base_name}", use_container_width=True)
    with c2:
        st.image(mod_view, caption=f"Modified (zoomed): {mod_name}", use_container_width=True)
    with c3:
        st.image(overlay_png, caption="Overlay view", use_container_width=True)

    st.markdown("#### Quantified diff (visual + data-driven)")
    threshold = st.slider("Diff sensitivity threshold", 5, 80, 25, 1)
    metrics = compute_visual_metrics(base_view, mod_view, threshold=int(threshold))
    st.json(metrics)

    tile = st.slider("Quantification tile size", 48, 256, 96, 16)
    top_k = st.slider("Top change regions", 5, 60, 20, 1)
    tile_df, ann_png = tile_change_analysis(base_view, mod_view, tile=int(tile), threshold=int(threshold), top_k=int(top_k))

    cc1, cc2 = st.columns(2)
    with cc1:
        st.image(diff_png, caption="Diff highlight (enhanced)", use_container_width=True)
        st.download_button("Download diff PNG", data=diff_png, file_name="diff.png")
    with cc2:
        st.image(ann_png, caption="Most-changed regions (quantified)", use_container_width=True)
        st.download_button("Download annotated PNG", data=ann_png, file_name="diff_regions.png")

    if not tile_df.empty:
        st.dataframe(tile_df, use_container_width=True, height=240)
    else:
        st.info("No significant changes detected at the current threshold/zoom.")

    st.markdown("#### Attribute-level change detection (OCR-based)")
    lang = st.selectbox("OCR language", ["eng"], index=0)
    run_ocr = st.checkbox("Run OCR attribute extraction", value=True, help="Recommended for drawings/title blocks and annotations.")
    base_ocr = mod_ocr = ""
    base_conf = mod_conf = 0.0
    attr_df = pd.DataFrame()

    if run_ocr:
        base_ocr, base_conf = _ocr_image_text(base_view, lang=lang)
        mod_ocr, mod_conf = _ocr_image_text(mod_view, lang=lang)
        base_attr = extract_engineering_attributes(base_ocr)
        mod_attr = extract_engineering_attributes(mod_ocr)
        attr_df = diff_attributes(base_attr, mod_attr, base_conf=base_conf, mod_conf=mod_conf)

        a1, a2 = st.columns(2)
        with a1:
            st.caption(f"Baseline OCR mean confidence: {base_conf:.2f}")
            st.text_area("Baseline OCR text (preview)", value=base_ocr[:2500], height=140)
        with a2:
            st.caption(f"Modified OCR mean confidence: {mod_conf:.2f}")
            st.text_area("Modified OCR text (preview)", value=mod_ocr[:2500], height=140)

        if not attr_df.empty:
            st.dataframe(attr_df, use_container_width=True, height=240)
        else:
            st.info("No attribute-level deltas detected from OCR at the current zoom region.")

    
    st.markdown("#### AI-powered classification & narrative")
    st.caption("AI is optional. The app always provides deterministic change insights; enable OpenAI in the sidebar for richer classification.")

    use_ai = st.checkbox("Use AI vision diff classification (requires OpenAI key)", value=False)
    model_name = (st.session_state.get("openai_model") or DEFAULT_MODEL).strip() or DEFAULT_MODEL

    # Build deterministic deltas for classification/narrative
    base_attr = extract_engineering_attributes(base_ocr) if (base_ocr or "").strip() else {}
    mod_attr = extract_engineering_attributes(mod_ocr) if (mod_ocr or "").strip() else {}
    attr_deltas: Dict[str, Any] = {"added": {}, "removed": {}}
    for k in set(list(base_attr.keys()) + list(mod_attr.keys())):
        b = list(dict.fromkeys((base_attr.get(k) or [])))
        m = list(dict.fromkeys((mod_attr.get(k) or [])))
        add = sorted(list(set(m) - set(b)))
        rem = sorted(list(set(b) - set(m)))
        if add:
            attr_deltas["added"][k] = add
        if rem:
            attr_deltas["removed"][k] = rem

    regions_list = tile_df.to_dict(orient="records") if "tile_df" in locals() and not tile_df.empty else []
    st.session_state["last_visual_regions"] = regions_list
    st.session_state["last_visual_attr_deltas"] = attr_deltas

    # AI call (optional)
    ai_out: Optional[Dict[str, Any]] = None
    if use_ai:
        base_view_png = io.BytesIO(); base_view.save(base_view_png, format="PNG")
        mod_view_png = io.BytesIO(); mod_view.save(mod_view_png, format="PNG")
        ai_out = llm_vision_diff(
            b_png=base_view_png.getvalue(),
            m_png=mod_view_png.getvalue(),
            base_ocr=base_ocr,
            mod_ocr=mod_ocr,
            metrics=metrics,
            model=model_name,
        )

        if isinstance(ai_out, dict) and ai_out.get("error"):
            st.warning(f"AI classification issue: {ai_out.get('message', 'Unknown error')}")
            if ai_out.get("raw"):
                with st.expander("Raw AI output (debug)", expanded=False):
                    st.code(str(ai_out.get("raw"))[:6000])
            ai_out = None
        elif ai_out:
            with st.expander("AI classification output (JSON)", expanded=False):
                st.json(ai_out)

    # Deterministic fallback classification (always available)
    if (not use_ai) or (ai_out is None):
        ai_out = fallback_visual_classification(
            {"pixel_change_ratio": metrics.get("pixel_change_ratio"), "ssim": metrics.get("ssim"), "n_regions": len(regions_list)},
            attr_deltas,
        )

    # Narrative (always available)
    narrative = fallback_visual_narrative(
        {"pixel_change_ratio": metrics.get("pixel_change_ratio"), "ssim": metrics.get("ssim"), "n_regions": len(regions_list)},
        base_ocr,
        mod_ocr,
        attr_deltas,
        regions_list,
    )
    with st.expander("Narrative of the compared documents", expanded=True):
        st.markdown(narrative)

    st.markdown("#### One-click production manual (PDF) from the modified design")
    include_images = st.checkbox("Include modified + diff images in the PDF", value=True)
    if st.button("Generate Production Manual PDF", type="primary"):
        manual = generate_production_manual_struct(modified_text=mod_ocr or "", diff_summary=diff_summary, model=model_name)
        out_dir = os.path.join(EXPORTS_DIR, "production_manuals")
        os.makedirs(out_dir, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_path = os.path.join(out_dir, f"production_manual_{stamp}.pdf")

        images = None
        if include_images:
            mod_view_png = io.BytesIO(); mod_view.save(mod_view_png, format="PNG")
            images = {"modified": mod_view_png.getvalue(), "diff": diff_png}

        write_pdf_production_manual(pdf_path, manual, images=images)
        with open(pdf_path, "rb") as f:
            st.download_button("Download Production Manual (PDF)", data=f.read(), file_name=os.path.basename(pdf_path))
        st.caption("Manual includes a model report (mode/confidence) and is generated from the modified document context.")

    with st.expander("Ask the AI assistant about this comparison", expanded=False):
        page_ai_assistant(context_mode="last_visual")




def deterministic_narrative(context_blocks: List[Dict[str, str]], question: str = "") -> str:
    """Best-effort narrative without an LLM."""
    lines: List[str] = []
    if question.strip():
        lines.append(f"Question: {question.strip()}\n")
    lines.append("Narrative (LLM unavailable, using extracted context):")
    for b in context_blocks[:12]:
        bid = b.get("id", "context")
        txt = (b.get("text", "") or "").strip()
        if not txt:
            continue
        snippet = txt[:900].replace("\n\n", "\n")
        lines.append(f"\n[{bid}]\n{snippet}")
    lines.append("\nTo enable AI reasoning/classification, set an OpenAI API key in the sidebar and re-run the question.")
    return "\n".join(lines)

def llm_answer_with_evidence(
    question: str,
    context_blocks: List[Dict[str, str]],
    model: str,
    fast_mode: bool = True,
    max_ctx_blocks: int = 8,
    max_chars_per_block: int = 1200,
) -> Tuple[str, List[Dict[str, str]]]:
    """Return (answer, evidence_used). Evidence identifiers must be cited like [id]."""
    client = openai_client_from_session()
    if client is None:
        return "", []

    ranked = rank_context_blocks(question, context_blocks, top_k=max_ctx_blocks, max_chars_per_block=max_chars_per_block)
    ctx = "\n\n".join([f"[{b['id']}]\n{b['text']}" for b in ranked])

    system = (
        "You are an AI engineering assistant. "
        "Answer ONLY using the provided context blocks. "
        "Cite sources using their identifiers in square brackets (e.g., [chk_doc_xxx_0001]). "
        "If the answer is not in context, say you don't know."
    )

    style = "Return a short, direct answer. Use bullets when helpful." if fast_mode else "Return a detailed answer with clear sections and bullets."
    user = f"Question:\n{question}\n\nContext blocks:\n{ctx}\n\n{style}"

    try:
        kwargs = {"temperature": 0.2 if fast_mode else 0.25}
        # Cap output tokens when supported
        kwargs["max_output_tokens"] = 450 if fast_mode else 900

        resp = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": [{"type": "input_text", "text": system}]},
                {"role": "user", "content": [{"type": "input_text", "text": user}]},
            ],
            **kwargs,
        )
        answer = (getattr(resp, "output_text", "") or "").strip()
    except TypeError:
        resp = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": [{"type": "input_text", "text": system}]},
                {"role": "user", "content": [{"type": "input_text", "text": user}]},
            ],
            temperature=0.2 if fast_mode else 0.25,
        )
        answer = (getattr(resp, "output_text", "") or "").strip()
    except Exception:
        return "", []

    used = []
    for b in ranked:
        if f"[{b['id']}]" in answer:
            used.append({"id": b["id"], "text": b["text"][:260]})
    return answer, used

def _ensure_png_from_upload(filename: str, file_bytes: bytes, page: int = 0, zoom: float = 3.0) -> Tuple[Optional[bytes], Optional[str]]:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        if fitz is None:
            return None, "pymupdf_missing"
        try:
            png = render_pdf_page_to_png_bytes(file_bytes, pageno=page, zoom=zoom)
            return png, None
        except Exception as e:
            return None, f"{type(e).__name__}: {e}"
    if ext in (".png", ".jpg", ".jpeg"):
        return file_bytes, None
    return None, "unsupported_diagram_type"


def _extract_procedure_text(filename: str, file_bytes: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes) if file_bytes else ""
    if ext == ".docx":
        return extract_text_from_docx(file_bytes) if file_bytes else ""
    if ext in (".txt", ".md", ".csv", ".json"):
        return extract_text_from_txt(file_bytes) if file_bytes else ""
    return extract_text_from_txt(file_bytes) if file_bytes else ""



def detect_geometry(png_bytes: bytes) -> Dict[str, Any]:
    """
    Best-effort geometry detection from a rasterized drawing image.
    Returns a structured summary (no rules about specific drawing standards).
    Uses OpenCV when available; otherwise returns available=False with an error.
    """
    if not CV2_AVAILABLE or cv2 is None:
        return {
            "available": False,
            "n_lines": 0,
            "angles_deg": [],
            "lines_sample": [],
            "edge_density": 0.0,
            "error": "OpenCV not installed. Install with `pip install opencv-python` (or opencv-python-headless).",
        }
    try:
        arr = np.frombuffer(png_bytes, dtype=np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None:
            return {"available": False, "n_lines": 0, "angles_deg": [], "lines_sample": [], "edge_density": 0.0, "error": "Could not decode image bytes."}

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # light denoise while preserving edges
        gray = cv2.GaussianBlur(gray, (3, 3), 0)

        edges = cv2.Canny(gray, threshold1=60, threshold2=160, apertureSize=3, L2gradient=True)
        edge_density = float(np.mean(edges > 0))

        # Detect line segments
        lines = cv2.HoughLinesP(
            edges,
            rho=1,
            theta=np.pi / 180,
            threshold=120,
            minLineLength=max(30, int(min(img.shape[:2]) * 0.04)),
            maxLineGap=8,
        )

        angles = []
        sample = []
        if lines is not None:
            for l in lines[:200]:
                x1, y1, x2, y2 = [int(v) for v in l[0]]
                dx = (x2 - x1)
                dy = (y2 - y1)
                ang = float(np.degrees(np.arctan2(dy, dx)))
                # normalize angle to [0,180)
                ang = ang % 180.0
                angles.append(ang)
                # keep small sample for display
                if len(sample) < 30:
                    length = float(np.hypot(dx, dy))
                    sample.append({"x1": x1, "y1": y1, "x2": x2, "y2": y2, "angle_deg": ang, "length_px": length})

        return {
            "available": True,
            "n_lines": int(0 if lines is None else len(lines)),
            "angles_deg": angles[:2000],
            "lines_sample": sample,
            "edge_density": edge_density,
            "image_shape": {"h": int(img.shape[0]), "w": int(img.shape[1])},
        }
    except Exception as e:
        return {"available": False, "n_lines": 0, "angles_deg": [], "lines_sample": [], "edge_density": 0.0, "error": str(e)}


def diagram_multimodal_representation(diagram_png_bytes: bytes, ocr_lang: str = "eng") -> Dict[str, Any]:
    rep: Dict[str, Any] = {"ocr": {}, "geometry": {}, "attributes": {}}
    ocr = ocr_image_to_data(diagram_png_bytes, lang=ocr_lang)
    rep["ocr"] = ocr
    txt = (ocr.get("text") or "").strip()
    rep["attributes"] = extract_engineering_attributes(txt) if txt else {"dimensions": [], "angles": [], "tolerances": [], "diameters": [], "threads": [], "notes": []}
    rep["geometry"] = detect_geometry(diagram_png_bytes)
    return rep


def llm_plan_diagram_change(rep: Dict[str, Any], procedure_text: str, change_request: str) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
    schema_hint = """{
  \"change_plan\": {
    \"summary\": \"\",
    \"detected_elements_used\": [\"...\"],
    \"operations\": [
      {
        \"type\": \"update_dimension|update_tolerance|update_angle|update_note|move_annotation|geometry_change|other\",
        \"target\": \"\",
        \"from\": \"\",
        \"to\": \"\",
        \"confidence\": 0.0,
        \"evidence\": [\"ocr:text snippet\", \"bbox:(x0,y0,x1,y1)\"]
      }
    ]
  },
  \"image_edit_prompt\": \"\",
  \"procedure_update_instructions\": {
    \"add_steps\": [],
    \"remove_steps\": [],
    \"modify_steps\": []
  },
  \"model_report\": {
    \"overall_confidence\": 0.0,
    \"limitations\": [],
    \"key_parameters\": []
  }
}"""
    system = (
        "You are an expert engineering change assistant specialized in interpreting technical drawings and procedures. "
        "Be conservative: do not invent dimensions; only propose changes consistent with the user's request. "
        "Return only valid JSON."
    )
    user = f"""Diagram representation (OCR + geometry + attributes):
{json_dumps_pretty(rep)}

Procedure (may be partial):
{procedure_text[:8000]}

User change request:
{change_request}

Task:
1) Interpret the requested changes.
2) Create a structured change plan with concrete operations.
3) Produce an image edit prompt that updates ONLY the requested items while preserving all unchanged elements and the overall style.
4) Provide procedure update instructions and a model report with confidence, limitations, and key parameters.
"""
    out, meta = llm_generate_json_debug(system=system, user=user, schema_hint=schema_hint, temperature=0.2, model=st.session_state.get("openai_model", DEFAULT_MODEL))
    if out is None:
        fallback = heuristic_change_plan(rep, change_request)
        return fallback, None
    return out, None


def llm_update_procedure(procedure_text: str, change_plan: Dict[str, Any], change_request: str) -> Tuple[Optional[str], Optional[str]]:
    system = (
        "You are an expert manufacturing engineer. Update procedures precisely and safely. "
        "Return a clean, structured procedure in markdown with numbered steps, plus a 'Controls & Validation' section."
    )
    user = f"""Original procedure:
{procedure_text[:12000]}

User change request:
{change_request}

Structured change plan:
{json_dumps_pretty(change_plan)}

Write the revised procedure. Requirements:
- Preserve unchanged steps.
- Modify/insert/remove only what is necessary.
- Include updated controls, inspection points, and validation evidence references.
"""
    return llm_generate_text(system=system, user=user, temperature=0.2)


def page_diagram_change_update():
    inject_theme()
    ensure_dirs()
    page_title(
        "Diagram Change Update (AI Multimodal)",
        "Upload a baseline diagram + procedure, enter a change request, generate an updated diagram image and revised procedure.",
    )

    # ----------------------------
    # Output management (delete / archive / regenerate)
    # ----------------------------
    with st.expander("Output controls (delete / archive / regenerate)", expanded=False):
        c1, c2, c3, c4 = st.columns([1, 1, 1, 2])
        with c1:
            if st.button("Regenerate", help="Clears generated outputs for a fresh run (keeps your uploads)."):
                dcu_clear_outputs(delete_files=False)
                st.success("Cleared generated outputs. Re-run plan / image / procedure generation.")
                st.rerun()
        with c2:
            if st.button("Archive", help="Copies current generated outputs to the archives folder (does not delete originals)."):
                folder = dcu_archive_outputs()
                if folder:
                    st.success(f"Archived to: {folder}")
                else:
                    st.info("No generated outputs to archive yet.")
        with c3:
            if st.button("Delete", help="Deletes generated output files and clears session outputs.", type="secondary"):
                dcu_clear_outputs(delete_files=True)
                st.success("Deleted generated outputs.")
                st.rerun()
        with c4:
            arts = st.session_state.get("dcu_artifacts", []) or []
            if arts:
                st.caption("Session artifacts")
                for p in arts[-8:]:
                    st.code(p, language="text")
                # One-click ZIP of artifacts
                try:
                    buf = io.BytesIO()
                    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
                        for p in arts:
                            if p and os.path.exists(p):
                                z.write(p, arcname=os.path.basename(p))
                    st.download_button("Download all outputs as ZIP", data=buf.getvalue(), file_name="diagram_change_outputs.zip")
                except Exception:
                    pass
            else:
                st.caption("No generated outputs in this session yet.")

    # ----------------------------
    # Inputs
    # ----------------------------
    st.markdown("#### Inputs")
    colA, colB = st.columns(2)
    with colA:
        base_diag = st.file_uploader(
            "Baseline diagram (PDF/PNG/JPG/CAD)",
            type=["pdf", "png", "jpg", "jpeg"] + CAD_FILE_EXTS,
            key="dcu_base_diag",
        )
    with colB:
        proc_file = st.file_uploader(
            "Procedure doc (PDF/DOCX/TXT/MD)",
            type=["pdf", "docx", "txt", "md"],
            key="dcu_proc",
        )


    # Persist uploads across navigation
    if base_diag is not None:
        try:
            st.session_state["dcu_base_diag_name"] = getattr(base_diag, "name", "baseline")
            st.session_state["dcu_base_diag_bytes"] = base_diag.getvalue() if hasattr(base_diag, "getvalue") else base_diag.read()
        except Exception:
            pass
    if proc_file is not None:
        try:
            st.session_state["dcu_proc_name"] = getattr(proc_file, "name", "procedure")
            st.session_state["dcu_proc_bytes"] = proc_file.getvalue() if hasattr(proc_file, "getvalue") else proc_file.read()
        except Exception:
            pass
    change_request = st.text_area(
        "Change request (textual / numeric / angular / geometric)",
        height=120,
        placeholder="Example: Increase Ø12 ±0.1 → Ø14 ±0.1; rotate feature by +5°; update note A to torque 35Nm.",
        key="dcu_change_request",
    )

    with st.expander("OCR & representation settings", expanded=False):
        ocr_lang = st.selectbox("OCR language", ["eng", "nld", "deu", "fra"], index=0, key="dcu_ocr_lang")
        pdf_page = st.number_input("PDF page (0-indexed)", min_value=0, value=0, step=1, key="dcu_pdf_page")
        pdf_zoom = st.slider("PDF render zoom", 1.0, 4.0, 3.0, 0.1, key="dcu_pdf_zoom")
        st.caption("Tip: increase render zoom for drawings to improve OCR + dimensional extraction.")

    if base_diag is None and not st.session_state.get("dcu_base_diag_bytes"):
        st.info("Upload a baseline diagram to begin.")
        return

    
    # ----------------------------
    # Load baseline diagram bytes -> baseline PNG bytes
    # ----------------------------
    base_bytes = st.session_state.get("dcu_base_diag_bytes", b"")
    base_name = st.session_state.get("dcu_base_diag_name", getattr(base_diag, "name", "baseline"))
    base_name = getattr(base_diag, "name", "baseline")
    base_ext = (base_name.split(".")[-1] if "." in base_name else "").lower()

    # If baseline is CAD, allow optional preview for rendering
    preview_bytes: Optional[bytes] = None
    preview_ext = ""
    if base_ext in CAD_FILE_EXTS:
        st.warning("CAD uploaded. For baseline vs updated visual comparison, upload a PDF/PNG/JPG preview below.")
        cad_preview = st.file_uploader("CAD preview for visual compare (optional)", type=["pdf","png","jpg","jpeg"], key="dcu_cad_preview")
        if cad_preview is None and st.session_state.get("dcu_cad_preview_bytes"):
            preview_bytes = st.session_state.get("dcu_cad_preview_bytes")
            preview_name = st.session_state.get("dcu_cad_preview_name","preview")
            preview_ext = (preview_name.split(".")[-1] if "." in preview_name else "").lower()
        if cad_preview is not None:
            preview_bytes = cad_preview.getvalue() if hasattr(cad_preview, "getvalue") else cad_preview.read()
            st.session_state["dcu_cad_preview_name"] = getattr(cad_preview, "name", "preview")
            st.session_state["dcu_cad_preview_bytes"] = preview_bytes
            preview_name = getattr(cad_preview, "name", "preview")
            preview_ext = (preview_name.split(".")[-1] if "." in preview_name else "").lower()

    render_bytes = preview_bytes if (base_ext in CAD_FILE_EXTS and preview_bytes is not None) else (base_bytes if base_ext not in CAD_FILE_EXTS else None)
    render_ext = preview_ext if (base_ext in CAD_FILE_EXTS and preview_bytes is not None) else (base_ext if base_ext not in CAD_FILE_EXTS else "")

    base_png: Optional[bytes] = None
    if render_bytes is None:
        base_png = None
    elif render_ext == "pdf":
        base_png = render_pdf_page_to_png_bytes(render_bytes, pageno=int(pdf_page), zoom=float(pdf_zoom))
        if not base_png:
            st.error("Could not render PDF. Install PyMuPDF (fitz).")
            return
    else:
        base_png = render_bytes

    # ----------------------------
    # Procedure text extraction
    # ----------------------------
    proc_text = ""
    if proc_file:
        pbytes = st.session_state.get("dcu_proc_bytes", b"")
        pname = getattr(proc_file, "name", "procedure")
        pext = (pname.split(".")[-1] if "." in pname else "").lower()
        if pext == "pdf":
            proc_text = extract_text_from_pdf(pbytes)
        elif pext == "docx":
            proc_text = extract_text_from_docx(pbytes)
        else:
            try:
                proc_text = pbytes.decode("utf-8", errors="ignore")
            except Exception:
                proc_text = ""

    # ----------------------------
    # Multimodal representation (baseline)
    # ----------------------------
    st.markdown("#### Multimodal representation (baseline)")
    rep = diagram_multimodal_representation(base_png, ocr_lang=ocr_lang)

    ocr = rep.get("ocr", {}) if isinstance(rep, dict) else {}
    geom = rep.get("geometry", {}) if isinstance(rep, dict) else {}
    attrs = rep.get("attributes", {}) if isinstance(rep, dict) else {}

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("OCR mean confidence", f"{float(ocr.get('mean_conf', 0.0)):.2f}")
    with c2:
        st.metric("OCR words", str(len(ocr.get("words") or [])))
    with c3:
        st.metric("Lines detected", str(int((geom or {}).get("n_lines") or 0)))
    with c4:
        st.metric("Dims extracted", str(len((attrs or {}).get("dimensions") or [])))

    with st.expander("Extracted attributes (dimensions/angles/tolerances/threads)", expanded=False):
        st.json(attrs or {})

    if not (change_request or "").strip():
        st.warning("Enter a change request to generate a plan, impact analysis, and a modified diagram for comparison.")
        return

    # ----------------------------
    # Planning
    # ----------------------------
    ai_ready = openai_client_from_session() is not None
    if not ai_ready:
        st.info("OpenAI key not set. AI planning/procedure update will use fallbacks. Diagram generation can still run via Local SD or Rule-based overlay.")

    cols = st.columns([1, 2])
    with cols[0]:
        if st.button("Generate change plan", type="secondary"):
            if ai_ready:
                plan, plan_err = llm_plan_diagram_change(rep, proc_text, change_request)
                if plan_err:
                    st.warning(f"AI planning had issues: {plan_err}. A fallback plan may still be produced.")
                if not plan:
                    plan = heuristic_change_plan(rep, change_request)
            else:
                plan = heuristic_change_plan(rep, change_request)
            st.session_state["dcu_plan"] = plan
            st.success("Change plan ready.")
            st.rerun()

    plan = st.session_state.get("dcu_plan")
    if not plan:
        st.stop()

    with cols[1]:
        conf = (plan.get("model_report", {}) or {}).get("overall_confidence", None) if isinstance(plan, dict) else None
        if isinstance(conf, (int, float)):
            st.metric("Plan confidence", f"{conf:.2f}")
        limitations = (plan.get("model_report", {}) or {}).get("limitations", []) if isinstance(plan, dict) else []
        if limitations:
            st.caption("Model limitations")
            st.write("• " + "\n• ".join([str(x) for x in limitations[:6]]))

    with st.expander("Change plan (JSON)", expanded=False):
        st.code(json_dumps_pretty(plan), language="json")

    # ----------------------------
    # Diagram regeneration
    # ----------------------------
    st.markdown("#### Updated diagram generation (for visual comparison)")
    prompt = plan.get("image_edit_prompt", "") if isinstance(plan, dict) else ""
    if not prompt:
        prompt = (
            "Modify this technical drawing to apply the change request exactly while preserving all unchanged geometry, "
            "dimensions, annotations, title block, line weights, and layout. Keep it crisp black-and-white.\n"
            f"Change request: {change_request}"
        )

    backend = st.selectbox(
        "Diagram regeneration backend",
        [
            "OpenAI GPT Image (requires verified org)",
            "Local Stable Diffusion (AUTOMATIC1111 img2img)",
            "Rule-based overlay (no AI, OCR-driven)",
        ],
        index=0,
        key="dcu_backend",
        help="If OpenAI image models are blocked, use Local SD or Rule-based overlay to still generate a modified image for comparison.",
    )

    size = st.selectbox("Output image size", ["1024x1024", "1536x1536"], index=0, key="dcu_size")
    w, h = [int(x) for x in size.split("x", 1)]

    # Local SD settings
    sd_url = st.session_state.get("dcu_sd_url", "http://127.0.0.1:7860")
    sd_denoise = float(st.session_state.get("dcu_sd_denoise", 0.45))
    sd_steps = int(st.session_state.get("dcu_sd_steps", 30))
    sd_cfg = float(st.session_state.get("dcu_sd_cfg", 7.0))

    if backend.startswith("Local Stable Diffusion"):
        with st.expander("Local Stable Diffusion settings", expanded=False):
            sd_url = st.text_input("WebUI API base URL", value=sd_url, key="dcu_sd_url")
            sd_denoise = st.slider("Denoising strength", 0.05, 0.90, sd_denoise, 0.01, key="dcu_sd_denoise")
            sd_steps = st.slider("Steps", 10, 80, sd_steps, 1, key="dcu_sd_steps")
            sd_cfg = st.slider("CFG scale", 1.0, 14.0, sd_cfg, 0.5, key="dcu_sd_cfg")
            st.caption("Tip: For drawings, denoise 0.30–0.55 preserves geometry/linework while applying edits.")

    can_generate = True
    if backend.startswith("OpenAI") and not ai_ready:
        can_generate = False
        st.warning("OpenAI key not set — switch backend to Local SD or Rule-based overlay, or set your key in the sidebar.")

    if st.button("Generate updated diagram image", type="primary", disabled=not can_generate):
        img_bytes, img_err = None, None

        if backend.startswith("OpenAI"):
            img_bytes, img_err = openai_image_edit_or_generate(base_png, prompt=prompt, size=size, prefer_edit=True)
            if img_err == "org_verification_required":
                st.error("OpenAI image generation is blocked because your OpenAI organization is not verified for GPT Image models.")
                st.info("Switch backend to Local Stable Diffusion or Rule-based overlay to proceed.")
                img_bytes = None

        elif backend.startswith("Local Stable Diffusion"):
            img_bytes, img_err = sd_webui_img2img(
                init_png=base_png,
                prompt=prompt,
                url=sd_url,
                denoising_strength=sd_denoise,
                steps=sd_steps,
                cfg_scale=sd_cfg,
                width=w,
                height=h,
            )
            if img_err == "requests_not_installed":
                st.error("Missing dependency: requests. Run `pip install requests` or reinstall requirements.txt.")

        else:
            img_bytes, img_err = overlay_rule_based_edit(
                base_png=base_png,
                rep=rep,
                plan=plan if isinstance(plan, dict) else {"change_plan": {"operations": []}},
                change_request=change_request,
            )

        if img_err or not img_bytes:
            st.error(f"Diagram generation failed: {img_err or 'no_image_returned'}")
        else:
            st.session_state["dcu_updated_diagram"] = img_bytes
            # persist artifact
            out_path = os.path.join(EXPORTS_DIR, f"updated_diagram_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
            try:
                with open(out_path, "wb") as f:
                    f.write(img_bytes)
                dcu_register_artifact(out_path)
            except Exception:
                pass
            st.success("Updated diagram generated.")
            st.rerun()

    updated_png = st.session_state.get("dcu_updated_diagram")

    # ----------------------------
    # Visual comparison (advanced + zoom + dimensional deltas)
    # ----------------------------
    if updated_png:
        st.markdown("### Visual comparison (Baseline vs Updated)")
        zoom_h = st.slider("Viewer height", 420, 960, 640, 20, help="Larger height = easier inspection of dimensions.")
        tab_zoom, tab_side, tab_overlay, tab_diff, tab_dim = st.tabs(
            ["Zoom & Inspect", "Side-by-side", "Overlay", "Diff + Metrics", "Dimensional delta"]
        )

        with tab_zoom:
            col1, col2 = st.columns(2)
            with col1:
                zoomable_image_viewer(base_png, height=int(zoom_h), label="Baseline (zoom/pan)")
            with col2:
                zoomable_image_viewer(updated_png, height=int(zoom_h), label="Updated (zoom/pan)")

        with tab_side:
            col1, col2 = st.columns(2)
            with col1:
                if base_png is None:
                    st.warning("Baseline preview unavailable (no baseline image/PDF rendered).")
                else:
                    safe_st_image(base_png, caption="Baseline", use_container_width=True)
            with col2:
                if updated_png is None:
                    st.warning("Updated preview unavailable (no updated image/PDF rendered).")
                else:
                    safe_st_image(updated_png, caption="Updated", use_container_width=True)

        with tab_overlay:
            alpha = st.slider("Overlay alpha (Updated over Baseline)", 0.0, 1.0, 0.55, 0.01)
            blended = blend_png_bytes(base_png, updated_png, alpha=alpha)
            if blended:
                zoomable_image_viewer(blended, height=int(zoom_h), label=f"Overlay (alpha={alpha:.2f})")
                st.download_button("Download overlay PNG", data=blended, file_name="overlay.png")
            else:
                st.info("Overlay unavailable (missing PIL).")

        with tab_diff:
            diff = image_diff(base_png, updated_png)
            if diff:
                zoomable_image_viewer(diff, height=int(zoom_h), label="Diff highlight (zoom/pan)")
                st.download_button("Download diff PNG", data=diff, file_name="diff_highlight.png")
                # persist diff
                try:
                    diff_path = os.path.join(EXPORTS_DIR, f"diff_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
                    with open(diff_path, "wb") as f:
                        f.write(diff)
                    dcu_register_artifact(diff_path)
                except Exception:
                    pass
            st.markdown("#### Quantified impact (visual)")
            metrics = compute_visual_metrics(base_png, updated_png) if (base_png and updated_png) else {"status":"skipped","reason":"Baseline or updated image not available."}
            st.json(metrics)

        with tab_dim:
            st.markdown("#### Detected dimensional changes (OCR-derived)")
            st.caption("This extracts dimension/angle/tolerance tokens from OCR and compares baseline vs updated.")
            if st.button("Analyze updated drawing (OCR + attributes)"):
                rep_u = diagram_multimodal_representation(updated_png, ocr_lang=ocr_lang)
                st.session_state["dcu_updated_rep"] = rep_u
                st.success("Updated drawing analyzed.")
                st.rerun()

            rep_u = st.session_state.get("dcu_updated_rep")
            if rep_u:
                ocr_u = rep_u.get("ocr", {}) if isinstance(rep_u, dict) else {}
                attrs_u = rep_u.get("attributes", {}) if isinstance(rep_u, dict) else {}
                base_conf = float((ocr or {}).get("mean_conf", 0.0))
                mod_conf = float((ocr_u or {}).get("mean_conf", 0.0))
                df = diff_attributes(attrs or {}, attrs_u or {}, base_conf=base_conf, mod_conf=mod_conf)
                if df is None or getattr(df, "empty", True):
                    st.info("No attribute deltas detected (or OCR did not extract comparable tokens).")
                else:
                    q = st.text_input("Search deltas", value="", placeholder="e.g., Ø14, ±0.1, 45")
                    dff = df.copy()
                    if q.strip():
                        mask = dff.apply(lambda r: q.lower() in (str(r.get("baseline")) + " " + str(r.get("modified"))).lower(), axis=1)
                        dff = dff[mask]
                    st.dataframe(dff, use_container_width=True, hide_index=True)
                    st.caption(f"Baseline OCR conf: {base_conf:.2f} • Updated OCR conf: {mod_conf:.2f}")

        st.download_button("Download updated diagram PNG", data=updated_png, file_name="updated_diagram.png")

    # ----------------------------
    # Revised procedure (advanced view + PDF export)
    # ----------------------------
    st.markdown("### Revised procedure")
    if st.button("Generate revised procedure", type="secondary"):
        revised = ""
        if ai_ready:
            revised, perr = llm_update_procedure(proc_text, plan.get("change_plan", plan), change_request)
            if perr:
                st.warning(f"AI procedure update had issues: {perr}. Falling back to heuristic update.")
                revised = ""
        if not revised:
            revised = (proc_text or "").strip()
            ops = (plan.get("change_plan", {}) if isinstance(plan, dict) else {}).get("operations", []) or []
            for op in ops:
                frm = (op.get("from") or "").strip()
                to = (op.get("to") or "").strip()
                if frm and to and revised:
                    revised = revised.replace(frm, to)
            header = f"# Revised Procedure (heuristic)\n\nChange request:\n- {change_request}\n\n---\n\n"
            revised = header + (revised if revised else "(No procedure text provided.)")
        st.session_state["dcu_revised_procedure"] = revised

        # persist revised procedure artifacts
        try:
            md_path = os.path.join(EXPORTS_DIR, f"revised_procedure_pack_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.md")
            with open(md_path, "wb") as f:
                f.write(revised.encode("utf-8"))
            dcu_register_artifact(md_path)
        except Exception:
            pass

        st.success("Revised procedure ready.")
        st.rerun()

    revised_proc = st.session_state.get("dcu_revised_procedure", "")
    if revised_proc:
        proc_tab1, proc_tab2 = st.tabs(["Procedure (clean)", "Procedure diff"])
        with proc_tab1:
            st.markdown(revised_proc)
        with proc_tab2:
            a = (proc_text or "").splitlines()
            b = (revised_proc or "").splitlines()
            # Compact diff rendering (highlight changed lines)
            diff_lines = list(difflib.unified_diff(a, b, fromfile="original", tofile="revised", lineterm=""))
            if diff_lines:
                st.code("\n".join(diff_lines[:2500]), language="diff")
            else:
                st.info("No textual difference detected (or no original procedure provided).")

        st.download_button("Download revised procedure (MD)", data=revised_proc.encode("utf-8"), file_name="revised_procedure.md")
        # Procedure PDF download (pack: includes change details + visuals)
        images_pack: Dict[str, bytes] = {}
        if base_png:
            images_pack["baseline"] = base_png
        if updated_png:
            images_pack["updated"] = updated_png
        if base_png and updated_png:
            try:
                images_pack["diff"] = image_diff(base_png, updated_png)
            except Exception:
                pass

        plan_for_pdf = None
        if isinstance(plan, dict):
            plan_for_pdf = plan.get("change_plan") or plan

        proc_pdf = pdf_bytes_procedure_pack(
            "Revised Procedure Pack",
            revised_proc,
            change_request=change_request,
            change_plan=plan_for_pdf if isinstance(plan_for_pdf, dict) else None,
            images=images_pack,
            user=st.session_state.get("user", ""),
        )
        st.download_button("Download revised procedure pack (PDF)", data=proc_pdf, file_name="revised_procedure_pack_pack.pdf")
        try:
            pdf_path = os.path.join(EXPORTS_DIR, f"revised_procedure_pack_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
            with open(pdf_path, "wb") as f:
                f.write(proc_pdf)
            dcu_register_artifact(pdf_path)
        except Exception:
            pass

    # ----------------------------
    # One-click production manual (PDF pack)
    # ----------------------------
    st.markdown("---")
    st.markdown("### One-click production manual (PDF pack)")
    if st.button("Generate Production Manual PDF", type="primary"):
        ensure_dirs()
        pdf_path = os.path.join(EXPORTS_DIR, f"production_manual_{uuid.uuid4().hex}.pdf")
        manual = {
            "title": "Production Manual — Diagram Change Update",
            "generated_at": now_iso(),
            "change_request": change_request,
            "diagram_representation_summary": {
                "ocr_mean_conf": float((ocr or {}).get("mean_conf", 0.0)),
                "attributes": attrs,
                "geometry": geom,
            },
            "change_plan": plan.get("change_plan") if isinstance(plan, dict) else plan,
            "revised_procedure": revised_proc or "(No revised procedure generated.)",
            "controls_and_validation": (
                "Update inspection points, torque specs, tolerances, and acceptance criteria. "
                "Perform first-article verification on changed features and update control plans if applicable."
            ),
        }
        imgs = {"baseline": base_png}
        if updated_png:
            imgs["updated"] = updated_png
            d = image_diff(base_png, updated_png)
            if d:
                imgs["diff"] = d
        try:
            write_pdf_production_manual(pdf_path, manual, images=imgs)
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            st.download_button("Download Production Manual (PDF)", data=pdf_bytes, file_name=os.path.basename(pdf_path))
            dcu_register_artifact(pdf_path)
            st.success("Production manual generated.")
        except Exception as e:
            st.error(f"PDF generation failed: {type(e).__name__}: {e}")




# ----------------------------
# Enterprise connectors (SAP / Teamcenter) — demo-ready scaffolding
# ----------------------------
def _sync_log(system: str, action: str, ref_id: str, status: str, payload: Dict[str, Any]) -> str:
    try:
        con = db()
        cur = con.cursor()
        sync_id = f"{system}-{dt.datetime.utcnow().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
        cur.execute(
            "INSERT INTO external_sync_log(id, system, action, ref_id, status, payload_json, created_at) VALUES (?,?,?,?,?,?,?)",
            (sync_id, system, action, ref_id, status, json.dumps(payload), dt.datetime.utcnow().isoformat()),
        )
        con.commit()
        con.close()
        return sync_id
    except Exception:
        return ""


def cache_external_bom(system: str, key: str, df: pd.DataFrame) -> str:
    try:
        con = db()
        cur = con.cursor()
        cache_id = f"{system}-BOM-{dt.datetime.utcnow().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
        cur.execute(
            "INSERT INTO external_bom_cache(id, system, key, payload_json, created_at) VALUES (?,?,?,?,?)",
            (cache_id, system, key, df.to_json(orient="records"), dt.datetime.utcnow().isoformat()),
        )
        con.commit()
        con.close()
        return cache_id
    except Exception:
        return ""


def load_external_bom_latest(system: str, key: str) -> Optional[pd.DataFrame]:
    try:
        con = db()
        q = "SELECT payload_json FROM external_bom_cache WHERE system=? AND key=? ORDER BY created_at DESC LIMIT 1"
        row = con.execute(q, (system, key)).fetchone()
        con.close()
        if not row or not row[0]:
            return None
        return pd.read_json(row[0])
    except Exception:
        return None


def sap_fetch_bom_odata(odata_url: str, auth: Optional[tuple] = None, headers: Optional[Dict[str, str]] = None, timeout_s: int = 30) -> pd.DataFrame:
    """
    Fetch BOM from SAP via an OData endpoint (customer-specific).
    Expected JSON shapes: { "d": { "results": [...] } } or { "value": [...] } or a direct list.
    Returns a normalized DataFrame with at least: parent_part_number, part_number, quantity, bom_description, unit_of_measure, cost
    """
    import requests
    hdr = {"Accept": "application/json"}
    if headers:
        hdr.update(headers)
    r = requests.get(odata_url, auth=auth, headers=hdr, timeout=timeout_s)
    r.raise_for_status()
    js = r.json()
    rows = None
    if isinstance(js, list):
        rows = js
    elif isinstance(js, dict):
        if "value" in js and isinstance(js["value"], list):
            rows = js["value"]
        elif "d" in js and isinstance(js["d"], dict) and "results" in js["d"]:
            rows = js["d"]["results"]
        elif "results" in js and isinstance(js["results"], list):
            rows = js["results"]
    rows = rows or []

    df = pd.DataFrame(rows)

    # Heuristic column mapping (customer SAP models vary)
    colmap = {
        "ParentMaterial": "parent_part_number",
        "ParentMaterialNumber": "parent_part_number",
        "BOMHeaderMaterial": "parent_part_number",
        "Material": "part_number",
        "MaterialNumber": "part_number",
        "Component": "part_number",
        "ComponentMaterial": "part_number",
        "ComponentMaterialNumber": "part_number",
        "ComponentDescription": "bom_description",
        "Description": "bom_description",
        "Quantity": "quantity",
        "ComponentQuantity": "quantity",
        "BaseUnit": "unit_of_measure",
        "Unit": "unit_of_measure",
        "UnitOfMeasure": "unit_of_measure",
        "StandardPrice": "cost",
        "Cost": "cost",
    }
    for src, dst in colmap.items():
        if src in df.columns and dst not in df.columns:
            df[dst] = df[src]

    if "parent_part_number" not in df.columns:
        df["parent_part_number"] = ""
    if "part_number" not in df.columns:
        for c in df.columns:
            if "material" in c.lower() or "component" in c.lower() or "part" in c.lower():
                df["part_number"] = df[c].astype(str)
                break
    if "part_number" not in df.columns:
        df["part_number"] = ""

    if "quantity" not in df.columns:
        df["quantity"] = 1
    if "bom_description" not in df.columns:
        df["bom_description"] = ""
    if "unit_of_measure" not in df.columns:
        df["unit_of_measure"] = ""
    if "cost" not in df.columns:
        df["cost"] = 0.0

    df["part_number"] = df["part_number"].astype(str).fillna("")
    df["parent_part_number"] = df["parent_part_number"].astype(str).fillna("")
    df["quantity"] = pd.to_numeric(df["quantity"], errors="coerce").fillna(1).astype(float)
    df["cost"] = pd.to_numeric(df["cost"], errors="coerce").fillna(0).astype(float)
    return df


def sap_update_bom_demo(bom_key: str, ecr_payload: Dict[str, Any]) -> str:
    """Demo update: logs an outbox event representing applying an ECR to SAP."""
    return _sync_log("SAP", "UPDATE_BOM", bom_key, "QUEUED", ecr_payload)


def teamcenter_update_ecr_demo(ecr_id: str, payload: Dict[str, Any]) -> str:
    """Demo update: logs an outbox event representing updating Teamcenter ECR/ECO + attachments."""
    return _sync_log("TEAMCENTER", "UPSERT_ECR", ecr_id, "QUEUED", payload)


def page_connectors() -> None:
    page_title("Enterprise Connectors (SAP BOM + Teamcenter ECR)", "Cohesive SAP → Copilot → Teamcenter workflow (demo-ready)")
    st.info(
        "This page is a cohesive integration scaffold. Use Mock mode for demos. "
        "For production, wire your customer-specific SAP (OData/BAPI/RFC) and Teamcenter (REST/SOA) endpoints."
    )

    tabs = st.tabs(["SAP BOM", "Teamcenter ECR/ECO", "End-to-End Flow", "CAD / Evidence"])

    with tabs[0]:
        st.subheader("SAP — BOM ingestion + update")
        mode = st.radio("Mode", ["Mock (demo)", "OData (HTTP)"], horizontal=True, key="sap_mode")
        bom_key = st.text_input("BOM key (Material/Assembly No.)", value=st.session_state.get("sap_bom_key", "IDEAL-5000-TWINSET"), key="sap_bom_key")

        if mode.startswith("Mock"):
            st.caption("Mock mode uses a BOM CSV from your Document Library or an uploaded CSV.")
            c1, c2 = st.columns(2)

            with c1:
                st.markdown("**Use an existing BOM CSV from the library**")
                docs_df = list_documents()
                if isinstance(docs_df, pd.DataFrame) and "filename" in docs_df.columns and "doc_id" in docs_df.columns:
                    csv_df = docs_df[docs_df["filename"].str.lower().str.endswith(".csv", na=False)]
                    options = ["(none)"] + [f'{row["doc_id"]} — {row["filename"]}' for _, row in csv_df.iterrows()]
                else:
                    options = ["(none)"]
                pick = st.selectbox("Select BOM CSV", options, key="sap_pick_csv")
                if st.button("Load selected BOM into SAP cache"):
                    if pick != "(none)":
                        doc_id = pick.split("—")[0].strip()
                        doc = get_document(doc_id)
                        stored_path = (doc or {}).get("stored_path")
                        if stored_path and os.path.exists(stored_path):
                            try:
                                with open(stored_path, "rb") as f:
                                    b = f.read()
                                df = pd.read_csv(io.BytesIO(b))
                                cache_external_bom("SAP", bom_key, df)
                                st.success("Cached BOM from library as SAP BOM.")
                            except Exception as e:
                                st.error(f"Could not parse CSV: {e}")
                        else:
                            st.error("Could not load the selected document.")
                    else:
                        st.warning("Pick a CSV first.")

            with c2:
                st.markdown("**Or upload a BOM CSV (mock SAP pull)**")
                up = st.file_uploader("Upload BOM CSV", type=["csv"], key="sap_upload_bom_csv")
                if up is not None:
                    try:
                        df = pd.read_csv(up)
                        st.dataframe(df, use_container_width=True, height=260)
                        if st.button("Cache uploaded BOM as SAP BOM"):
                            cache_external_bom("SAP", bom_key, df)
                            st.success("Cached uploaded BOM as SAP BOM.")
                    except Exception as e:
                        st.error(f"Failed to read CSV: {e}")

        else:
            st.caption("OData mode fetches a BOM from an SAP OData endpoint returning JSON.")
            odata_url = st.text_input("SAP OData URL", placeholder="https://<sap-host>/sap/opu/odata/sap/<SERVICE>/...?$format=json", key="sap_odata_url")
            c1, c2 = st.columns(2)
            user = c1.text_input("SAP user (optional)", key="sap_user")
            pwd = c2.text_input("SAP password (optional)", type="password", key="sap_pwd")

            if st.button("Fetch BOM from SAP (OData)"):
                if not odata_url.strip():
                    st.error("Please provide an OData URL.")
                else:
                    try:
                        auth = (user, pwd) if user and pwd else None
                        df = sap_fetch_bom_odata(odata_url, auth=auth)
                        st.success(f"Fetched {len(df)} BOM rows.")
                        st.dataframe(df, use_container_width=True, height=350)
                        cache_external_bom("SAP", bom_key, df)
                        # store evidence
                        csv_bytes = df.to_csv(index=False).encode("utf-8")
                        upsert_document(f"SAP_BOM_{bom_key}.csv", csv_bytes, mime="text/csv")
                        st.info("Saved fetched BOM to library and cached as SAP BOM.")
                    except Exception as e:
                        st.error(f"SAP fetch failed: {type(e).__name__}: {e}")

        st.markdown("---")
        st.subheader("Apply ECR updates back to SAP (demo outbox)")
        st.caption("For demo we log an integration event. In production, call SAP OData/BAPI to update BOM/master data.")
        change_request = st.text_area("Change request (ECR) summary", height=110, key="sap_ecr_summary")
        if st.button("Queue SAP BOM update from ECR"):
            payload = {
                "bom_key": bom_key,
                "change_request": change_request,
                "generated_diagram_present": bool(st.session_state.get("dcu_updated_diagram")),
                "revised_procedure_present": bool(st.session_state.get("dcu_revised_procedure")),
            }
            sync_id = sap_update_bom_demo(bom_key, payload)
            st.success(f"Queued SAP update event: {sync_id or '(logged)'}")

    with tabs[1]:
        st.subheader("Teamcenter — ECR/ECO update + attachments")
        tc_mode = st.radio("Mode", ["Mock (demo)", "REST/SOA (placeholder)"], horizontal=True, key="tc_mode")
        ecr_id = st.text_input("Teamcenter ECR/ECO ID (or desired ID)", value=st.session_state.get("tc_ecr_id", "ECR-DEMO-001"), key="tc_ecr_id")

        st.caption("Push ECR metadata + AI-generated updated diagram + revised procedure to Teamcenter.")
        payload = {
            "ecr_id": ecr_id,
            "change_request": st.session_state.get("sap_ecr_summary", ""),
            "diagram_png_bytes_b64": base64.b64encode(st.session_state.get("dcu_updated_diagram", b"")).decode("utf-8") if st.session_state.get("dcu_updated_diagram") else "",
            "revised_procedure_md": st.session_state.get("dcu_revised_procedure", ""),
            "timestamp": dt.datetime.utcnow().isoformat(),
        }

        if tc_mode.startswith("Mock"):
            if st.button("Queue Teamcenter ECR update (mock)"):
                sync_id = teamcenter_update_ecr_demo(ecr_id, payload)
                st.success(f"Queued Teamcenter update event: {sync_id or '(logged)'}")
                with st.expander("Payload preview"):
                    preview = dict(payload)
                    if preview.get("diagram_png_bytes_b64"):
                        preview["diagram_png_bytes_b64"] = preview["diagram_png_bytes_b64"][:200] + "…"
                    if preview.get("revised_procedure_md"):
                        preview["revised_procedure_md"] = preview["revised_procedure_md"][:500] + "…"
                    st.json(preview)
        else:
            st.warning("Teamcenter REST/SOA wiring is customer-specific. Use Mock for demo, or implement your Teamcenter endpoints here.")
            st.code("""
# Placeholder integration:
# - Authenticate to Teamcenter
# - Create / find ECR item
# - Update ECR form fields
# - Attach updated diagram PNG as a Dataset
# - Attach procedure pack PDF/MD
# - Start workflow / release
""")

        st.markdown("---")
        st.subheader("Integration outbox log")
        try:
            con = db()
            df = pd.read_sql_query("SELECT system, action, ref_id, status, created_at FROM external_sync_log ORDER BY created_at DESC LIMIT 30", con)
            con.close()
            st.dataframe(df, use_container_width=True, height=240)
        except Exception as e:
            st.error(f"Could not read sync log: {e}")

    with tabs[2]:
        st.subheader("End-to-End: SAP BOM → ECR → Impact → Diagram/Procedure → Teamcenter")
        st.write(
            "Use this guided flow to demo a cohesive tool: pull BOM from SAP, enter an ECR, run BOM-based impact analysis, "
            "generate updated diagram + revised procedure, then push the ECR pack to Teamcenter."
        )

        bom_key = st.session_state.get("sap_bom_key", "IDEAL-5000-TWINSET")
        cached = load_external_bom_latest("SAP", bom_key)

        st.markdown("**Step 1 — SAP BOM cached**")
        if cached is None:
            st.warning("No SAP BOM cached yet. Go to the SAP BOM tab and cache a BOM (Mock or OData).")
        else:
            st.success(f"SAP BOM cached: {len(cached)} rows for {bom_key}")
            st.dataframe(cached.head(20), use_container_width=True, height=230)

        st.markdown("**Step 2 — ECR text**")
        ecr_text = st.text_area("ECR text", value=st.session_state.get("sap_ecr_summary", ""), height=110, key="e2e_ecr_text")

        st.markdown("**Step 3 — Impact analysis (BOM-based)**")
        if st.button("Run impact analysis from SAP BOM + ECR", key="e2e_run_impact"):
            if cached is None:
                st.error("Cache a BOM first.")
            else:
                df = cached.copy()
                # Normalize schema
                if "part_number" not in df.columns:
                    for c in df.columns:
                        if "part" in c.lower() or "component" in c.lower() or "material" in c.lower():
                            df["part_number"] = df[c].astype(str)
                            break
                if "quantity" not in df.columns:
                    df["quantity"] = 1
                if "cost" not in df.columns:
                    df["cost"] = 0.0

                df["part_number"] = df["part_number"].astype(str).fillna("")
                parts = [pn for pn in df["part_number"].unique().tolist() if pn]
                mentioned = sorted({pn for pn in parts if pn in (ecr_text or "")})
                impacted = df[df["part_number"].isin(mentioned)].copy() if mentioned else df.head(0).copy()

                roll = {
                    "bom_rows": int(len(df)),
                    "mentioned_parts_in_ecr": int(len(mentioned)),
                    "impacted_rows": int(len(impacted)),
                    "cost_rollup": float((pd.to_numeric(df["cost"], errors="coerce").fillna(0) * pd.to_numeric(df["quantity"], errors="coerce").fillna(1)).sum()),
                }
                st.session_state["e2e_last_impact"] = {"rollup": roll, "mentioned_parts": mentioned}
                c1, c2 = st.columns(2)
                c1.metric("Impacted rows", roll["impacted_rows"])
                c2.metric("Mentioned parts", roll["mentioned_parts_in_ecr"])
                st.json(roll)
                if len(impacted) > 0:
                    st.dataframe(impacted, use_container_width=True, height=240)
                else:
                    st.info("No explicit part numbers matched in ECR text. Tip: include a part number from the BOM in the ECR.")

        st.markdown("**Step 4 — Check generated artifacts**")
        d_ok = bool(st.session_state.get("dcu_updated_diagram"))
        p_ok = bool(st.session_state.get("dcu_revised_procedure"))
        st.write(f"- Updated diagram available: {'✅' if d_ok else '❌'}")
        st.write(f"- Revised procedure available: {'✅' if p_ok else '❌'}")
        st.caption("Generate these in the 'Diagram Change Update' tab. This page will then include them in the Teamcenter payload.")

        st.markdown("**Step 5 — Queue external updates (demo outbox)**")
        colA, colB = st.columns(2)
        if colA.button("Queue SAP update", key="e2e_queue_sap"):
            sync_id = sap_update_bom_demo(bom_key, {"bom_key": bom_key, "ecr": ecr_text, "impact": st.session_state.get("e2e_last_impact", {})})
            st.success(f"SAP update queued: {sync_id or '(logged)'}")
        if colB.button("Queue Teamcenter ECR update", key="e2e_queue_tc"):
            ecr_id = st.session_state.get("tc_ecr_id", "ECR-DEMO-001")
            sync_id = teamcenter_update_ecr_demo(ecr_id, {
                "ecr_id": ecr_id,
                "ecr": ecr_text,
                "impact": st.session_state.get("e2e_last_impact", {}),
                "diagram_present": d_ok,
                "procedure_present": p_ok,
            })
            st.success(f"Teamcenter update queued: {sync_id or '(logged)'}")

        st.markdown("---")
        st.subheader("Outbox (last 20 integration events)")
        try:
            con = db()
            df = pd.read_sql_query("SELECT system, action, ref_id, status, created_at FROM external_sync_log ORDER BY created_at DESC LIMIT 20", con)
            con.close()
            st.dataframe(df, use_container_width=True, height=240)
        except Exception as e:
            st.error(f"Could not read sync log: {e}")

    with tabs[3]:
        st.subheader("CAD / Evidence upload")
        st.caption("Upload CAD and other engineering evidence to use as context for change packs and Q&A.")
        cad = st.file_uploader(
            "Upload CAD file (STEP/STP/IGES/JT/STL/DWG/DXF/OBJ/SLDPRT/SLDASM)",
            type=CAD_FILE_EXTS,
            key="cad_upload",
        )
        if cad is not None:
            try:
                data = cad.read()
                doc_id = upsert_document(cad.name, data, mime=cad.type or "")
                st.success(f"CAD stored in library: {doc_id[:8]} — {cad.name}")
            except Exception as e:
                st.error(f"Failed to store CAD: {type(e).__name__}: {e}")

def page_ai_assistant(context_mode: str = "library") -> None:
    context_mode = (context_mode or "library").strip().lower()
    page_title("AI Assistant", "Ask contextual questions across your ingested evidence. Answers are traceable.")
    st.session_state.setdefault("ai_chat", [])

    # Light-weight readiness line (fast)
    ai_enabled = openai_client_from_session() is not None
    docs = list_documents()
    st.caption(f"AI enabled: {'Yes' if ai_enabled else 'No'} · Documents: {len(docs)} · Mode: {context_mode}")
    if not ai_enabled:
        st.warning("AI is currently unavailable. Enter your OpenAI key in the sidebar (Settings) to enable GPT answers. The app will still provide deterministic summaries.")


    # Quick upload (optional)
    show_quick = st.checkbox("Quick upload to library (PDF/Word/Excel/CSV/CAD)", value=False, key="ai_quick_upload_show")
    if show_quick:
        st.caption("Upload files to the document library for Q&A. They remain available across tabs and sessions.")
        quick_files = st.file_uploader(
            "Upload files for Q&A",
            type=["pdf", "docx", "xlsx", "xls", "csv", "txt", "json", "md"] + CAD_FILE_EXTS,
            accept_multiple_files=True,
            key="ai_quick_upload",
        )
        if quick_files:
            for f in quick_files:
                try:
                    b = f.getvalue()
                    doc_id = upsert_document(f.name, b, mime=getattr(f, "type", "") or "")
                    st.success(f"Ingested {f.name} → {doc_id}")
                except Exception as e:
                    st.error(f"Failed to ingest {getattr(f,'name','file')}: {e}")
            try:
                st.rerun()
            except Exception:
                try:
                    st.experimental_rerun()
                except Exception:
                    pass

    # Conversation (persisted)
    with st.container():
        st.markdown("### Conversation (persists across tabs)")
        if st.session_state["ai_chat"]:
            for msg in st.session_state["ai_chat"][-20:]:
                role = msg.get("role", "assistant")
                content = msg.get("content", "")
                if role == "user":
                    st.markdown(f"**You:** {content}")
                else:
                    st.markdown(f"**Copilot:** {content}")
        else:
            st.caption("No messages yet. Ask a question below.")

        c_chat1, c_chat2, c_chat3 = st.columns([1, 1, 1])
        if c_chat1.button("Clear chat", key="ai_clear_chat"):
            st.session_state["ai_chat"] = []
            try:
                st.rerun()
            except Exception:
                pass
        if c_chat2.button("Archive session", key="ai_archive_session"):
            arc_id = archive_current_session(st.session_state.get("session_archive_name", "AI session"))
            if arc_id:
                st.success("Session archived.")
            else:
                st.error("Archive failed.")
        if c_chat3.button("Copy last answer", key="ai_copy_last"):
            last = ""
            for msg in reversed(st.session_state["ai_chat"]):
                if msg.get("role") == "assistant":
                    last = msg.get("content", "")
                    break
            if last:
                st.code(last)

    st.markdown("---")

    # Scope / context selection
    scope = st.selectbox("Scope", ["All documents", "Selected document"], index=0, key="ai_scope")
    selected_doc_id: Optional[str] = None
    if scope == "Selected document" and not docs.empty:
        sel = st.selectbox("Document", list(docs["filename"]), key="ai_sel_doc")
        row = docs[docs["filename"] == sel].iloc[0]
        selected_doc_id = str(row["doc_id"])

    q = st.text_area(
        "Question",
        placeholder="e.g., What changed between baseline and modified? Which dimensions/tolerances are impacted? What are acceptance criteria?",
        height=120,
        key="ai_question",
    )

    c_fast1, c_fast2, c_fast3, c_fast4 = st.columns([1, 1, 1, 1])
    with c_fast1:
        fast_mode = st.toggle("Fast mode", value=True, help="Uses fewer context blocks and shorter output.", key="ai_fast_mode")
    with c_fast2:
        max_ctx_blocks = st.slider("Context blocks", min_value=3, max_value=18, value=8, step=1, help="Fewer blocks = faster.", key="ai_ctx_blocks")
    with c_fast3:
        max_chars = st.slider("Chars per block", min_value=300, max_value=2000, value=1000, step=100, help="Lower = faster.", key="ai_chars")
    with c_fast4:
        st.toggle("Streaming", value=True, help="Stream tokens as they arrive (feels faster).", key="ai_stream_answer")

    # Build context lazily only when asking (prevents slow load when switching tabs)
    if st.button("Ask", type="primary", key="ai_ask_btn"):
        if not (q or "").strip():
            st.warning("Please enter a question.")
            return


        # Fast path for tabular docs (CSV/XLSX): answer common aggregation questions instantly
        try:
            if st.session_state.get("ai_scope") == "Selected document" and selected_doc_id:
                _doc = get_document(selected_doc_id)
                if _doc and str(_doc.get("file_ext","")).lower() in [".csv", ".xlsx", ".xls"] and _doc.get("stored_path"):
                    _p = _doc["stored_path"]
                    if os.path.exists(_p):
                        if _p.lower().endswith(".csv"):
                            _df = pd.read_csv(_p, engine="python")
                        else:
                            _df = pd.read_excel(_p)
                        # Normalize common numeric/date columns
                        for _c in _df.columns:
                            if any(k in str(_c).lower() for k in ["date", "time", "effective", "effectivity"]):
                                _df[_c] = pd.to_datetime(_df[_c], errors="coerce")
                        # Heuristic QA
                        _q = (q or "").strip().lower()
                        _ans = ""
                        _fig = None
                        def _pick(cands):
                            for cand in cands:
                                lc = str(cand).lower()
                                for cc in _df.columns:
                                    lcc = str(cc).lower()
                                    if lcc == lc or lc in lcc:
                                        return cc
                            return None
                        if ("average" in _q or "avg" in _q) and "lead" in _q:
                            c = _pick(["lead_time", "leadtime", "lead_time_days", "lead_time_day"])
                            if c is not None:
                                v = pd.to_numeric(_df[c], errors="coerce")
                                _ans = f"Average lead time: **{float(np.nanmean(v)):.2f}** (from column `{c}`)."
                        elif ("total" in _q or "sum" in _q) and "lead" in _q:
                            c = _pick(["lead_time", "leadtime", "lead_time_days", "lead_time_day"])
                            if c is not None:
                                v = pd.to_numeric(_df[c], errors="coerce")
                                _ans = f"Total lead time: **{float(np.nansum(v)):.2f}** (from column `{c}`)."
                        elif ("unique" in _q and "part" in _q):
                            c = _pick(["part_number", "child_part_number", "parent_part_number"])
                            if c is not None:
                                _ans = f"Unique part numbers: **{int(pd.Series(_df[c]).nunique(dropna=True))}** (from `{c}`)."
                        elif ("unique" in _q and "serial" in _q):
                            c = _pick(["serial_number", "serial_no", "serial", "sn"])
                            if c is not None:
                                _ans = f"Unique serial numbers: **{int(pd.Series(_df[c]).nunique(dropna=True))}** (from `{c}`)."
                        elif ("sum" in _q and "extended" in _q and "cost" in _q):
                            c = _pick(["extended_cost_eur", "extended_cost", "extended_cost_euro"])
                            if c is not None:
                                v = pd.to_numeric(_df[c], errors="coerce")
                                _ans = f"Sum extended cost (€): **{float(np.nansum(v)):.2f}** (from `{c}`)."
                        # Simple plot: "plot <metric> over <date>"
                        if (not _fig) and ("plot" in _q or "chart" in _q) and ("over" in _q or "by" in _q):
                            x = _pick(["effectivity_date", "effective_from", "effective_to", "date", "creation_time", "creation date"])
                            y = _pick(["extended_cost_eur", "quantity", "unit_cost_eur", "weight_kg", "lead_time_days", "lead_time"])
                            if x is not None and y is not None and "px" in globals():
                                g = _df.copy()
                                g = g[[x, y]].dropna()
                                if pd.api.types.is_datetime64_any_dtype(g[x]):
                                    g["_x"] = g[x].dt.date
                                else:
                                    g["_x"] = g[x].astype(str)
                                g["_y"] = pd.to_numeric(g[y], errors="coerce")
                                grp = g.groupby("_x")["_y"].sum().reset_index()
                                _fig = px.line(grp, x="_x", y="_y", title=f"Sum({y}) over {x}")
                        if _ans:
                            st.markdown("### Answer")
                            st.markdown(_ans)
                            if _fig is not None:
                                st.plotly_chart(_fig, use_container_width=True)
                            st.session_state["ai_chat"].append({"role": "user", "content": q})
                            st.session_state["ai_chat"].append({"role": "assistant", "content": re.sub(r"\*\*|`", "", _ans)})
                            return
        except Exception:
            pass

        context_blocks: List[Dict[str, str]] = []

        # Special context from Visual Comparison page
        if context_mode == "last_visual":
            regions = st.session_state.get("last_visual_regions", []) or []
            deltas = st.session_state.get("last_visual_attr_deltas", {}) or {}
            context_blocks.append({"id": "visual_attr_deltas", "text": json.dumps(deltas, ensure_ascii=False)[:6000]})
            context_blocks.append({"id": "visual_regions", "text": json.dumps(regions, ensure_ascii=False)[:6000]})
            # Also add the last diff summary if present
            if st.session_state.get("dcu_last_diff"):
                context_blocks.append({"id": "dcu_last_diff", "text": json.dumps(st.session_state.get("dcu_last_diff"), ensure_ascii=False)[:6000]})
        else:
            if docs.empty:
                st.info("Upload documents in 'Ingestion & Library' or via Quick upload to get started.")
                return

            try:
                if st.session_state.get("ai_scope") == "Selected document" and selected_doc_id:
                    chunks = get_document_chunks_cached(selected_doc_id)
                    context_blocks = [{"id": c["chunk_id"], "text": c["text"]} for c in chunks[:90]]
                else:
                    # Recent docs only; ranking chooses the best chunks
                    for _, r in docs.head(8).iterrows():
                        doc_id = str(r["doc_id"])
                        chunks = get_document_chunks_cached(doc_id)
                        context_blocks.extend([{"id": c["chunk_id"], "text": c["text"]} for c in chunks[:4]])
                    context_blocks = context_blocks[:90]
            except Exception:
                context_blocks = []

        if not context_blocks:
            st.warning("No context blocks available. Upload documents first.")
            return

        model = st.session_state.get("openai_model", DEFAULT_MODEL)

        # Rank once for both LLM and fallback narrative
        ranked_for_fallback = rank_context_blocks(q, context_blocks, top_k=max_ctx_blocks, max_chars_per_block=max_chars)

        ans = ""
        ranked = ranked_for_fallback

        if ai_enabled:
            # Streaming output (feels faster) + evidence panel
            ctx = "\n\n".join([f"[{b['id']}]\n{b['text']}" for b in ranked])

            system = (
                "You are an AI engineering assistant. "
                "Answer ONLY using the provided context blocks. "
                "Cite sources using their identifiers in square brackets (e.g., [chk_doc_xxx_0001]). "
                "If the answer is not in context, say you don't know."
            )
            style = "Return a short, direct answer. Use bullets where useful." if fast_mode else "Return a detailed answer with clear sections and bullets."
            user_prompt = f"Question:\n{q}\n\nContext blocks:\n{ctx}\n\n{style}"

            st.markdown("### Answer")
            if st.session_state.get("ai_stream_answer", True):
                with st.spinner("Streaming answer…"):
                    gen, meta, _err = llm_stream_text(
                        system=system,
                        user=user_prompt,
                        temperature=0.2 if fast_mode else 0.25,
                        model=model,
                        max_output_tokens=450 if fast_mode else 900,
                    )
                    try:
                        if hasattr(st, "write_stream"):
                            ans = st.write_stream(gen()) or ""
                        else:
                            ph = st.empty()
                            acc = ""
                            for chunk in gen():
                                acc += chunk
                                ph.markdown(acc)
                            ans = acc
                    except Exception:
                        ans = ""

                    if not ans.strip():
                        ans = (meta.get("text", "") or "").strip()
            else:
                with st.spinner("Generating answer…"):
                    ans, _ev = llm_answer_with_evidence(
                        q,
                        ranked_for_fallback,
                        model=model,
                        fast_mode=fast_mode,
                        max_ctx_blocks=max_ctx_blocks,
                        max_chars_per_block=max_chars,
                    )
                st.write(ans)
        else:
            st.markdown("### Answer (LLM unavailable)")
            ans = deterministic_narrative(ranked_for_fallback, question=q)
            st.write(ans)

        if not (ans or "").strip():
            ans = deterministic_narrative(ranked_for_fallback, question=q)

        # Persist chat (so you don't lose it when switching pages)
        try:
            st.session_state.setdefault("ai_chat", [])
            st.session_state["ai_chat"].append({"role": "user", "content": q})
            st.session_state["ai_chat"].append({"role": "assistant", "content": ans})
        except Exception:
            pass

        # Evidence panel (based on citations present in the final text)
        ev: List[Dict[str, str]] = []
        for b in ranked:
            if f"[{b['id']}]" in ans:
                ev.append({"id": b["id"], "text": b["text"][:260]})
        with st.container():
            st.markdown("### Evidence used")
            st.json(ev)



def page_risk_scenario(identity: Dict[str, str]):
    page_title("Risk & Scenario", "Quick what-if simulation and readiness checks (rule-based; AI optional).")
    changes = list_changes()
    if changes.empty:
        st.info("Create a change first in 'Change Workspace (ECR/ECO)' or 'Diagram Change Update'.")
        return

    change_id = st.selectbox("Select change", list(changes["change_id"]))
    ch = get_change(change_id)
    payload = ch.get("payload", {}) if isinstance(ch, dict) else {}

    st.markdown("#### Change summary")
    st.json(payload)

    st.markdown("#### What-if sliders")
    c1, c2, c3 = st.columns(3)
    with c1:
        schedule_delta = st.slider("Schedule impact (weeks)", 0, 52, 6)
    with c2:
        cost_delta = st.slider("Cost impact (€k)", 0, 2000, 150)
    with c3:
        complexity = st.slider("Complexity", 1, 10, 6)

    risk_score = min(100, int(schedule_delta * 1.4 + cost_delta * 0.03 + complexity * 6))
    st.metric("Scenario risk score", f"{risk_score}/100")

    if risk_score >= 70:
        st.warning("High risk: consider phased release, extra validation, and manufacturing trials.")
    elif risk_score >= 40:
        st.info("Medium risk: validate critical dimensions/tolerances and update work instructions.")
    else:
        st.success("Low risk: proceed with standard review + validation gates.")

    st.markdown("#### Recommended validation gates")
    st.write(
        "- Dimensional inspection plan\n"
        "- Tooling / fixture verification\n"
        "- First Article Inspection (FAI)\n"
        "- Process capability check (Cp/Cpk)\n"
        "- Acceptance criteria sign-off"
    )

def page_export(identity: Dict[str, str]):
    page_title("Export Center", "Download PDFs, procedure packs, and archives generated by the app.")
    ensure_dirs()

    st.markdown("#### Exports")
    ex_files = sorted([f for f in os.listdir(EXPORTS_DIR) if os.path.isfile(os.path.join(EXPORTS_DIR, f))])
    if not ex_files:
        st.info("No exports yet.")
    else:
        for fn in ex_files[::-1][:30]:
            fp = os.path.join(EXPORTS_DIR, fn)
            with open(fp, "rb") as f:
                st.download_button(f"Download {fn}", data=f.read(), file_name=fn)

    st.markdown("#### Archives")
    ar_files = sorted([f for f in os.listdir(ARCHIVES_DIR) if os.path.isfile(os.path.join(ARCHIVES_DIR, f))])
    if not ar_files:
        st.info("No archives yet.")
    else:
        for fn in ar_files[::-1][:30]:
            fp = os.path.join(ARCHIVES_DIR, fn)
            with open(fp, "rb") as f:
                st.download_button(f"Download {fn}", data=f.read(), file_name=fn)

# ----------------------------
# Defensible Moat Showcase (demo)
# ----------------------------
def _db_count(query: str, params: tuple = ()) -> int:
    try:
        con = db()
        cur = con.cursor()
        cur.execute(query, params)
        row = cur.fetchone()
        con.close()
        return int(row[0] if row and row[0] is not None else 0)
    except Exception:
        return 0


def seed_synthetic_change_outcomes(n: int = 120) -> int:
    """
    Demo helper: seed the change_history_training table with synthetic labeled examples
    so the predictor can show measurable performance for a demo without customer data.
    """
    try:
        con = db()
        cur = con.cursor()
        # avoid over-seeding
        existing = _db_count("SELECT COUNT(*) FROM change_history_training")
        if existing >= 300:
            con.close()
            return 0

        import random
        random.seed(42)

        templates = [
            "Increase bracket thickness from {a}mm to {b}mm; replace fasteners with A2 stainless; update torque spec.",
            "Change seal material from EPDM to FKM; add leak test step; update acceptance criteria to <{leak} ml/min.",
            "Relocate sensor harness routing; add strain relief; update work instruction and inspection plan {insp}.",
            "Modify hole diameter from {a}mm to {b}mm; tolerance tighten to ±{tol}mm; update CMM program.",
            "Switch supplier for polymer profile; update RoHS/REACH docs; revalidate fitment on assembly {asm}.",
            "Add chamfer {a}x45°; remove interference on assembly {asm}; update drawing notes and PFMEA."
        ]

        def rand_pn():
            return f"PN-{random.randint(100000,999999)}"

        inserted = 0
        for i in range(n):
            a = random.choice([4,5,6,8,10,12,14,16,18,20])
            b = a + random.choice([1,2,3,4,5])
            leak = random.choice([1.0, 2.5, 5.0, 7.5])
            tol = random.choice([0.05,0.1,0.2,0.3])
            insp = f"INSP-{random.randint(1000,9999)}"
            asm = f"ASM-{random.randint(1000,9999)}"
            t = random.choice(templates).format(a=a, b=b, leak=leak, tol=tol, insp=insp, asm=asm)
            t2 = f"{t} Affected parts: {rand_pn()}, {rand_pn()}, {rand_pn()}."
            # synthetic labels: bigger deltas and tighter tolerances correlate with delay/quality risk
            outcome_delay = 1 if (b-a) >= 4 or ("supplier" in t.lower()) or ("revalidate" in t.lower()) else 0
            outcome_quality = 1 if ("tolerance tighten" in t.lower()) or (tol <= 0.1) or ("leak test" in t.lower()) else 0

            cur.execute(
                "INSERT INTO change_history_training(change_id, text, outcome_delay, outcome_quality_issue, created_at) VALUES (?,?,?,?,?)",
                (f"SYN-{dt.datetime.utcnow().strftime('%Y%m%d')}-{i:04d}", t2, outcome_delay, outcome_quality, dt.datetime.utcnow().isoformat()),
            )
            inserted += 1

        con.commit()
        con.close()
        return inserted
    except Exception:
        return 0


def train_predictor_detailed() -> Optional[Dict[str, Any]]:
    """
    Train the demo predictors and return metrics + top signals for transparency.
    Leaves the original train_predictor() untouched for backwards compatibility.
    """
    if TfidfVectorizer is None or LogisticRegression is None or train_test_split is None:
        return None

    con = db()
    df = pd.read_sql_query("SELECT text, outcome_delay, outcome_quality_issue FROM change_history_training", con)
    con.close()
    if len(df) < 30:
        return None

    X_text = df["text"].fillna("").tolist()
    vec = TfidfVectorizer(stop_words="english", max_features=8000, ngram_range=(1,2))
    X = vec.fit_transform(X_text)

    out: Dict[str, Any] = {"n_samples": int(len(df)), "targets": {}}
    feature_names = np.array(vec.get_feature_names_out())

    for target in ["outcome_delay", "outcome_quality_issue"]:
        y = df[target].fillna(0).astype(int).values
        strat = y if len(np.unique(y)) > 1 else None
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.25, random_state=42, stratify=strat)

        clf = LogisticRegression(max_iter=2000)
        try:
            clf.fit(X_train, y_train)
            proba = clf.predict_proba(X_test)[:, 1]
            auc = float(roc_auc_score(y_test, proba)) if roc_auc_score is not None else 0.0

            coefs = clf.coef_.ravel()
            top_pos_idx = np.argsort(coefs)[-12:][::-1]
            top_neg_idx = np.argsort(coefs)[:12]
            top_pos = [(feature_names[i], float(coefs[i])) for i in top_pos_idx]
            top_neg = [(feature_names[i], float(coefs[i])) for i in top_neg_idx]

            out["targets"][target] = {
                "auc": auc,
                "positive_rate": float(np.mean(y)),
                "top_positive_signals": top_pos,
                "top_negative_signals": top_neg,
            }
        except Exception:
            out["targets"][target] = {
                "auc": 0.0,
                "positive_rate": float(np.mean(y)),
                "top_positive_signals": [],
                "top_negative_signals": [],
            }
    return out


def page_moat_showcase() -> None:
    page_title("Defensible Moat Showcase (Demo)")
    st.caption("Demo-ready ML moat: data network effects + predictive outcome models + multimodal diff intelligence.")

    c1, c2, c3, c4 = st.columns(4)
    n_docs = _db_count("SELECT COUNT(*) FROM documents")
    n_changes = _db_count("SELECT COUNT(*) FROM changes")
    n_chunks = _db_count("SELECT COUNT(*) FROM chunks")
    n_labels = _db_count("SELECT COUNT(*) FROM change_history_training")

    c1.metric("Documents ingested", n_docs)
    c2.metric("Changes created", n_changes)
    c3.metric("Indexed chunks", n_chunks)
    c4.metric("Labeled outcomes", n_labels)

    st.markdown("---")

    left, right = st.columns([1, 1])

    with left:
        st.subheader("1) Data Network Effects")
        st.write(
            "Every ingestion, change request, approval/edit, and actual outcome can be captured as training data. "
            "That creates compounding advantage over time (faster cycle time, better risk prediction, better impact coverage)."
        )

        with st.expander("Add a labeled outcome", expanded=True):
            txt = st.text_area("Change description / ECO narrative", height=130, placeholder="Paste a change narrative here…")
            colA, colB = st.columns(2)
            delay = colA.checkbox("Outcome: delay occurred", value=False)
            qual = colB.checkbox("Outcome: quality issue occurred", value=False)
            if st.button("Add labeled example"):
                if txt.strip():
                    try:
                        con = db()
                        cur = con.cursor()
                        cur.execute(
                            "INSERT INTO change_history_training(change_id, text, outcome_delay, outcome_quality_issue, created_at) VALUES (?,?,?,?,?)",
                            (f"MAN-{dt.datetime.utcnow().strftime('%Y%m%d%H%M%S')}", txt.strip(), int(delay), int(qual), dt.datetime.utcnow().isoformat()),
                        )
                        con.commit()
                        con.close()
                        st.success("Added. Re-train the predictor to see updated metrics.")
                        try:
                            st.rerun()
                        except Exception:
                            pass
                    except Exception as e:
                        st.error(f"Could not write training row: {e}")
                else:
                    st.warning("Please paste a change narrative first.")

        if st.button("Seed synthetic training data (demo)"):
            added = seed_synthetic_change_outcomes(120)
            if added > 0:
                st.success(f"Seeded {added} labeled examples.")
                try:
                    st.rerun()
                except Exception:
                    pass
            else:
                st.info("No seeding needed (already enough examples) or DB write not available.")

    with right:
        st.subheader("2) Predictive Change Outcomes (ML)")
        st.write("Train a lightweight predictor to estimate **Delay Risk** and **Quality Issue Risk** from historical ECO narratives.")
        metrics = train_predictor_detailed()
        if metrics is None:
            st.warning("Predictor unavailable or not enough labeled outcomes. Use seeding for demo.")
        else:
            t1, t2 = st.columns(2)
            auc_delay = metrics["targets"].get("outcome_delay", {}).get("auc", 0.0)
            auc_qual = metrics["targets"].get("outcome_quality_issue", {}).get("auc", 0.0)
            t1.metric("Delay-risk AUC", f"{auc_delay:.2f}")
            t2.metric("Quality-risk AUC", f"{auc_qual:.2f}")

            st.caption("Transparent signals (top model coefficients):")
            tab1, tab2 = st.tabs(["Delay risk signals", "Quality risk signals"])
            with tab1:
                d = metrics["targets"].get("outcome_delay", {})
                st.markdown("**Top positive signals:**")
                st.table(pd.DataFrame(d.get("top_positive_signals", []), columns=["signal", "weight"]).head(12))
                st.markdown("**Top negative signals:**")
                st.table(pd.DataFrame(d.get("top_negative_signals", []), columns=["signal", "weight"]).head(12))
            with tab2:
                q = metrics["targets"].get("outcome_quality_issue", {})
                st.markdown("**Top positive signals:**")
                st.table(pd.DataFrame(q.get("top_positive_signals", []), columns=["signal", "weight"]).head(12))
                st.markdown("**Top negative signals:**")
                st.table(pd.DataFrame(q.get("top_negative_signals", []), columns=["signal", "weight"]).head(12))

    st.markdown("---")
    st.subheader("3) Multimodal Diff Intelligence")
    st.write(
        "Your Visual Diff pipeline already produces quantitative deltas (SSIM, pixel deltas, OCR coverage) and change highlights. "
        "For a demo moat: treat those outputs as measurable signals and show transparent confidence reporting. "
        "For a real moat: collect confirmed deltas over time and train a learned model for dimension/tolerance/annotation classification."
    )
    st.info("Demo flow: run Visual Diff → then ask the AI Assistant contextual questions about the highlighted deltas.")




def page_semantic_trace_diff(identity: Dict[str,str]):
    page_title("Semantic Trace & Diff", "Heuristic feature-level mapping CAD/BIM/DOC assets to BOM parts and requirement-like statements.")

    # Choose a BOM dataframe from UEDM if available, else try latest imported BOM doc
    uedm = st.session_state.get("uedm")
    bom_df = None
    if isinstance(uedm, dict) and isinstance(uedm.get("attributes"), pd.DataFrame):
        # Try reconstruct wide part table
        items = uedm.get("items", pd.DataFrame())
        attrs = uedm.get("attributes", pd.DataFrame())
        if isinstance(items, pd.DataFrame) and isinstance(attrs, pd.DataFrame) and not items.empty and not attrs.empty:
            part_ids = items[items.get("item_type","") == "part"]["item_id"].astype(str).tolist() if "item_id" in items.columns else []
            a_part = attrs[attrs["item_id"].astype(str).isin(part_ids)].copy() if "item_id" in attrs.columns else pd.DataFrame()
            if not a_part.empty and {"item_id","key","value"}.issubset(a_part.columns):
                try:
                    bom_df = a_part.pivot_table(index="item_id", columns="key", values="value", aggfunc="first").reset_index()
                except Exception:
                    bom_df = None

    docs_df = list_documents()
    if docs_df is None or docs_df.empty:
        st.info("No documents ingested yet.")
        return

    # If no BOM in UEDM, attempt load a selected BOM document
    if bom_df is None or (isinstance(bom_df, pd.DataFrame) and bom_df.empty):
        st.warning("No BOM found in UEDM. Select a BOM file from your library to build a mapping.")
        bom_docs = docs_df[docs_df["file_ext"].astype(str).str.lower().isin([".csv",".xlsx",".xls"])].copy()
        if bom_docs.empty:
            st.info("Ingest a BOM CSV/XLSX first.")
            return
        pick = st.selectbox("BOM source", bom_docs["filename"].tolist(), index=0)
        doc_id = str(bom_docs[bom_docs["filename"]==pick].iloc[0]["doc_id"])
        doc = get_document(doc_id)
        pth = str((doc or {}).get("stored_path") or "")
        if not pth or not os.path.exists(pth):
            st.error("Missing BOM file on disk.")
            return
        try:
            bom_df = pd.read_csv(pth) if pth.lower().endswith('.csv') else pd.read_excel(pth)
        except Exception as e:
            st.error(f"Failed to read BOM: {e}")
            return

    st.subheader("1) Semantic mapping (assets -> BOM parts)")
    mapping = semantic_map_assets_to_bom(bom_df, docs_df)
    if mapping.empty:
        st.info("No part-number matches found yet. Tip: include part numbers in filenames or requirement text.")
    else:
        st.dataframe(mapping, use_container_width=True)

    st.subheader("2) Requirement-like statements linked to parts")
    # Pull requirement-like statements from extracted text (simple heuristics)
    req_rows = []
    for _, r in docs_df.iterrows():
        doc_id = str(r.get("doc_id"))
        doc = get_document(doc_id)
        txt = str((doc or {}).get("extracted_text") or "")
        if not txt.strip():
            continue
        # pick lines that look like requirements
        lines = [ln.strip() for ln in txt.splitlines() if len(ln.strip())>20]
        for ln in lines[:500]:
            if any(k in ln.lower() for k in ["shall ","must ","tolerance", "acceptance", "inspection", "verify", "validate"]):
                tokens = extract_part_tokens(ln)
                if tokens:
                    for t in tokens[:5]:
                        req_rows.append({"part_number": t, "statement": ln[:220], "doc_id": doc_id, "filename": str(r.get("filename"))})
    req_df = pd.DataFrame(req_rows)
    if req_df.empty:
        st.info("No requirement-like statements with part tokens found in the ingested text yet.")
    else:
        st.dataframe(req_df.head(300), use_container_width=True)

    st.subheader("3) Semantic delta: pick a part and see linked assets")
    part_candidates = sorted(set(mapping["part_number"].tolist())) if not mapping.empty else []
    if part_candidates:
        part = st.selectbox("Part", part_candidates, index=0)
        st.markdown("**Linked assets**")
        st.dataframe(mapping[mapping["part_number"]==part][["filename","kind","doc_id"]], use_container_width=True)
        st.markdown("**Linked requirement statements**")
        if not req_df.empty:
            st.dataframe(req_df[req_df["part_number"]==part][["statement","filename"]].head(50), use_container_width=True)

def main():
    ensure_dirs()
    init_db()
    inject_theme()

    identity = sidebar_identity()
    session_persistence_sidebar()
    st.sidebar.markdown("---")
    page = st.sidebar.radio("Navigation", [
        "Home",
        "Ingestion & Library",
        "BOM & Graph",
        "BOM Intelligence Dashboard",
        "Unified Data Hub (UEDM)",
        "Change Intelligence Dashboard",
        "Semantic Trace & Diff",
        "Auto KPI Discovery",
        "Snapshot Compare",
        "Geometry & BIM Intelligence",
        "BI Dashboard Builder",
        "Change Workspace (ECR/ECO)",
        "Visual Comparison",
        "Diagram Change Update",
        "AI Assistant",
        "Defensible Moat Showcase (Demo)",
        "Risk & Scenario",
        "Export Center",
        "Connectors (PLM/CAD)",
    ])

    if page == "Home":
        page_home()
    elif page == "Ingestion & Library":
        page_ingest()
    elif page == "BOM & Graph":
        page_bom_graph()
    elif page == "BOM Intelligence Dashboard":
        page_bom_dashboard()
    elif page == "Unified Data Hub (UEDM)":
        page_uedm_hub()
    elif page == "Change Intelligence Dashboard":
        page_change_intelligence_dashboard(identity)
    elif page == "Snapshot Compare":
        page_snapshot_compare(identity)
    elif page == "Auto KPI Discovery":
        page_auto_kpi_discovery(identity)

    elif page == "Semantic Trace & Diff":
        page_semantic_trace_diff(identity)

    elif page == "Geometry & BIM Intelligence":
        page_geometry_bim_intelligence()
    elif page == "BI Dashboard Builder":
        page_bi_dashboard_builder()
    elif page == "Change Workspace (ECR/ECO)":
        page_change_workspace(identity)
    elif page == "Visual Comparison":
        page_visual_diff()
    elif page == "Diagram Change Update":
        page_diagram_change_update()
    elif page == "AI Assistant":
        page_ai_assistant()
    elif page == "Defensible Moat Showcase (Demo)":
        page_moat_showcase()
    elif page == "Risk & Scenario":
        page_risk_scenario(identity)
    elif page == "Export Center":
        page_export(identity)
    elif page == "Connectors (PLM/CAD)":
        page_connectors()
    else:
        page_home()



# -----------------------------
# Auto-discovered KPIs by source type
# -----------------------------

CAD_EXTS = {".stp", ".step", ".iges", ".igs", ".stl", ".obj", ".glb", ".gltf"}
BIM_EXTS = {".ifc"}
BOM_EXTS = {".csv", ".xlsx", ".xls"}
DOC_EXTS = {".pdf", ".docx", ".txt", ".md"}
IMG_EXTS = {".png", ".jpg", ".jpeg", ".webp"}


def infer_source_kind(filename: str, file_ext: str | None = None) -> str:
    ext = (file_ext or os.path.splitext(filename or "")[1] or "").lower().strip()
    if ext in BOM_EXTS:
        return "bom"
    if ext in CAD_EXTS:
        return "cad"
    if ext in BIM_EXTS:
        return "bim"
    if ext in IMG_EXTS:
        return "image"
    if ext in DOC_EXTS:
        return "doc"
    return "other"


def _safe_read_bytes(stored_path: str) -> bytes:
    try:
        with open(stored_path, "rb") as f:
            return f.read()
    except Exception:
        return b""


def compute_doc_kpis(doc: Dict[str, Any]) -> Dict[str, Any]:
    """Lightweight KPIs for documents/images."""
    out: Dict[str, Any] = {}
    txt = str((doc or {}).get("extracted_text") or "")
    out["text_chars"] = len(txt)
    out["text_words"] = len(re.findall(r"\w+", txt))
    out["has_text"] = bool(txt.strip())
    return out


def compute_bim_kpis(doc: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    pth = str((doc or {}).get("stored_path") or "")
    if not pth or not os.path.exists(pth):
        return out
    if ifcopenshell is None:
        # Fallback: rough counts from the raw IFC text
        raw = _safe_read_bytes(pth).decode("utf-8", errors="ignore")
        out["ifc_lines"] = len(raw.splitlines())
        out["ifc_entities_rough"] = len(re.findall(r"^#\d+\s*=\s*IFC", raw, flags=re.M))
        return out
    try:
        model = ifcopenshell.open(pth)
        # Common entity counts
        for ent in ["IfcWall", "IfcSlab", "IfcBeam", "IfcColumn", "IfcDoor", "IfcWindow", "IfcPipeSegment", "IfcFlowFitting", "IfcBuildingStorey"]:
            try:
                out[ent] = len(model.by_type(ent))
            except Exception:
                pass
        try:
            out["ifc_total_entities"] = len(model.by_type("IfcRoot"))
        except Exception:
            pass
    except Exception as e:
        out["bim_error"] = str(e)
    return out


def compute_cad_kpis(doc: Dict[str, Any]) -> Tuple[Dict[str, Any], Optional[Any]]:
    """Returns (metrics, optional_figure)."""
    out: Dict[str, Any] = {}
    fig = None
    pth = str((doc or {}).get("stored_path") or "")
    fn = str((doc or {}).get("filename") or "")
    if not pth or not os.path.exists(pth):
        return out, None
    b = _safe_read_bytes(pth)
    # Reuse existing analyzer (now guarded by TRIMESH availability)
    try:
        metrics, fig = analyze_3d_model_bytes(b, fn)
        if isinstance(metrics, dict):
            out.update(metrics)
    except Exception as e:
        out["cad_error"] = str(e)
    return out, fig


def extract_part_tokens(s: str) -> List[str]:
    """Extract likely part-number tokens from text/filenames."""
    s = (s or "")
    tokens = set()
    # Common patterns like ABC-12345, AIB1234, 123-456-789
    for m in re.findall(r"\b[A-Z]{2,6}[-_]?\d{2,8}\b", s.upper()):
        tokens.add(m)
    for m in re.findall(r"\b\d{3,6}[-_]\d{2,6}[-_]\d{2,6}\b", s):
        tokens.add(m)
    # Alnum long tokens
    for m in re.findall(r"\b[A-Z0-9]{6,}\b", s.upper()):
        if any(ch.isdigit() for ch in m) and any(ch.isalpha() for ch in m):
            tokens.add(m)
    return sorted(tokens)


def semantic_map_assets_to_bom(bom_df: pd.DataFrame, docs_df: pd.DataFrame) -> pd.DataFrame:
    """Heuristic mapping from CAD/BIM/DOC assets to BOM parts via part tokens in filename/text."""
    if bom_df is None or bom_df.empty:
        return pd.DataFrame()
    # Pick a part-number column
    part_col = None
    for c in ["part_number", "part_no", "pn", "item", "item_id"]:
        if c in bom_df.columns:
            part_col = c
            break
    if part_col is None:
        return pd.DataFrame()

    bom_parts = bom_df[[part_col]].dropna().astype(str)
    bom_set = set(bom_parts[part_col].str.upper().tolist())

    rows = []
    for _, r in docs_df.iterrows():
        fn = str(r.get("filename") or "")
        doc_id = str(r.get("doc_id") or "")
        doc = get_document(doc_id)
        txt = str((doc or {}).get("extracted_text") or "")
        tokens = set(extract_part_tokens(fn) + extract_part_tokens(txt))
        hits = sorted([t for t in tokens if t in bom_set])
        if hits:
            for h in hits:
                rows.append({"part_number": h, "doc_id": doc_id, "filename": fn, "kind": infer_source_kind(fn, r.get("file_ext"))})
    return pd.DataFrame(rows)


def page_auto_kpi_discovery(identity: Dict[str,str]):
    page_title("Auto KPI Discovery", "Auto-generate KPIs and previews per source type (BOM vs CAD vs BIM vs docs).")

    docs = list_documents()
    if docs is None or docs.empty:
        st.info("Ingest files first in 'Ingestion & Library'.")
        return

    docs = docs.copy()
    docs["kind"] = [infer_source_kind(r.get("filename",""), r.get("file_ext")) for _, r in docs.iterrows()]

    kind = st.selectbox("Source type", sorted(docs["kind"].unique().tolist()), index=0)
    dfk = docs[docs["kind"] == kind].copy()

    st.caption(f"{len(dfk)} file(s) detected as '{kind}'.")
    pick = st.selectbox("Select file", dfk["filename"].tolist(), index=0)
    row = dfk[dfk["filename"] == pick].iloc[0].to_dict()
    doc = get_document(str(row.get("doc_id")))

    if kind == "bom":
        st.subheader("BOM preview + auto KPIs")
        pth = str((doc or {}).get("stored_path") or "")
        if not pth or not os.path.exists(pth):
            st.warning("Missing stored file path.")
            return
        try:
            if pth.lower().endswith('.csv'):
                bdf = pd.read_csv(pth)
            else:
                bdf = pd.read_excel(pth)
        except Exception as e:
            st.error(f"Failed to read BOM: {e}")
            return

        st.dataframe(bdf.head(200), use_container_width=True)

        # Auto KPI discovery
        num_cols = [c for c in bdf.columns if pd.api.types.is_numeric_dtype(bdf[c]) or c.lower() in ["quantity","unit_cost_eur","extended_cost_eur","weight_kg","lead_time_days"]]
        cat_cols = [c for c in bdf.columns if c not in num_cols]

        # basic cards
        c1,c2,c3,c4 = st.columns(4)
        # part count
        part_col = next((c for c in ["part_number","part_no","pn","item"] if c in bdf.columns), None)
        serial_col = next((c for c in ["serial_number","serial_no","sn"] if c in bdf.columns), None)
        c1.metric("Unique parts", int(bdf[part_col].nunique()) if part_col else len(bdf))
        c2.metric("Unique serials", int(bdf[serial_col].nunique()) if serial_col else 0)
        qty_col = next((c for c in ["quantity","bom_quantity","qty"] if c in bdf.columns), None)
        c3.metric("Sum Qty", f"{pd.to_numeric(bdf[qty_col], errors='coerce').fillna(0).sum():,.0f}" if qty_col else "—")
        cost_col = next((c for c in ["extended_cost_eur","extended_cost","cost"] if c in bdf.columns), None)
        c4.metric("Sum Ext Cost (€)", f"{pd.to_numeric(bdf[cost_col], errors='coerce').fillna(0).sum():,.2f}" if cost_col else "—")

        # Effectivity line if available
        eff = next((c for c in ["effectivity_date","effective_from","created_at","creation_time"] if c in bdf.columns), None)
        if eff and cost_col:
            st.subheader("Effectivity timeline")
            tmp = bdf.copy()
            tmp[eff] = _as_datetime_series(tmp[eff])
            tmp[cost_col] = pd.to_numeric(tmp[cost_col], errors='coerce')
            tmp = tmp[tmp[eff].notna()].copy()
            if not tmp.empty:
                grp = tmp.groupby(tmp[eff].dt.date)[cost_col].sum().reset_index()
                grp.columns=["date","sum"]
                st.plotly_chart(px.line(grp, x="date", y="sum"), use_container_width=True)

    elif kind in ["cad","bim"]:
        st.subheader("Model KPIs")
        if kind == "cad":
            metrics, fig = compute_cad_kpis(doc or {})
            st.json(metrics)
            if fig is not None:
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("3D preview requires 'trimesh' + plotly. Metrics will still work if trimesh is installed.")
        else:
            metrics = compute_bim_kpis(doc or {})
            st.json(metrics)
            if "bim_error" in metrics:
                st.warning(metrics["bim_error"])

    else:
        st.subheader("Document KPIs")
        metrics = compute_doc_kpis(doc or {})
        st.json(metrics)
        # Show extracted text preview
        txt = str((doc or {}).get("extracted_text") or "")
        if txt:
            st.text_area("Extracted text (preview)", txt[:5000], height=260)




if __name__ == "__main__":
    main()
